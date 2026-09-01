# -*- coding: utf-8 -*-
"""
scraper_engine.py — движок парсера goszakup.gov.kz, вынесенный из
исходного finder.py в библиотечный модуль для веб-интерфейса (app.py).

Логика парсинга/классификации тендеров не менялась по существу. Изменения
по сравнению с finder.py:

  1. Все ключевые слова (strong/medium/weak/negative/context) теперь
     ЦЕЛИКОМ хранятся в it_keywords.json и редактируются через веб-интерфейс
     (вкладка "Ключевые слова") — руками JSON лучше не трогать. Если файла
     нет при первом запуске, он создаётся из встроенных значений по
     умолчанию (см. _DEFAULT_KEYWORDS ниже).
  2. print(...) заменён на log(...) — сообщения складываются в потокобезопасный
     буфер LOG_BUFFER и отдаются во фронтенд через /api/run/logs (polling).
  3. Интерактивные input()/sys.exit() убраны — вместо этогоState-machine
     (класс ScraperRunner), которым управляет app.py по кнопкам в браузере:
     launch_browser() -> find_tab() -> start_run() -> stop_run().
  4. Остальной код (parse_list_page, parse_detail_page, classify_is_it,
     ExcelSink, чекпоинт и т.д.) перенесён без изменений в поведении.
"""

import json
import os
import random
import re
import subprocess
import sys
import threading
import time
import traceback
import uuid
from datetime import date
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode, urljoin

import requests
from bs4 import BeautifulSoup
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter

try:
    from selenium import webdriver
    from selenium.webdriver.edge.options import Options as EdgeOptions
    from selenium.webdriver.edge.service import Service as EdgeService
    SELENIUM_AVAILABLE = True
except Exception:
    SELENIUM_AVAILABLE = False

try:
    import anthropic
    ANTHROPIC_AVAILABLE = True
except Exception:
    ANTHROPIC_AVAILABLE = False

BASE_DIR = os.path.dirname(os.path.abspath(__file__))

# ============================================================================
# ============================ ЛОГ ДЛЯ ВЕБ-ИНТЕРФЕЙСА ========================
# ============================================================================

_log_lock = threading.Lock()
LOG_BUFFER = []          # list of {"id": int, "ts": str, "msg": str, "level": str}
_log_counter = 0
MAX_LOG_LINES = 5000


def _classify_log_level(msg: str) -> str:
    """Определяет 'смысл' строки лога для дружелюбного отображения в
    интерфейсе (иконка/цвет) - без необходимости трогать каждый из
    полусотни вызовов log() по отдельности. Основано на реально
    используемых в коде префиксах/фразах."""
    m = msg.strip()
    if not m:
        return "info"
    if set(m) == {"="}:
        return "divider"
    if m.startswith("[!]") or "Ошибка" in m or "ошибка" in m or m.startswith("Traceback") or 'File "' in m:
        return "error"
    if m.startswith("[IT"):
        return "success"
    if m.startswith("[= дубль]"):
        return "neutral"
    if (m.startswith("[авто-перезапуск]") or "Автоматический перезапуск" in m
            or "Останавлив" in m or "остановлен" in m.lower()):
        return "warning"
    return "info"


def log(msg: str):
    global _log_counter
    line = str(msg)
    with _log_lock:
        _log_counter += 1
        LOG_BUFFER.append({
            "id": _log_counter,
            "ts": time.strftime("%H:%M:%S"),
            "msg": line,
            "level": _classify_log_level(line),
        })
        if len(LOG_BUFFER) > MAX_LOG_LINES:
            del LOG_BUFFER[: len(LOG_BUFFER) - MAX_LOG_LINES]
    print(line)


def get_log_since(since_id: int = 0):
    with _log_lock:
        return [l for l in LOG_BUFFER if l["id"] > since_id]


# ============================================================================
# ================================ НАСТРОЙКИ ================================
# ============================================================================

EDGE_DEBUG_ADDRESS = "127.0.0.1:9222"
EDGE_DRIVER_PATH = os.path.join(BASE_DIR, "msedgedriver.exe")
EDGE_BROWSER_PATH = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
AUTO_LAUNCH_EDGE = True

AUTOMATION_PROFILE_DIR = os.path.join(BASE_DIR, ".edge_automation_profile")

COUNT_RECORD_PER_PAGE = 2000
MAX_PAGES_SAFETY_LIMIT = 2000
EMPTY_PAGE_MAX_RETRIES = 3
REQUEST_DELAY_SEC = (0.3, 0.7)
DETAIL_PAGE_DELAY_SEC = (0.15, 0.4)
REQUEST_TIMEOUT_SEC = 20
DETAIL_WORKERS = 4

SCAN_ATTACHED_DOCUMENTS = True
MAX_DOCS_TO_SCAN_IF_NO_MATCH = 3
MAX_DOCUMENT_SIZE_MB = 25
DOWNLOAD_TMP_DIR = os.path.join(BASE_DIR, "tmp_downloads")
DELETE_DOWNLOADED_DOCS_AFTER_SCAN = True

XLSX_SAVE_EVERY_N_MATCHES = 5

OUTPUT_XLSX = os.path.join(BASE_DIR, "goszakup_it_tenders.xlsx")
OUTPUT_XLSX_ASTANA = os.path.join(BASE_DIR, "goszakup_it_tenders_astana.xlsx")
CHECKPOINT_FILE = os.path.join(BASE_DIR, "checkpoint_seen_ids.json")
VIEWED_FILE = os.path.join(BASE_DIR, "viewed_tenders.json")
RUN_STATS_FILE = os.path.join(BASE_DIR, "run_stats.json")
FAVORITES_FILE = os.path.join(BASE_DIR, "favorites.json")
FAVORITES_XLSX = os.path.join(BASE_DIR, "izbrannye_tendery.xlsx")
ERROR_LOG_FILE = os.path.join(BASE_DIR, "errors.log")
PROFILES_FILE = os.path.join(BASE_DIR, "keyword_profiles.json")
KEYWORDS_FILE = os.path.join(BASE_DIR, "it_keywords.json")  # старый формат - только для миграции
BORDERLINE_LOG_FILE = os.path.join(BASE_DIR, "borderline_candidates.log")

DEBUG_MODE = False
DEBUG_DUMP_DIR = os.path.join(BASE_DIR, "debug_html")

IT_SCORE_THRESHOLD = 12
BORDERLINE_SCORE_MIN = 6

# ============================================================================
# =================== КЛЮЧЕВЫЕ СЛОВА: ПРОФИЛИ ТЕМАТИК =========================
# ============================================================================
# Раньше был один плоский набор из 5 категорий (it_keywords.json) - под одну
# тематику (IT). Теперь ключевые слова разбиты на именованные ПРОФИЛИ
# (например 'IT' и 'Проектирование (AutoCAD)'), каждый со своими 5
# категориями. Можно включать НЕСКОЛЬКО профилей одновременно - тогда их
# списки объединяются (без дублей) и классификатор ищет тендеры сразу по
# всем включённым тематикам. Управляется с вкладки "Ключевые слова" в
# интерфейсе (выбор активных профилей + редактирование текста внутри
# профиля), хранится в keyword_profiles.json.

_DEFAULT_KEYWORDS = {
    "strong": [
        "разработка программного обеспечения", "разработка информационной системы",
        "информационная система", "автоматизация", "внедрение информационной системы",
        "доработка по", "модернизация информационной системы", "crm", "erp", "bpm", "api",
        "веб-портал", "веб-приложени", "мобильное приложени", "sql", "postgresql", "mysql",
        "oracle", "база данных", "1с", "1c", "кибербезопасност", "информационная безопасност",
        "эцп", "искусственный интеллект", "машинное обучение", "нейросет", "цифровизация",
        "цифровая трансформация", "цифровая платформа", "цифровая экосистема",
        "бағдарламалық қамтамасыз етуді әзірлеу", "ақпараттық жүйені әзірлеу",
        "ақпараттық жүйе", "автоматтандыру", "ақпараттық жүйені енгізу",
        "ақпараттық жүйені жаңғырту", "веб қосымша", "мобильді қосымша", "дерекқор",
        "киберқауіпсіздік", "ақпараттық қауіпсіздік", "эцқ", "жасанды интеллект",
        "машиналық оқыту", "нейрондық желі", "нейрожелі", "цифрландыру",
        "цифрлық трансформация", "цифрлық платформа", "цифрлық экожүйе",
    ],
    "medium": [
        "программн", "программное обеспечение", "сопровождение по",
        "техническая поддержка по", "сервер", "серверное оборудование",
        "сетевое оборудование", "виртуализац", "облачн", "cloud", "хостинг", "домен",
        "firewall", "helpdesk", "программист", "системный администратор",
        "бағдарламалық", "бағдарламалық жасақтама", "бқ бойынша сүйемелдеу",
        "техникалық қолдау", "серверлік жабдық", "желілік жабдық", "виртуалдандыру",
        "бұлттық", "бағдарламашы", "жүйелік әкімші",
    ],
    "weak": [
        "лицензи", "информационн", "цифров", "электронн", "экосистем",
        "лицензия", "ақпараттық", "цифрлық", "электрондық", "экожүйе",
    ],
    "negative": [
        "мебель", "стол", "стул", "шкаф", "дверь", "окно", "жалюзи", "строительство",
        "ремонт здания", "капитальный ремонт", "асфальт", "бетон", "кирпич", "труба",
        "водоснабжение", "канализация", "отопление", "спецодежда", "бумага",
        "канцелярские товары", "хозтовары",
        "жиһаз", "үстел", "орындық", "сөре", "есік", "терезе", "құрылыс",
        "күрделі жөндеу", "кірпіш", "құбыр", "сумен жабдықтау", "кәріз", "жылыту",
        "арнайы киім", "қағаз", "кеңсе тауарлары", "шаруашылық тауарлары",
    ],
    "context": [
        "техническое задание", "разработка информационной системы",
        "создание информационной системы", "оказание ит услуг",
        "оказание услуг по сопровождению", "оказание услуг по разработке",
        "внедрение программного обеспечения", "модернизация информационной системы",
        "автоматизация процессов", "цифровую экосистему", "цифровой трансформации",
        "цифровой экосистемы",
        "техникалық тапсырма", "ақпараттық жүйені әзірлеу", "ақпараттық жүйені құру",
        "ат қызметтерін көрсету", "сүйемелдеу бойынша қызметтер көрсету",
        "әзірлеу бойынша қызметтер көрсету", "бағдарламалық қамтамасыз етуді енгізу",
        "ақпараттық жүйені жаңғырту", "процестерді автоматтандыру",
        "цифрлық экожүйесін дамыту", "цифрлық білім беру экожүйесі",
        "цифрлық трансформациясы",
    ],
}

# Профиль "Проектирование (AutoCAD)" - проектно-изыскательские работы любого
# масштаба: от эскизного проекта до рабочей документации, генплана,
# инженерных изысканий и экспертизы. Не привязан к конкретному ПО (AutoCAD
# упомянут явно, но так же ловит Civil 3D, Revit, ArchiCAD, Renga, nanoCAD,
# КОМПАС - тендеры редко называют софт по имени, обычно пишут вид работ).
_AUTOCAD_KEYWORDS = {
    "strong": [
        "проектно-изыскательские работы", "пир", "проектно-сметная документация", "псд",
        "разработка псд", "разработка проектно-сметной документации",
        "рабочая документация", "рабочий проект", "эскизный проект", "технический проект",
        "архитектурно-строительное проектирование", "архитектурное проектирование",
        "генеральный план", "генплан", "проект планировки территории", "ппт",
        "проект межевания территории", "инженерные изыскания",
        "инженерно-геологические изыскания", "инженерно-геодезические изыскания",
        "топографическая съемка", "топосъемка", "проект организации строительства", "пос",
        "государственная экспертиза", "негосударственная экспертиза",
        "экспертиза проекта", "экспертное заключение",
        "технико-экономическое обоснование", "тэо",
        "autocad", "autodesk", "civil 3d", "revit", "archicad", "renga", "nanocad",
        "компас-3d", "microstation", "bim", "тим",
        "информационное моделирование", "технологии информационного моделирования",
        "проектирование реконструкции", "вертикальная планировка",
        "жобалау", "жобалық-сметалық құжаттама", "жсқ",
        "жұмыс жобасы", "эскиздік жоба", "техникалық жоба",
        "сәулет-құрылыс жобалау", "бас жоспар",
        "инженерлік іздестіру", "инженерлік-геологиялық іздестіру",
        "инженерлік-геодезиялық іздестіру", "топографиялық түсіру",
        "жобаға сараптама", "мемлекеттік сараптама",
        "ақпараттық модельдеу технологиялары",
    ],
    "medium": [
        "проектирование", "проектные работы", "разработка проекта", "разработка чертежей",
        "чертежи", "разделы проекта", "архитектурные решения", "конструктивные решения",
        "стройгенплан", "строительный генеральный план",
        "сметная документация", "объектная смета", "локальная смета", "сметчик",
        "план, разрез, фасад", "благоустройство территории",
        "жобалау жұмыстары", "жобаны әзірлеу", "сызбаларды әзірлеу",
        "сметалық құжаттама", "аумақты абаттандыру",
    ],
    "weak": [
        "проект", "чертёж", "смета", "изыскания", "планировка",
        "жоба", "сызба", "смета", "жоспарлау",
    ],
    "negative": [
        "поставка мебели", "поставка товаров", "охранные услуги", "услуги связи",
        "клининговые услуги", "уборка помещений", "канцелярские товары",
        "текущий ремонт", "поставка стройматериалов", "аренда техники",
        "жиһаз жеткізу", "күзет қызметтері", "байланыс қызметтері",
        "үй-жайларды тазалау", "ағымдағы жөндеу",
    ],
    "context": [
        "оказание услуг по разработке проектно-сметной документации",
        "разработка рабочего проекта", "выполнение проектно-изыскательских работ",
        "техническое задание на проектирование", "разработка генерального плана",
        "проведение инженерных изысканий", "прохождение государственной экспертизы",
        "жобалық-сметалық құжаттаманы әзірлеу бойынша қызметтер көрсету",
        "жұмыс жобасын әзірлеу", "жобалық-іздестіру жұмыстарын орындау",
        "жобалауға арналған техникалық тапсырма",
    ],
}

# Профиль "Дизайн интерьера" - узкая тематика: дизайн интерьеров и
# внутренних помещений, рабочие чертежи по дизайн-проекту, 3D-визуализация,
# комплектация (мебель/материалы/оборудование по дизайн-проекту),
# авторский надзор дизайнера. Специально уже - чтобы не путаться с общим
# "Проектирование (AutoCAD)" (то больше про АС/КР/инженерные разделы и
# генпланы, это - именно про интерьеры).
_INTERIOR_DESIGN_KEYWORDS = {
    "strong": [
        "дизайн интерьера", "дизайн-проект интерьера", "дизайн-проект помещения",
        "дизайн внутренних помещений", "дизайн помещений", "дизайн внутренней отделки",
        "интерьерные решения", "разработка дизайн-проекта",
        "3d визуализация", "3д визуализация", "3d-визуализация", "3д-визуализация",
        "визуализация интерьера", "фотореалистичная визуализация",
        "комплектация мебели", "комплектация мебелью", "комплектация оборудования",
        "комплектация материалов", "комплектация помещений",
        "авторский надзор дизайнера", "шеф-надзор дизайнера", "шеф-монтаж",
        "эскизный дизайн-проект", "концепция интерьера", "рабочие чертежи по дизайн-проекту",
        "интерьер дизайны", "ішкі жасау дизайны", "ішкі кеңістік дизайны",
        "жиһаздандыру", "авторлық қадағалау", "3d визуализациясы",
    ],
    "medium": [
        "дизайн-проект", "рабочие чертежи", "авторский надзор", "комплектация",
        "меблировка", "декоративная отделка", "отделочные материалы подбор",
        "жиһаздандыру жобасы", "сәндік әрлеу",
    ],
    "weak": [
        "интерьер", "визуализация", "меблировка", "отделка",
        "интерьер", "әрлеу",
    ],
    "negative": [
        "поставка канцелярской мебели", "охранные услуги", "услуги связи",
        "клининговые услуги", "уборка помещений", "капитальный ремонт кровли",
        "текущий ремонт без дизайна", "поставка стройматериалов",
        "күзет қызметтері", "байланыс қызметтері", "үй-жайларды тазалау",
    ],
    "context": [
        "разработка дизайн-проекта интерьера", "выполнение дизайн-проекта с комплектацией",
        "услуги по дизайну интерьера и авторскому надзору",
        "3д визуализация интерьера", "дизайн-проект с рабочими чертежами",
        "комплектация мебелью и оборудованием по дизайн-проекту",
        "интерьер дизайн-жобасын әзірлеу", "жиһаздандыру бойынша қызметтер көрсету",
    ],
}

_kw_lock = threading.Lock()
IT_STRONG_KEYWORDS = []
IT_MEDIUM_KEYWORDS = []
IT_WEAK_KEYWORDS = []
IT_NEGATIVE_KEYWORDS = []
IT_CONTEXT_PHRASES = []

# Встроенные профили - при появлении нового встроенного профиля в новой
# версии приложения он будет автоматически добавлен (и включён) даже в уже
# существующий keyword_profiles.json пользователя, см. _ensure_builtin_profiles().
_BUILTIN_PROFILES = {
    "IT": _DEFAULT_KEYWORDS,
    "Проектирование (AutoCAD)": _AUTOCAD_KEYWORDS,
    "Дизайн интерьера": _INTERIOR_DESIGN_KEYWORDS,
}
_BUILTIN_ACTIVE_BY_DEFAULT = {"IT", "Проектирование (AutoCAD)", "Дизайн интерьера"}


def _ensure_builtin_profiles(payload: dict) -> bool:
    """Дополняет payload отсутствующими встроенными профилями. Возвращает
    True, если что-то реально добавили (тогда файл на диске нужно
    пересохранить) - используется и для полностью нового файла, и для
    донасыщения уже существующего файла новыми встроенными профилями,
    появившимися в более новой версии приложения."""
    changed = False
    for name, defaults in _BUILTIN_PROFILES.items():
        if name not in payload["profiles"]:
            payload["profiles"][name] = _clean_profile(defaults)
            if name in _BUILTIN_ACTIVE_BY_DEFAULT and name not in payload["active"]:
                payload["active"].append(name)
            changed = True
    return changed

_profiles_lock = threading.Lock()
_profiles_cache = None
_KW_CATEGORIES = ("strong", "medium", "weak", "negative", "context")


def _clean_profile(data: dict) -> dict:
    return {cat: list(data.get(cat, [])) for cat in _KW_CATEGORIES}


def _default_profiles_payload() -> dict:
    payload = {"active": [], "profiles": {}}
    _ensure_builtin_profiles(payload)
    return payload


def _load_profiles_from_disk() -> dict:
    if os.path.exists(PROFILES_FILE):
        try:
            with open(PROFILES_FILE, "r", encoding="utf-8") as f:
                data = json.load(f)
            if isinstance(data.get("profiles"), dict) and isinstance(data.get("active"), list):
                if _ensure_builtin_profiles(data):
                    _save_profiles_to_disk(data)
                return data
        except Exception:
            pass
    # Миграция со старого однопрофильного it_keywords.json, если он есть -
    # его содержимое становится профилем "IT", плюс сразу добавляются
    # остальные встроенные профили (см. _BUILTIN_ACTIVE_BY_DEFAULT).
    if os.path.exists(KEYWORDS_FILE):
        try:
            with open(KEYWORDS_FILE, "r", encoding="utf-8") as f:
                old = json.load(f)
            payload = {"active": ["IT"], "profiles": {"IT": _clean_profile(old)}}
            _ensure_builtin_profiles(payload)
            _save_profiles_to_disk(payload)
            return payload
        except Exception:
            pass
    payload = _default_profiles_payload()
    _save_profiles_to_disk(payload)
    return payload


def _save_profiles_to_disk(payload: dict):
    with open(PROFILES_FILE, "w", encoding="utf-8") as f:
        json.dump(payload, f, ensure_ascii=False, indent=2)


def get_profiles_state() -> dict:
    """Полное состояние: {"active": [...], "profiles": {name: {5 категорий}}}."""
    global _profiles_cache
    with _profiles_lock:
        if _profiles_cache is None:
            _profiles_cache = _load_profiles_from_disk()
        return json.loads(json.dumps(_profiles_cache))


def _apply_active_keywords():
    """Пересчитывает глобальные IT_STRONG_KEYWORDS/... как объединение (без
    дублей, без учёта регистра) одноимённых категорий из ВСЕХ активных
    профилей. classify_is_it()/detect_it_type()/detect_priority() продолжают
    читать эти же глобальные списки, как и раньше - им не важно, из скольких
    профилей они собраны."""
    global IT_STRONG_KEYWORDS, IT_MEDIUM_KEYWORDS, IT_WEAK_KEYWORDS
    global IT_NEGATIVE_KEYWORDS, IT_CONTEXT_PHRASES

    state = get_profiles_state()
    merged = {cat: [] for cat in _KW_CATEGORIES}
    seen = {cat: set() for cat in _KW_CATEGORIES}
    for name in state["active"]:
        prof = state["profiles"].get(name)
        if not prof:
            continue
        for cat in _KW_CATEGORIES:
            for kw in prof.get(cat, []):
                low = kw.lower()
                if low not in seen[cat]:
                    seen[cat].add(low)
                    merged[cat].append(kw)

    with _kw_lock:
        IT_STRONG_KEYWORDS = merged["strong"]
        IT_MEDIUM_KEYWORDS = merged["medium"]
        IT_WEAK_KEYWORDS = merged["weak"]
        IT_NEGATIVE_KEYWORDS = merged["negative"]
        IT_CONTEXT_PHRASES = merged["context"]


def save_profile(name: str, keywords: dict) -> dict:
    """Создаёт профиль (если его не было) или полностью перезаписывает его
    5 категорий. Активность профиля этим не меняется - см. set_active_profiles()."""
    global _profiles_cache
    name = (name or "").strip()
    if not name:
        raise ValueError("Пустое имя профиля")
    with _profiles_lock:
        if _profiles_cache is None:
            _profiles_cache = _load_profiles_from_disk()
        _profiles_cache["profiles"][name] = _clean_profile(keywords)
        _save_profiles_to_disk(_profiles_cache)
    _apply_active_keywords()
    return get_profiles_state()


def delete_profile(name: str) -> dict:
    global _profiles_cache
    with _profiles_lock:
        if _profiles_cache is None:
            _profiles_cache = _load_profiles_from_disk()
        _profiles_cache["profiles"].pop(name, None)
        if name in _profiles_cache["active"]:
            _profiles_cache["active"].remove(name)
        _save_profiles_to_disk(_profiles_cache)
    _apply_active_keywords()
    return get_profiles_state()


def rename_profile(old_name: str, new_name: str) -> dict:
    global _profiles_cache
    new_name = (new_name or "").strip()
    if not new_name:
        raise ValueError("Пустое имя профиля")
    with _profiles_lock:
        if _profiles_cache is None:
            _profiles_cache = _load_profiles_from_disk()
        if old_name not in _profiles_cache["profiles"] or old_name == new_name:
            return get_profiles_state()
        _profiles_cache["profiles"][new_name] = _profiles_cache["profiles"].pop(old_name)
        if old_name in _profiles_cache["active"]:
            _profiles_cache["active"] = [
                new_name if n == old_name else n for n in _profiles_cache["active"]
            ]
        _save_profiles_to_disk(_profiles_cache)
    _apply_active_keywords()
    return get_profiles_state()


def set_active_profiles(names) -> dict:
    global _profiles_cache
    with _profiles_lock:
        if _profiles_cache is None:
            _profiles_cache = _load_profiles_from_disk()
        valid = [n for n in names if n in _profiles_cache["profiles"]]
        _profiles_cache["active"] = valid
        _save_profiles_to_disk(_profiles_cache)
    _apply_active_keywords()
    return get_profiles_state()


_apply_active_keywords()

# ============================================================================
# ==================== ИИ-ПРОВЕРКА РЕЛЕВАНТНОСТИ (опционально) ===============
# ============================================================================
# Ключевые слова (даже с профилями) — это сопоставление подстрок: они не
# понимают падежи, синонимы и непривычные формулировки ("дизайн-проектА
# интерьера" уже не совпадёт с фразой "дизайн-проект интерьера" из-за
# окончания). Обучать отдельную ML-модель под это избыточно — правильный
# инструмент здесь: сам Claude, которому просто на человеческом языке
# описывают критерий, и он смотрит на КАЖДЫЙ конкретный тендер осмысленно,
# а не по буквам.
#
# Работает как ДОПОЛНИТЕЛЬНЫЙ слой поверх ключевых слов, а не замена:
#   - режим "borderline" (по умолчанию) - ИИ дозывается только для тендеров,
#     набравших заметный, но недостаточный балл (та же зона, что раньше шла
#     в borderline_candidates.log) - дёшево, ловит то, что ключевые слова
#     почти нашли, но не дотянули из-за формулировки.
#   - режим "all" - ИИ дополнительно перепроверяет и то, что уже прошло по
#     ключевым словам, и может отклонить ложные срабатывания. Дороже
#     (больше вызовов API), но точнее.
#
# ДВА ПРОВАЙДЕРА НА ВЫБОР:
#   - "gemini" (по умолчанию) - у Google Gemini API есть по-настоящему
#     бесплатный тариф (модели Flash/Flash-Lite через Google AI Studio,
#     без привязки карты, без истечения - см. https://aistudio.google.com).
#     ВАЖНО: это не то же самое, что бесплатный лимит сообщений в веб-чате
#     chatgpt.com/gemini.google.com - там ограничения только для человека,
#     печатающего в браузере, к программным вызовам API это не относится.
#     У ChatGPT (OpenAI) такого постоянного бесплатного тарифа для API нет -
#     только один небольшой пробный кредит на новый аккаунт, дальше как и у
#     Anthropic - платите за каждый токен.
#   - "anthropic" (Claude, платно) - оставлен как опция, если уже есть
#     платный доступ и хочется более сильную модель.

AI_SETTINGS_FILE = os.path.join(BASE_DIR, "ai_settings.json")

DEFAULT_AI_CRITERIA = (
    "Тендеры, связанные с дизайном интерьеров и внутренних помещений: "
    "разработка дизайн-проекта, рабочие чертежи по дизайн-проекту, "
    "3D-визуализация интерьера, комплектация (мебель/материалы/оборудование "
    "по дизайн-проекту), авторский надзор дизайнера."
)

_ai_settings_lock = threading.Lock()
_ai_settings_cache = None


def _default_ai_settings() -> dict:
    return {
        "enabled": False,
        "provider": "gemini",           # "gemini" (бесплатно) | "anthropic" (платно)
        "model": "gemini-2.5-flash",
        "criteria": DEFAULT_AI_CRITERIA,
        "verify_mode": "borderline",     # "borderline" | "all"
        "gemini_api_key": "",
        "anthropic_api_key": "",
    }


_PROVIDER_DEFAULT_MODELS = {
    "gemini": "gemini-2.5-flash",
    "anthropic": "claude-sonnet-5",
}


def _load_ai_settings_from_disk() -> dict:
    defaults = _default_ai_settings()
    if not os.path.exists(AI_SETTINGS_FILE):
        return defaults
    try:
        with open(AI_SETTINGS_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        defaults.update({k: v for k, v in data.items() if k in defaults})
        return defaults
    except Exception:
        return defaults


def _save_ai_settings_to_disk(data: dict):
    with open(AI_SETTINGS_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)


def get_ai_settings() -> dict:
    global _ai_settings_cache
    with _ai_settings_lock:
        if _ai_settings_cache is None:
            _ai_settings_cache = _load_ai_settings_from_disk()
        return dict(_ai_settings_cache)


def save_ai_settings(data: dict) -> dict:
    global _ai_settings_cache, _anthropic_client
    with _ai_settings_lock:
        current = _ai_settings_cache or _load_ai_settings_from_disk()
        merged = dict(current)
        merged["enabled"] = bool(data.get("enabled", current["enabled"]))
        provider = data.get("provider", current["provider"])
        merged["provider"] = provider if provider in ("gemini", "anthropic") else current["provider"]
        requested_model = str(data.get("model") or "").strip()
        merged["model"] = requested_model or _PROVIDER_DEFAULT_MODELS.get(merged["provider"], current["model"])
        merged["criteria"] = str(data.get("criteria", current["criteria"]))
        verify_mode = data.get("verify_mode", current["verify_mode"])
        merged["verify_mode"] = verify_mode if verify_mode in ("borderline", "all") else current["verify_mode"]
        # Ключи, вставленные прямо в интерфейс - хранятся ТОЛЬКО в
        # ai_settings.json (этот файл в .gitignore, в репозиторий никогда
        # не попадёт), никогда не пишутся в исходный код. Если поле в
        # запросе не передано - оставляем то, что уже было сохранено
        # раньше (чтобы сохранение других настроек не стирало ключ).
        if "gemini_api_key" in data:
            merged["gemini_api_key"] = str(data.get("gemini_api_key") or "").strip()
        if "anthropic_api_key" in data:
            merged["anthropic_api_key"] = str(data.get("anthropic_api_key") or "").strip()
            _anthropic_client = None  # сбрасываем закэшированный клиент - вдруг ключ поменялся
        _ai_settings_cache = merged
        _save_ai_settings_to_disk(merged)
        return dict(merged)


def _resolve_api_key(provider: str) -> str:
    """Переменная окружения имеет приоритет (для тех, кто предпочитает
    настраивать через неё) - но если её нет, используется ключ, вставленный
    прямо во вкладке 'ИИ-проверка' и сохранённый в ai_settings.json."""
    env_var = "GEMINI_API_KEY" if provider == "gemini" else "ANTHROPIC_API_KEY"
    key = os.environ.get(env_var)
    if key:
        return key
    settings = get_ai_settings()
    return settings.get(f"{provider}_api_key", "") or ""


def ai_provider_ready(provider: str) -> bool:
    """Есть ли всё нужное для реального вызова этого провайдера (ключ,
    модуль) - используется и в интерфейсе (чтобы показать честный статус),
    и внутри process_tender (чтобы не пытаться звонить туда, где заведомо
    ничего не выйдет)."""
    if provider == "gemini":
        return bool(_resolve_api_key("gemini"))
    if provider == "anthropic":
        return ANTHROPIC_AVAILABLE and bool(_resolve_api_key("anthropic"))
    return False


_anthropic_client = None
_anthropic_client_key = None
_anthropic_client_lock = threading.Lock()


def _get_anthropic_client(api_key: str):
    global _anthropic_client, _anthropic_client_key
    if not ANTHROPIC_AVAILABLE or not api_key:
        return None
    with _anthropic_client_lock:
        if _anthropic_client is None or _anthropic_client_key != api_key:
            _anthropic_client = anthropic.Anthropic(api_key=api_key)
            _anthropic_client_key = api_key
        return _anthropic_client


GEMINI_API_BASE = "https://generativelanguage.googleapis.com/v1beta/models"

_AI_RELEVANCE_PROMPT = """Ты помогаешь отобрать релевантные тендеры среди объявлений о госзакупках Казахстана.

Критерий релевантности (сформулирован пользователем):
---
{criteria}
---

Название лота: {title}

Описание/детали объявления (может быть обрезано):
---
{description}
---

Определи, подходит ли этот тендер под критерий выше. Учитывай смысл, а не только точные слова -
формулировки в реальных тендерах бывают самые разные (падежи, синонимы, канцелярит).

Ответь СТРОГО в формате JSON, без markdown-разметки и пояснений вокруг:
{{"relevant": true или false, "reasoning": "одно короткое предложение почему"}}
"""


def _parse_ai_json_response(text: str):
    text = re.sub(r"^```(?:json)?|```$", "", text.strip(), flags=re.MULTILINE).strip()
    data = json.loads(text)
    return {"relevant": bool(data.get("relevant")), "reasoning": str(data.get("reasoning", ""))[:300]}


# Причина последнего неудавшегося вызова ИИ - читается вызывающим кодом
# (например, "Пересмотр ИИ"), чтобы показать пользователю ПОЧЕМУ проверка
# не сработала, а не просто молча пропустить строку. ai_check_relevance()
# по контракту всё равно возвращает None при ошибке (это не меняется, чтобы
# не ломать process_tender) - это отдельный, дополнительный канал именно
# для диагностики.
_last_ai_error = None


def _ai_check_relevance_gemini(title: str, description: str, criteria: str, model: str, api_key: str):
    global _last_ai_error
    if not api_key:
        _last_ai_error = "Не задан ключ Gemini"
        return None
    prompt = _AI_RELEVANCE_PROMPT.format(
        criteria=criteria.strip(), title=(title or "")[:300], description=(description or "")[:4000],
    )
    url = f"{GEMINI_API_BASE}/{model}:generateContent?key={api_key}"
    try:
        resp = requests.post(
            url, json={"contents": [{"parts": [{"text": prompt}]}]},
            timeout=30,
        )
        if not resp.ok:
            # Тело ответа Gemini при ошибке обычно прямо объясняет причину
            # (неверный ключ, модель не найдена, превышена квота и т.д.) -
            # это самое ценное, что можно показать пользователю.
            _last_ai_error = f"Gemini HTTP {resp.status_code}: {resp.text[:300]}"
            return None
        data = resp.json()
        text = data["candidates"][0]["content"]["parts"][0]["text"]
        result = _parse_ai_json_response(text)
        _last_ai_error = None
        return result
    except Exception as e:
        _last_ai_error = f"Gemini: {e}"
        log_error_traceback(f"ai_check_relevance (gemini): {e}")
        return None


def _ai_check_relevance_anthropic(title: str, description: str, criteria: str, model: str, api_key: str):
    global _last_ai_error
    client = _get_anthropic_client(api_key)
    if client is None:
        _last_ai_error = "Не задан ключ Anthropic или модуль 'anthropic' не установлен"
        return None
    try:
        resp = client.messages.create(
            model=model,
            max_tokens=200,
            messages=[{
                "role": "user",
                "content": _AI_RELEVANCE_PROMPT.format(
                    criteria=criteria.strip(),
                    title=(title or "")[:300],
                    description=(description or "")[:4000],
                ),
            }],
        )
        text = "".join(getattr(block, "text", "") for block in resp.content)
        result = _parse_ai_json_response(text)
        _last_ai_error = None
        return result
    except Exception as e:
        _last_ai_error = f"Anthropic: {e}"
        log_error_traceback(f"ai_check_relevance (anthropic): {e}")
        return None


def ai_check_relevance(title: str, description: str, criteria: str, model: str, provider: str = "gemini"):
    """Спрашивает у ИИ (Gemini по умолчанию, либо Claude), подходит ли
    тендер под свободно описанный критерий. Возвращает
    dict {"relevant": bool, "reasoning": str} или None, если проверка не
    удалась (нет ключа/модуля, сетевая ошибка, кривой ответ) - в этом
    случае вызывающий код должен просто ничего не менять (не отклонять и
    не продвигать тендер на основании неудавшейся проверки)."""
    api_key = _resolve_api_key(provider)
    if provider == "anthropic":
        return _ai_check_relevance_anthropic(title, description, criteria, model, api_key)
    return _ai_check_relevance_gemini(title, description, criteria, model, api_key)

# Слова-исключения "продление лицензии" — не редактируются через UI
# (это не тематические ключевые слова, а грамматические паттерны),
# но их тоже можно вынести в JSON при желании — см. README.
LICENSE_EXCLUDE_PATTERNS = [
    "продление лиценз", "продление права использования", "продление права пользования",
    "продление подписки", "подписка", "renewal", "subscription", "annual support",
    "техническая поддержка программного обеспечения", "продление технической поддержки",
    "продление сопровождения", "продление доступа", "право использования",
    "неисключительная лицензия", "лицензия на использование", "продление лицензии",
    "продление лицензий", "продление лицензии на", "предоставление лицензии",
    "право пользования по", "техническая поддержка и сопровождение",
    "лицензияны ұзарту", "жазылымды ұзарту", "жазылым", "пайдалану құқығын ұзарту",
    "техникалық қолдауды ұзарту", "сүйемелдеуді ұзарту", "қолжетімділікті ұзарту",
    "пайдалану құқығы", "айрықша емес лицензия", "пайдалануға арналған лицензия",
]

_LICENSE_RENEWAL_EXCEPTION_WORDS = [
    "разработка", "доработка", "модернизация", "внедрение", "создание",
    "разработка информационной системы", "создание информационной системы",
    "әзірлеу", "жетілдіру", "пысықтау", "жаңғырту", "енгізу", "құру",
]

_KAZAKH_LICENSE_SUBJECT_STEMS = ["лицензи", "жазылым"]
_KAZAKH_RENEWAL_STEM = "ұзарт"

ANNOUNCE_LINK_RE = re.compile(r"/ru/announce/index/(\d+)")
NUMBER_ANNO_RE = re.compile(r"\b(\d{5,}-\d+)\b")
_STOP_LOOKAHEAD = (
    r"(?=\s*\d{4}-\d{2}-\d{2}|\s*[\d\s\u00a0]{1,20}[.,]\d{2}\s*(?:$|[А-ЯЁ])"
    r"|Заказчик:|Организатор:|Способ[^:\n]{0,30}:|Статус:|$)"
)
ORGANIZER_RE = re.compile(r"Организатор:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
CUSTOMER_RE = re.compile(r"Заказчик:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
DOC_EXTENSIONS = (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".rtf", ".zip", ".rar", ".txt")

# "Юр. адрес организатора" (иногда встречается как "Юридический адрес
# организатора") - поле в блоке "Общие сведения" на детальной странице
# объявления. Само название/наименование организатора не всегда содержит
# город (особенно если это министерство/ведомство с обезличенным
# названием), тогда как юридический адрес почти всегда его указывает -
# поэтому используется как дополнительный (более надёжный) признак при
# определении тендеров по Астане, см. is_astana_tender().
ORGANIZER_LEGAL_ADDRESS_RE = re.compile(
    r"Юр(?:\.|идический)?\s*адрес\s*организатора\s*[:\s]*\n?\s*([^\n]{1,200})", re.I,
)

HEADERS = [
    "№ объявления", "Название лота/объявления", "Заказчик", "Организатор",
    "Способ закупки", "Сумма, тг.", "Статус", "Начало приёма заявок",
    "Окончание приёма заявок", "Ссылка на объявление", "Совпавшие IT-ключевые слова",
    "Документы (названия и ссылки)", "Тип IT", "Приоритет",
    "Краткое описание/ТЗ (выдержка)", "Обоснование ИИ",
]

KNOWN_STATUSES = [
    "Опубликовано (прием ценовых предложений)", "Опубликовано (ожидание проведения аукциона)",
    "Опубликовано (дополнение заявок)", "Опубликовано (прием заявок)",
    "Опубликовано (проверка заявок)", "Ожидание камерального контроля",
    "Ожидание проведения контроля качества", "Отправлено на контроль качества",
    "Пройдено контроль качества", "Рассмотрение дополнений заявок",
    "Формирование протокола промежуточных итогов", "Формирование протокола преддопуска",
    "Формирование протокола допуска", "Формирование протокола итогов",
    "Принятие решение о пересмотре итогов", "Принятие решение об исполнении уведомления",
    "Изменена документация", "Пересмотр итогов", "На обжаловании", "Приостановлено",
    "Отказ от закупки", "Рассмотрение заявок", "Опубликовано", "Завершено", "Отменено",
]

KNOWN_METHODS = [
    "Второй этап конкурса с использованием рамочного соглашения",
    "Первый этап конкурса с использованием рамочного соглашения",
    "Из одного источника по несостоявшимся закупкам не ГЗ",
    "Из одного источника по несостоявшимся закупкам",
    "Конкурс с предварительным квалификационным отбором",
    "Конкурс с применением двухэтапных процедур",
    "Конкурс с использованием рейтингово-балльной системы",
    "Конкурс с применением специального порядка",
    "Тендер с использованием рейтингово-балльной системы",
    "Тендер с применением особого порядка",
    "Государственные закупки с применением особого порядка",
    "Запрос ценовых предложений (не ГЗ new)", "Запрос ценовых предложений",
    "Открытый конкурс", "Аукцион (до 2022)", "Аукцион (с 2022)",
    "Аукцион (не ГЗ) с 2022 г", "Аукцион", "Закупка жилища",
    "Закупка по государственному социальному заказу",
    "Конкурс по строительству «под ключ»", "Тендер по строительству «под ключ»", "Тендер",
]

ASTANA_CITY_RE = re.compile(
    r"\b(?:г\.?\s*)?астан[аеу]\b|\bnur[\s\-]?sultan\b|\bнур[\s\-]?султан\w*\b", re.I,
)

EXCLUDE_FALSE_POSITIVE_PHRASES = [
    "информационные услуги", "информационное сопровождение",
    "информационно-разъяснительная работа", "информационно-консультационные услуги",
    "информационное обеспечение мероприятий", "размещение информации",
    "изготовление информационных стендов", "информационные материалы",
    "полиграфическая продукция", "освещение деятельности", "публикация материалов",
    "рекламные услуги", "услуги по информированию населения",
    "ақпараттық қызметтер", "ақпараттық сүйемелдеу",
    "іс-шараларды ақпараттық қамтамасыз ету", "ақпаратты орналастыру",
    "ақпараттық стендтер дайындау", "ақпараттық материалдар", "полиграфия өнімдері",
    "қызметті жария ету", "материалдарды жариялау", "жарнама қызметтері",
    "халықты хабардар ету қызметтері",
]


# ============================================================================
# ============================ ВСПОМОГАТЕЛЬНОЕ ==============================
# ============================================================================

def polite_sleep(delay_range):
    time.sleep(random.uniform(*delay_range))


def dump_debug_html(name: str, html: str):
    if not DEBUG_MODE:
        return
    os.makedirs(DEBUG_DUMP_DIR, exist_ok=True)
    with open(os.path.join(DEBUG_DUMP_DIR, name), "w", encoding="utf-8") as f:
        f.write(html)


def log_error_traceback(header: str):
    try:
        with open(ERROR_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(f"\n--- {time.strftime('%Y-%m-%d %H:%M:%S')} | {header} ---\n")
            f.write(traceback.format_exc())
            f.write("\n")
    except Exception:
        pass


# ============================================================================
# ============================== БРАУЗЕР =====================================
# ============================================================================

def _edge_debug_port_alive() -> bool:
    try:
        r = requests.get(f"http://{EDGE_DEBUG_ADDRESS}/json/version", timeout=1.5)
        return r.ok
    except Exception:
        return False


def try_auto_launch_edge():
    if _edge_debug_port_alive():
        log("Порт отладки уже отвечает - использую уже открытое окно автоматизации.")
        return True, "already_running"

    if not AUTO_LAUNCH_EDGE:
        log("[!] AUTO_LAUNCH_EDGE выключен. Откройте Edge вручную с флагом --remote-debugging-port")
        return False, "auto_launch_disabled"

    if not os.path.exists(EDGE_BROWSER_PATH):
        log(f"[!] Не нашёл Edge по пути {EDGE_BROWSER_PATH}.")
        return False, "edge_not_found"

    os.makedirs(AUTOMATION_PROFILE_DIR, exist_ok=True)
    port = EDGE_DEBUG_ADDRESS.split(":")[-1]
    log(f"Открываю отдельное окно Edge для автоматизации: {AUTOMATION_PROFILE_DIR}")
    subprocess.Popen([
        EDGE_BROWSER_PATH,
        f"--remote-debugging-port={port}",
        f"--user-data-dir={AUTOMATION_PROFILE_DIR}",
        "--profile-directory=Default",
        "--no-first-run",
        "--no-default-browser-check",
        "https://goszakup.gov.kz/ru/announce/index",
    ])

    for i in range(120):
        time.sleep(0.5)
        if _edge_debug_port_alive():
            log("Отдельное окно Edge запущено, порт отладки открылся успешно.")
            return True, "launched"
        if i > 0 and i % 10 == 0:
            log(f"    ...жду запуск Edge ({i // 2} сек)")

    log("[!] Порт отладки не открылся. Возможно, окно автоматизации уже висит в фоне.")
    return False, "timeout"


def attach_to_running_edge():
    options = EdgeOptions()
    options.add_experimental_option("debuggerAddress", EDGE_DEBUG_ADDRESS)
    service = EdgeService(executable_path=EDGE_DRIVER_PATH)
    return webdriver.Edge(service=service, options=options)


def get_cookies_dict(driver) -> dict:
    return {c["name"]: c["value"] for c in driver.get_cookies()}


def find_and_switch_to_goszakup_tab(driver):
    found_url = None
    found_handle = None
    for handle in driver.window_handles:
        try:
            driver.switch_to.window(handle)
            url = driver.current_url
        except Exception:
            continue
        if "goszakup.gov.kz" in url:
            found_url = url
            found_handle = handle
            if "search/announce" in url or "announce" in url:
                break
    if found_handle is not None:
        driver.switch_to.window(found_handle)
    return found_url


# ============================================================================
# ======================== ПАРСИНГ СТРАНИЦЫ СПИСКА ==========================
# ============================================================================

def build_page_url(base_url: str, page: int) -> str:
    parts = urlsplit(base_url)
    params = parse_qsl(parts.query, keep_blank_values=True)
    new_params = [(k, v) for k, v in params if k not in ("page", "count_record")]
    new_params.append(("page", str(page)))
    new_params.append(("count_record", str(COUNT_RECORD_PER_PAGE)))
    return urlunsplit((parts.scheme, parts.netloc, parts.path, urlencode(new_params), ""))


def _find_status(row_text: str) -> str:
    for status in KNOWN_STATUSES:
        if status in row_text:
            return status
    return ""


def _find_method(row_text: str) -> str:
    for method in KNOWN_METHODS:
        if method in row_text:
            return method
    return ""


_DETAIL_STATUS_LABEL_RE = re.compile(r"Статус:\s*\n?\s*([^\n]{1,120})")


def _extract_detail_status(description_text: str) -> str:
    """Определяет текущий статус тендера по тексту детальной страницы -
    используется при периодической проверке уже сохранённых тендеров (см.
    ScraperRunner._sweep_existing_tenders). В отличие от _find_status()
    (который просто ищет первое совпадение из KNOWN_STATUSES где угодно в
    тексте - для строки таблицы этого достаточно), детальная страница может
    содержать историю смены статусов, поэтому сначала пытаемся найти именно
    значение сразу после метки 'Статус:', и только если это не удалось -
    откатываемся на менее точный поиск по всему тексту."""
    if not description_text:
        return ""
    m = _DETAIL_STATUS_LABEL_RE.search(description_text)
    if m:
        candidate = m.group(1).strip()
        canonical = _find_status(candidate)
        if canonical:
            return canonical
        return candidate
    return _find_status(description_text)


def _clean_amount(raw: str) -> str:
    if not raw:
        return raw
    s = raw.strip()
    m = re.match(r"^0{1,3}[\s\u00a0]+(\d{1,3}(?:[\s\u00a0]\d{3})*[.,]\d{2})$", s)
    if m:
        return m.group(1)
    return s


def _extract_amount(row, row_text: str) -> str:
    money_cell_re = re.compile(r"^[\d\s\u00a0]{1,20}[.,]\d{2}$")
    if row is not None and getattr(row, "name", None) == "tr":
        cells = row.find_all(["td", "th"])
        for cell in reversed(cells):
            cell_text = cell.get_text(strip=True)
            if money_cell_re.match(cell_text):
                return _clean_amount(cell_text)

    amount_matches = re.findall(r"[\d\s\u00a0]{1,20}[.,]\d{2}", row_text)
    if amount_matches:
        return _clean_amount(amount_matches[-1].strip())
    return ""


def _truncate_leaked_text(text: str, *markers: str) -> str:
    if not text:
        return text
    cut_at = len(text)
    for marker in markers:
        if marker:
            idx = text.find(marker)
            if idx != -1:
                cut_at = min(cut_at, idx)
    return text[:cut_at].strip(" ,.-\u00a0")


def parse_list_page(html: str) -> list:
    soup = BeautifulSoup(html, "html.parser")
    results = []
    seen_ids_on_page = set()

    for a_tag in soup.find_all("a", href=ANNOUNCE_LINK_RE):
        match = ANNOUNCE_LINK_RE.search(a_tag.get("href", ""))
        if not match:
            continue
        tender_id = match.group(1)
        if tender_id in seen_ids_on_page:
            continue
        seen_ids_on_page.add(tender_id)

        title = a_tag.get_text(strip=True)
        if not title:
            continue

        row = a_tag.find_parent("tr") or a_tag.find_parent(["div", "li"])
        row_text = row.get_text(" ", strip=True) if row else a_tag.get_text(" ", strip=True)

        number_anno_match = NUMBER_ANNO_RE.search(row_text)
        number_anno = number_anno_match.group(1) if number_anno_match else tender_id

        organizer_match = ORGANIZER_RE.search(row_text)
        organizer_name = organizer_match.group(1).strip() if organizer_match else ""

        customer_match = CUSTOMER_RE.search(row_text)
        customer_name = customer_match.group(1).strip() if customer_match else ""

        status = _find_status(row_text)
        method = _find_method(row_text)

        organizer_name = _truncate_leaked_text(organizer_name, method, status)
        customer_name = _truncate_leaked_text(customer_name, method, status)

        amount = _extract_amount(row, row_text)

        dates = re.findall(r"\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(?::\d{2})?", row_text)
        start_date = dates[0] if len(dates) >= 1 else ""
        end_date = dates[1] if len(dates) >= 2 else ""

        results.append({
            "id": tender_id, "number_anno": number_anno, "title": title,
            "url": f"https://goszakup.gov.kz/ru/announce/index/{tender_id}",
            "organizer_name": organizer_name, "customer_name": customer_name,
            "status": status, "method": method, "amount": amount,
            "start_date": start_date, "end_date": end_date,
        })

    return results


# ============================================================================
# ======================= ПАРСИНГ ДЕТАЛЬНОЙ СТРАНИЦЫ =========================
# ============================================================================

_CHROME_TAGS = ("script", "style", "nav", "header", "footer", "select", "option", "noscript", "iframe")
_CHROME_HINT_RE = re.compile(
    r"(menu|navbar|nav-|footer|header|sidebar|breadcrumb|filter|search-form|"
    r"cookie|lang-switch|top-panel|topline|langs)", re.I,
)


def _strip_page_chrome(soup: BeautifulSoup) -> None:
    for tag_name in _CHROME_TAGS:
        for tag in soup.find_all(tag_name):
            tag.decompose()
    for tag in soup.find_all(True):
        if getattr(tag, "decomposed", False):
            continue
        cls = " ".join(tag.get("class", []) or [])
        tag_id = tag.get("id", "") or ""
        if (cls and _CHROME_HINT_RE.search(cls)) or (tag_id and _CHROME_HINT_RE.search(tag_id)):
            tag.decompose()


def parse_detail_page(html: str, page_url: str) -> dict:
    soup = BeautifulSoup(html, "html.parser")

    documents = []
    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        href_lower = href.lower()
        if href_lower.endswith(DOC_EXTENSIONS) or "download" in href_lower or "/file/" in href_lower:
            full_url = urljoin(page_url, href)
            doc_name = a_tag.get_text(strip=True) or full_url.split("/")[-1]
            documents.append({"name": doc_name, "url": full_url})

    _strip_page_chrome(soup)

    main_container = (
        soup.find("div", attrs={"id": re.compile(r"content", re.I)})
        or soup.find("main") or soup.find("table") or soup.find("body")
    )
    description_text = main_container.get_text("\n", strip=True) if main_container else ""

    customer_bin = ""
    bin_match = re.search(r"БИН[:\s]*([0-9]{12})", description_text)
    if bin_match:
        customer_bin = bin_match.group(1)

    organizer_legal_address = ""
    addr_match = ORGANIZER_LEGAL_ADDRESS_RE.search(description_text)
    if addr_match:
        organizer_legal_address = addr_match.group(1).strip()

    return {
        "description": description_text, "documents": documents,
        "customer_bin": customer_bin, "organizer_legal_address": organizer_legal_address,
    }


# ============================================================================
# ==================== СКАЧИВАНИЕ И ЧТЕНИЕ ДОКУМЕНТОВ ========================
# ============================================================================

def _safe_filename(url: str) -> str:
    name = url.split("?")[0].rstrip("/").split("/")[-1]
    name = re.sub(r"[^\w\.\-]+", "_", name)
    if not name:
        name = uuid.uuid4().hex
    return f"{uuid.uuid4().hex[:8]}_{name}"


def download_file(url: str, session: requests.Session, cookies: dict):
    os.makedirs(DOWNLOAD_TMP_DIR, exist_ok=True)
    local_path = os.path.join(DOWNLOAD_TMP_DIR, _safe_filename(url))
    try:
        with session.get(url, cookies=cookies, stream=True, timeout=30) as resp:
            resp.raise_for_status()
            total = 0
            max_bytes = MAX_DOCUMENT_SIZE_MB * 1024 * 1024
            with open(local_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    total += len(chunk)
                    if total > max_bytes:
                        f.close()
                        os.remove(local_path)
                        return None
                    f.write(chunk)
        return local_path
    except Exception:
        if os.path.exists(local_path):
            try:
                os.remove(local_path)
            except OSError:
                pass
        return None


def _extract_pdf(path):
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docx(path):
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_doc_legacy(path):
    try:
        import textract
        return textract.process(path).decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_xlsx(path):
    try:
        import openpyxl as _openpyxl
        wb = _openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        parts.append(str(cell))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_xls_legacy(path):
    try:
        import xlrd
        wb = xlrd.open_workbook(path)
        parts = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for cell in sheet.row(row_idx):
                    parts.append(str(cell.value))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".doc":
            return _extract_doc_legacy(path)
        elif ext in (".xlsx", ".xlsm"):
            return _extract_xlsx(path)
        elif ext == ".xls":
            return _extract_xls_legacy(path)
        elif ext in (".txt", ".rtf"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        return ""
    except Exception:
        return ""


def get_document_text(url: str, session: requests.Session, cookies: dict) -> str:
    path = download_file(url, session, cookies)
    if not path:
        return ""
    try:
        return extract_text_from_file(path)
    finally:
        if DELETE_DOWNLOADED_DOCS_AFTER_SCAN and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass


# ============================================================================
# ================================ КЛАССИФИКАЦИЯ =============================
# ============================================================================

def _normalize(text: str) -> str:
    return (text or "").lower().replace("ё", "е")


def is_license_renewal(*texts):
    text = _normalize(" ".join(filter(None, texts)))

    is_renewal_phrase = any(p in text for p in LICENSE_EXCLUDE_PATTERNS)
    is_renewal_kz_stems = (
        _KAZAKH_RENEWAL_STEM in text
        and any(s in text for s in _KAZAKH_LICENSE_SUBJECT_STEMS)
    )

    if is_renewal_phrase or is_renewal_kz_stems:
        if not any(x in text for x in _LICENSE_RENEWAL_EXCEPTION_WORDS):
            return True

    return False


def _collapse_repeats(s: str) -> str:
    return re.sub(r"(.)\1+", r"\1", s)


def classify_is_it(title="", description="", documents=""):
    title = _normalize(title)
    description = _normalize(description)
    documents = _normalize(documents)

    score = 0
    matched = []

    for phrase in EXCLUDE_FALSE_POSITIVE_PHRASES:
        title = title.replace(_normalize(phrase), "")
        description = description.replace(_normalize(phrase), "")
        documents = documents.replace(_normalize(phrase), "")

    def _kw_hit(text: str, text_collapsed: str, kw: str) -> bool:
        kw_norm = _normalize(kw)
        if kw_norm in text:
            return True
        return _collapse_repeats(kw_norm) in text_collapsed

    def scan(text, strong, medium, weak, multiplier):
        nonlocal score, matched
        text_collapsed = _collapse_repeats(text)

        for kw in strong:
            if _kw_hit(text, text_collapsed, kw):
                matched.append(kw)
                score += 10 * multiplier
        for kw in medium:
            if _kw_hit(text, text_collapsed, kw):
                matched.append(kw)
                score += 4 * multiplier
        for kw in weak:
            if _kw_hit(text, text_collapsed, kw):
                matched.append(kw)
                score += 1 * multiplier

    with _kw_lock:
        strong, medium, weak = list(IT_STRONG_KEYWORDS), list(IT_MEDIUM_KEYWORDS), list(IT_WEAK_KEYWORDS)
        negative, context = list(IT_NEGATIVE_KEYWORDS), list(IT_CONTEXT_PHRASES)

    scan(title, strong, medium, weak, 3)
    scan(description, strong, medium, weak, 2)
    scan(documents, strong, medium, weak, 1)

    title_and_desc = title + " " + description
    title_and_desc_collapsed = _collapse_repeats(title_and_desc)
    for phrase in context:
        if _kw_hit(title_and_desc, title_and_desc_collapsed, phrase):
            matched.append(phrase)
            score += 8

    combined_text_for_negatives = " ".join([title, description, documents])
    for kw in negative:
        if kw in combined_text_for_negatives:
            score -= 5

    matched = sorted(set(matched))
    return score >= IT_SCORE_THRESHOLD, score, matched


def detect_it_type(*text_fragments):
    text = _normalize(" ".join(t for t in text_fragments if t))

    if any(x in text for x in [
        "разработка", "доработка", "внедрение", "crm", "erp", "api", "веб",
        "мобильное приложени", "әзірлеу", "жетілдіру", "енгізу", "мобильді қосымша",
    ]):
        return "ПО"

    if any(x in text for x in [
        "сервер", "коммутатор", "маршрутизатор", "сетевое оборудование", "скс", "лвс",
        "firewall", "желілік жабдық", "жергілікті желі",
    ]):
        return "Оборудование"

    if any(x in text for x in [
        "техническая поддержка", "сопровождение", "аутсорсинг", "администрирование",
        "helpdesk", "техникалық қолдау", "сүйемелдеу", "әкімшілендіру",
    ]):
        return "Услуги"

    return "Другое"


def detect_priority(*text_fragments):
    text = _normalize(" ".join(t for t in text_fragments if t))
    score = 0

    high = [
        "разработка", "внедрение", "информационная система", "erp", "crm", "api",
        "искусственный интеллект", "модернизация", "әзірлеу", "енгізу",
        "ақпараттық жүйе", "жасанды интеллект", "жаңғырту",
    ]
    medium = [
        "сервер", "сетевое оборудование", "база данных", "виртуализация", "sql",
        "postgresql", "1с", "желілік жабдық", "дерекқор", "виртуалдандыру",
    ]
    low = [
        "техническая поддержка", "сопровождение", "лицензи",
        "техникалық қолдау", "сүйемелдеу", "лицензия",
    ]

    score += sum(5 for w in high if w in text)
    score += sum(3 for w in medium if w in text)
    score += sum(1 for w in low if w in text)

    if score >= 10:
        return "Высокий"
    if score >= 5:
        return "Средний"
    return "Низкий"


_borderline_lock = threading.Lock()


def log_borderline_candidate(row: dict, score: int, matched: list):
    # Все поля после title идут как самоописанные key=value (amount=,
    # end_date=, keywords=) - это позволяет добавлять новые поля в будущем
    # без ломки обратной совместимости чтения старых строк лога (парсинг в
    # app.py._read_borderline_rows разбирает их по тегу, а не по позиции).
    try:
        with _borderline_lock:
            with open(BORDERLINE_LOG_FILE, "a", encoding="utf-8") as f:
                f.write(
                    f"{time.strftime('%Y-%m-%d %H:%M:%S')}\t"
                    f"score={score}\t{row.get('number_anno', '')}\t{row.get('url', '')}\t"
                    f"{row.get('title', '')[:120]}\tamount={row.get('amount', '')}\t"
                    f"start_date={row.get('start_date', '')}\tend_date={row.get('end_date', '')}\t"
                    f"keywords={', '.join(matched)}\n"
                )
    except Exception:
        pass


def is_astana_tender(row: dict, detail: dict) -> bool:
    """Определяет, относится ли тендер к городу Астана.

    Смотрим на заказчика/организатора из списка, на юридический адрес
    организатора из блока "Общие сведения" детальной страницы (см.
    ORGANIZER_LEGAL_ADDRESS_RE - часто это самый надёжный источник, т.к.
    само наименование организатора город может не упоминать) и, как
    резервный вариант, на начало текста самого объявления - там иногда
    указывается адрес заказчика."""
    haystack = " ".join([
        row.get("customer_name", "") or "",
        row.get("organizer_name", "") or "",
        detail.get("organizer_legal_address", "") or "",
        (detail.get("description", "") or "")[:4000],
    ])
    return bool(ASTANA_CITY_RE.search(haystack))


# ============================================================================
# ================================ EXCEL =====================================
# ============================================================================

class ExcelSink:
    def __init__(self, path: str, sheet_title: str = "IT-тендеры"):
        self.path = path
        self.sheet_title = sheet_title
        self.lock = threading.Lock()
        self.workbook = None
        self.worksheet = None
        self.rows_since_save = 0
        self.existing_ids = set()

    def _build_new_workbook(self):
        wb = Workbook()
        ws = wb.active
        ws.title = self.sheet_title
        ws.append(HEADERS)
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        for col_idx in range(1, len(HEADERS) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center", wrap_text=True)
        widths = [16, 45, 35, 30, 22, 16, 24, 18, 18, 30, 30, 18, 18, 45, 50, 55]
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A2"
        return wb, ws

    def init_for_run(self):
        if os.path.exists(self.path):
            self.workbook = load_workbook(self.path)
            self.worksheet = (
                self.workbook[self.sheet_title]
                if self.sheet_title in self.workbook.sheetnames
                else self.workbook.active
            )
        else:
            self.workbook, self.worksheet = self._build_new_workbook()
            self.workbook.save(self.path)

        self.existing_ids = set()
        url_col_idx = HEADERS.index("Ссылка на объявление") + 1
        for row_cells in self.worksheet.iter_rows(min_row=2, values_only=True):
            if url_col_idx - 1 < len(row_cells):
                url_value = row_cells[url_col_idx - 1]
                tid = _extract_tender_id_from_url(url_value)
                if tid:
                    self.existing_ids.add(tid)

    def save_now(self):
        with self.lock:
            if self.workbook is not None:
                self.workbook.save(self.path)
                self.rows_since_save = 0

    def append_row(self, row: dict) -> bool:
        tender_id = _extract_tender_id_from_url(row.get("url", ""))

        with self.lock:
            if tender_id and tender_id in self.existing_ids:
                return False

            values = [
                row.get("number_anno", ""), row.get("title", ""), row.get("customer_name", ""),
                row.get("organizer_name", ""), row.get("method", ""), row.get("amount", ""),
                row.get("status", ""), row.get("start_date", ""), row.get("end_date", ""),
                row.get("url", ""), ", ".join(row.get("matched_keywords", [])),
                "; ".join(row.get("documents", [])), row.get("it_type", ""),
                row.get("priority", ""), (row.get("description", "") or "")[:1500],
                row.get("ai_reasoning", ""),
            ]

            self.worksheet.append(values)
            last_row = self.worksheet.max_row
            url_cell = self.worksheet.cell(row=last_row, column=HEADERS.index("Ссылка на объявление") + 1)
            if row.get("url"):
                url_cell.hyperlink = row["url"]
                url_cell.font = Font(color="0563C1", underline="single")
            if tender_id:
                self.existing_ids.add(tender_id)
            self.rows_since_save += 1
            need_save = self.rows_since_save >= XLSX_SAVE_EVERY_N_MATCHES

        if need_save:
            self.save_now()

        return True


def _extract_tender_id_from_url(url) -> str:
    if not url:
        return ""
    m = ANNOUNCE_LINK_RE.search(str(url))
    return m.group(1) if m else ""


def is_expired(end_date_str: str) -> bool:
    """Проверяет, истёк ли срок приёма заявок. Даты на сайте приходят в
    формате 'YYYY-MM-DD HH:MM' или 'YYYY-MM-DD HH:MM:SS' - это позволяет
    сравнивать их с текущим временем как обычные строки (лексикографическое
    сравнение ISO-подобных дат даёт тот же результат, что и сравнение
    настоящих datetime), без риска ошибок парсинга разных форматов."""
    end_date_str = (end_date_str or "").strip()
    if not end_date_str:
        return False
    now_str = time.strftime("%Y-%m-%d %H:%M:%S")
    if len(end_date_str) == 16:  # 'YYYY-MM-DD HH:MM' без секунд
        end_date_str = end_date_str + ":00"
    return end_date_str < now_str


def days_left(end_date_str: str):
    """Сколько календарных дней осталось до окончания приёма заявок (по
    датам, без учёта времени суток) - None, если дата неизвестна/не
    распознана. Может быть отрицательным для уже просроченных тендеров -
    используется вместе с is_expired() для жёлтого предупреждения
    'осталось N дней' (см. DEADLINE_WARNING_DAYS ниже), которое не должно
    само по себе означать 'уже истёк'."""
    end_date_str = (end_date_str or "").strip()
    if len(end_date_str) < 10:
        return None
    try:
        y, m, d = (int(x) for x in end_date_str[:10].split("-"))
        end_date = date(y, m, d)
    except Exception:
        return None
    return (end_date - date.fromtimestamp(time.time())).days


DEADLINE_WARNING_DAYS = 3


def delete_tender_row(tender_id: str) -> bool:
    """Удаляет тендер из goszakup_it_tenders.xlsx/_astana.xlsx по id -
    вызывается пользователем вручную с кнопки 'Удалить' (например, для
    просроченных тендеров, автоматика их не трогает). В отличие от
    ScraperRunner._sweep_existing_tenders (которая чистит ЗАВЕРШЁННЫЕ по
    статусу тендеры автоматически при каждом запуске сбора), это разовое
    ручное действие, доступное в любой момент через веб-интерфейс -
    работает даже если сбор ни разу не запускался в этом процессе
    (тогда сначала лениво подгружает книгу с диска)."""
    if not tender_id:
        return False
    url_col = HEADERS.index("Ссылка на объявление") + 1
    for sink in (sink_main, sink_astana):
        with sink.lock:
            if sink.worksheet is None:
                sink.init_for_run()
            if tender_id not in sink.existing_ids:
                continue
            found_row = None
            for row_idx in range(2, sink.worksheet.max_row + 1):
                url_val = sink.worksheet.cell(row=row_idx, column=url_col).value
                if _extract_tender_id_from_url(url_val) == tender_id:
                    found_row = row_idx
                    break
            if found_row is None:
                continue
            sink.worksheet.delete_rows(found_row)
            sink.existing_ids.discard(tender_id)
        sink.save_now()
        remove_viewed(tender_id)
        return True
    return False


def delete_borderline_by_tender_id(tender_id: str) -> int:
    """Удаляет из borderline_candidates.log все строки, относящиеся к
    данному тендеру (обычно одна, но их может быть несколько, если тендер
    несколько раз подряд оставался пограничным на разных прогонах)."""
    if not tender_id or not os.path.exists(BORDERLINE_LOG_FILE):
        return 0
    removed = 0
    with _borderline_lock:
        with open(BORDERLINE_LOG_FILE, "r", encoding="utf-8") as f:
            lines = f.readlines()
        kept = []
        for line in lines:
            stripped = line.rstrip("\n")
            parts = stripped.split("\t") if stripped else []
            url = parts[3] if len(parts) > 3 else ""
            if _extract_tender_id_from_url(url) == tender_id:
                removed += 1
                continue
            kept.append(line)
        if removed:
            with open(BORDERLINE_LOG_FILE, "w", encoding="utf-8") as f:
                f.writelines(kept)
    return removed


sink_main = ExcelSink(OUTPUT_XLSX, sheet_title="IT-тендеры")
sink_astana = ExcelSink(OUTPUT_XLSX_ASTANA, sheet_title="IT-тендеры (Астана)")


def init_workbooks_for_run():
    sink_main.init_for_run()
    sink_astana.init_for_run()


def save_workbooks_now():
    sink_main.save_now()
    sink_astana.save_now()


def append_it_result(result: dict, astana: bool) -> bool:
    if astana:
        return sink_astana.append_row(result)
    else:
        return sink_main.append_row(result)


# ============================================================================
# ============================== ЧЕКПОИНТ ====================================
# ============================================================================

_seen_ids_lock = threading.Lock()


def load_seen_ids() -> set:
    if not os.path.exists(CHECKPOINT_FILE):
        return set()
    try:
        with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("seen_ids", []))
    except Exception:
        return set()


def save_seen_ids(seen_ids: set):
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump({"seen_ids": sorted(seen_ids)}, f, ensure_ascii=False, indent=2)


# ============================================================================
# ==================== "ПРОСМОТРЕНО" (галочка в веб-интерфейсе) ==============
# ============================================================================
# Хранится отдельно от Excel/чекпоинта в своём JSON-файле, ключ - числовой id
# тендера (тот же, что используется для дедупликации в ExcelSink). Переживает
# перезапуски и app.py, и самого сбора. Когда тендер удаляется из Excel (см.
# ScraperRunner._sweep_existing_tenders - завершённые/отменённые тендеры),
# соответствующая отметка тоже удаляется отсюда, чтобы файл не рос бесконечно.

_viewed_lock = threading.Lock()
_viewed_ids_cache = None


def _load_viewed_ids_from_disk() -> set:
    if not os.path.exists(VIEWED_FILE):
        return set()
    try:
        with open(VIEWED_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("viewed_ids", []))
    except Exception:
        return set()


def _save_viewed_ids_to_disk(ids: set):
    with open(VIEWED_FILE, "w", encoding="utf-8") as f:
        json.dump({"viewed_ids": sorted(ids)}, f, ensure_ascii=False, indent=2)


def get_viewed_ids() -> set:
    global _viewed_ids_cache
    with _viewed_lock:
        if _viewed_ids_cache is None:
            _viewed_ids_cache = _load_viewed_ids_from_disk()
        return set(_viewed_ids_cache)


def set_viewed(tender_id: str, viewed: bool):
    global _viewed_ids_cache
    if not tender_id:
        return get_viewed_ids()
    with _viewed_lock:
        if _viewed_ids_cache is None:
            _viewed_ids_cache = _load_viewed_ids_from_disk()
        if viewed:
            _viewed_ids_cache.add(tender_id)
        else:
            _viewed_ids_cache.discard(tender_id)
        _save_viewed_ids_to_disk(_viewed_ids_cache)
        return set(_viewed_ids_cache)


def remove_viewed(tender_id: str):
    """Убирает отметку 'просмотрено', если она была - вызывается при
    удалении тендера из Excel (тендер завершён/отменён), чтобы файл
    viewed_tenders.json не накапливал id тендеров, которых больше нет
    в таблице."""
    global _viewed_ids_cache
    if not tender_id:
        return
    with _viewed_lock:
        if _viewed_ids_cache is None:
            _viewed_ids_cache = _load_viewed_ids_from_disk()
        if tender_id in _viewed_ids_cache:
            _viewed_ids_cache.discard(tender_id)
            _save_viewed_ids_to_disk(_viewed_ids_cache)


# ============================================================================
# ================================ ИЗБРАННОЕ =================================
# ============================================================================
# В отличие от 'Просмотрено' (флаг рядом с записью в основном Excel),
# избранное хранит СНИМОК данных тендера на момент добавления - отдельно от
# goszakup_it_tenders*.xlsx. Это намеренно: когда тендер завершается и
# удаляется из основной таблицы (см. ScraperRunner._sweep_existing_tenders),
# избранное продолжает его показывать - "просроченные" из избранного не
# удаляются. remove_viewed() вызывается при удалении из основной таблицы,
# а вот избранное там НИКОГДА не трогается.

FAVORITES_SNAPSHOT_FIELDS = [
    "№ объявления", "Название лота/объявления", "Заказчик", "Организатор",
    "Город", "Сумма, тг.", "Статус", "Тип IT", "Приоритет",
    "Начало приёма заявок", "Окончание приёма заявок", "Ссылка на объявление",
    "Документы (названия и ссылки)",
]
FAVORITES_XLSX_HEADERS = FAVORITES_SNAPSHOT_FIELDS + ["Заметка", "Добавлено"]

_favorites_lock = threading.Lock()
_favorites_cache = None


def _load_favorites_from_disk() -> dict:
    if not os.path.exists(FAVORITES_FILE):
        return {}
    try:
        with open(FAVORITES_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return {}


def _regenerate_favorites_xlsx(data: dict):
    """Полностью перестраивает izbrannye_tendery.xlsx из текущего
    favorites.json - вызывается при каждом изменении (добавление, удаление,
    правка заметки), так что xlsx-файл на диске всегда актуален и его
    можно в любой момент открыть напрямую, не только через кнопку
    'Скачать Excel' в интерфейсе."""
    try:
        wb = Workbook()
        ws = wb.active
        ws.title = "Избранное"
        ws.append(FAVORITES_XLSX_HEADERS)
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        for col_idx in range(1, len(FAVORITES_XLSX_HEADERS) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center", wrap_text=True)

        url_col_idx = FAVORITES_XLSX_HEADERS.index("Ссылка на объявление") + 1
        entries = sorted(data.values(), key=lambda e: e.get("Добавлено", ""), reverse=True)
        for entry in entries:
            ws.append([entry.get(h, "") for h in FAVORITES_XLSX_HEADERS])
            url_val = entry.get("Ссылка на объявление", "")
            if url_val:
                cell = ws.cell(row=ws.max_row, column=url_col_idx)
                cell.hyperlink = url_val
                cell.font = Font(color="0563C1", underline="single")

        widths = [16, 45, 35, 30, 14, 16, 24, 14, 14, 18, 18, 32, 40, 40, 18]
        for i, w in enumerate(widths, start=1):
            if i <= len(FAVORITES_XLSX_HEADERS):
                ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A2"
        wb.save(FAVORITES_XLSX)
    except Exception:
        log_error_traceback("regenerate favorites xlsx")


def _save_favorites_to_disk(data: dict):
    with open(FAVORITES_FILE, "w", encoding="utf-8") as f:
        json.dump(data, f, ensure_ascii=False, indent=2)
    _regenerate_favorites_xlsx(data)


def get_favorites() -> dict:
    global _favorites_cache
    with _favorites_lock:
        if _favorites_cache is None:
            _favorites_cache = _load_favorites_from_disk()
        return json.loads(json.dumps(_favorites_cache))  # deep copy for safe read


def get_favorite_ids() -> set:
    return set(get_favorites().keys())


def add_favorite(tender_id: str, row: dict) -> dict:
    global _favorites_cache
    if not tender_id:
        return get_favorites()
    with _favorites_lock:
        if _favorites_cache is None:
            _favorites_cache = _load_favorites_from_disk()
        existing = _favorites_cache.get(tender_id, {})
        entry = {key: row.get(key, "") for key in FAVORITES_SNAPSHOT_FIELDS}
        entry["Заметка"] = existing.get("Заметка", "")
        entry["Добавлено"] = existing.get("Добавлено") or time.strftime("%Y-%m-%d %H:%M:%S")
        _favorites_cache[tender_id] = entry
        _save_favorites_to_disk(_favorites_cache)
        return json.loads(json.dumps(_favorites_cache))


def remove_favorite(tender_id: str) -> dict:
    global _favorites_cache
    with _favorites_lock:
        if _favorites_cache is None:
            _favorites_cache = _load_favorites_from_disk()
        if tender_id in _favorites_cache:
            del _favorites_cache[tender_id]
            _save_favorites_to_disk(_favorites_cache)
        return json.loads(json.dumps(_favorites_cache))


def set_favorite_note(tender_id: str, note: str) -> dict:
    global _favorites_cache
    with _favorites_lock:
        if _favorites_cache is None:
            _favorites_cache = _load_favorites_from_disk()
        if tender_id in _favorites_cache:
            _favorites_cache[tender_id]["Заметка"] = note
            _save_favorites_to_disk(_favorites_cache)
        return json.loads(json.dumps(_favorites_cache))


# ============================================================================
# ================================ ПОТОКИ ====================================
# ============================================================================

_thread_local = threading.local()


def get_thread_session(base_session: requests.Session) -> requests.Session:
    sess = getattr(_thread_local, "session", None)
    if sess is None:
        sess = requests.Session()
        sess.headers.update(dict(base_session.headers))
        _thread_local.session = sess
    sess.cookies.update(base_session.cookies.get_dict())
    return sess


def fetch_html(session: requests.Session, url: str, attempts: int = 6) -> str:
    last_exc = None
    for attempt in range(attempts):
        try:
            resp = session.get(url, timeout=REQUEST_TIMEOUT_SEC)
            resp.raise_for_status()
            resp.encoding = resp.encoding or "utf-8"
            return resp.text
        except Exception as e:
            last_exc = e
            if attempt < attempts - 1:
                wait_sec = 5 * (attempt + 1)
                log(f"    [!] Сетевая заминка ({e.__class__.__name__}), жду {wait_sec} сек "
                    f"и пробую снова (попытка {attempt + 2}/{attempts})...")
                time.sleep(wait_sec)
    raise last_exc


def process_tender(session: requests.Session, row: dict):
    if "заверш" in row.get("status", "").lower():
        return None
    thread_session = get_thread_session(session)

    polite_sleep(DETAIL_PAGE_DELAY_SEC)
    html = fetch_html(thread_session, row["url"])
    dump_debug_html(f"detail_{row['id']}.html", html)

    detail = parse_detail_page(html, row["url"])

    if is_license_renewal(
        row.get("title", ""), detail.get("description", ""),
        " ".join(d["name"] for d in detail["documents"]),
    ):
        return None

    doc_display = [f"{d['name']} ({d['url']})" for d in detail["documents"]]
    doc_names = " ".join(d["name"] for d in detail["documents"])

    is_it, score, matched_keywords = classify_is_it(
        row.get("title", ""), detail.get("description", ""), doc_names
    )

    if not is_it and SCAN_ATTACHED_DOCUMENTS and detail["documents"]:
        cookies = thread_session.cookies.get_dict()
        for doc in detail["documents"][:MAX_DOCS_TO_SCAN_IF_NO_MATCH]:
            text = get_document_text(doc["url"], thread_session, cookies)
            if text:
                is_it, score, matched_keywords = classify_is_it(doc["name"], text, doc["name"])
                if is_it:
                    break

    ai_settings = get_ai_settings()
    ai_enabled = bool(ai_settings.get("enabled")) and ai_provider_ready(ai_settings.get("provider"))
    ai_reasoning = ""

    if not is_it:
        in_borderline_zone = BORDERLINE_SCORE_MIN <= score < IT_SCORE_THRESHOLD
        if ai_enabled and in_borderline_zone:
            # Ключевые слова почти нашли (пограничный балл), но не дотянули -
            # спрашиваем у ИИ, подходит ли тендер под свободно описанный
            # критерий (устойчиво к падежам/синонимам, в отличие от
            # сопоставления подстрок).
            ai_result = ai_check_relevance(
                row.get("title", ""), detail.get("description", ""),
                ai_settings["criteria"], ai_settings["model"], ai_settings["provider"],
            )
            if ai_result and ai_result["relevant"]:
                is_it = True
                matched_keywords = list(matched_keywords) + ["[ИИ] релевантно по критерию"]
                ai_reasoning = ai_result["reasoning"]
        if not is_it:
            if in_borderline_zone:
                log_borderline_candidate(row, score, matched_keywords)
            return None
    elif ai_enabled and ai_settings.get("verify_mode") == "all":
        # Режим "перепроверять всё": ИИ может отклонить то, что уже прошло
        # по ключевым словам, если по смыслу это не соответствует критерию
        # (ложное срабатывание). Отклонённое уходит в пограничные, а не
        # молча теряется - можно посмотреть и поправить критерий/слова.
        ai_result = ai_check_relevance(
            row.get("title", ""), detail.get("description", ""),
            ai_settings["criteria"], ai_settings["model"], ai_settings["provider"],
        )
        if ai_result:
            if ai_result["relevant"]:
                ai_reasoning = ai_result["reasoning"]
            else:
                log_borderline_candidate(row, score, matched_keywords + ["[ИИ отклонил]"])
                return None

    it_type = detect_it_type(row.get("title", ""), detail.get("description", ""), doc_names)
    priority = detect_priority(row.get("title", ""), detail.get("description", ""), doc_names)

    return {
        "number_anno": row.get("number_anno", ""), "title": row.get("title", ""),
        "customer_name": row.get("customer_name") or row.get("organizer_name", ""),
        "customer_bin": detail.get("customer_bin", ""), "organizer_name": row.get("organizer_name", ""),
        "method": row.get("method", ""), "amount": row.get("amount", ""),
        "status": row.get("status", ""), "start_date": row.get("start_date", ""),
        "end_date": row.get("end_date", ""), "url": row.get("url", ""),
        "matched_keywords": matched_keywords, "documents": doc_display,
        "description": detail.get("description", ""), "it_type": it_type, "priority": priority,
        "ai_reasoning": ai_reasoning,
        "is_astana": is_astana_tender(row, detail),
    }


def browser_keepalive_loop(driver, base_url: str, stop_event: threading.Event, interval_sec: int = 180):
    while not stop_event.wait(interval_sec):
        try:
            driver.execute_script(
                "fetch(arguments[0], {credentials: 'include', cache: 'no-store'}).catch(function(){});",
                base_url,
            )
        except Exception:
            pass


# ============================================================================
# ================== СОХРАНЕНИЕ СТАТИСТИКИ СБОРА НА ДИСК =====================
# ============================================================================
# Раньше счётчики (проверено/найдено/ошибок и т.д.) жили только в памяти
# процесса Flask - если приложение (python app.py) перезапускалось, вкладка
# "Сбор данных" после этого показывала одни нули, хотя по факту предыдущий
# сбор прошёл успешно. Теперь после каждого обновления счётчиков состояние
# сохраняется в run_stats.json и подгружается заново при следующем запуске
# app.py - статистика последнего сбора остаётся видна всегда, даже если
# ничего сейчас не запущено.

_run_stats_lock = threading.Lock()


def _load_run_stats():
    if not os.path.exists(RUN_STATS_FILE):
        return None
    try:
        with open(RUN_STATS_FILE, "r", encoding="utf-8") as f:
            return json.load(f)
    except Exception:
        return None


def _save_run_stats(counts: dict, state: str, finished_at: str = None):
    try:
        with _run_stats_lock:
            with open(RUN_STATS_FILE, "w", encoding="utf-8") as f:
                json.dump({
                    "counts": counts,
                    "state": state,
                    "finished_at": finished_at,
                }, f, ensure_ascii=False, indent=2)
    except Exception:
        pass


# ============================================================================
# ========================== STATE MACHINE ДЛЯ ВЕБ-UI ========================
# ============================================================================

class ScraperRunner:
    """Управляет жизненным циклом одного "сеанса" сбора данных, вызывается
    из app.py по кнопкам в браузере вместо консольных input()/print()."""

    def __init__(self):
        self.lock = threading.Lock()
        self.state = "idle"          # idle|launching|waiting_tab|ready|running|stopping|stopped|error
        self.driver = None
        self.base_url = None
        self.thread = None
        self.stop_event = threading.Event()
        self.keepalive_stop = threading.Event()
        self.counts = {
            "checked": 0, "found_it": 0, "found_astana": 0,
            "errors": 0, "skipped_duplicates": 0, "page": 0,
            "removed_completed": 0, "updated_statuses": 0,
        }
        self.error_message = None
        self.last_finished_at = None
        self.auto_restart = True
        self.auto_restart_count = 0
        self._restart_timer = None

        saved = _load_run_stats()
        if saved and isinstance(saved.get("counts"), dict):
            self.counts.update({k: v for k, v in saved["counts"].items() if k in self.counts})
            self.last_finished_at = saved.get("finished_at")

    def snapshot(self):
        with self.lock:
            return {
                "state": self.state,
                "base_url": self.base_url,
                "counts": dict(self.counts),
                "error": self.error_message,
                "selenium_available": SELENIUM_AVAILABLE,
                "last_finished_at": self.last_finished_at,
                "auto_restart": self.auto_restart,
                "auto_restart_count": self.auto_restart_count,
                "restart_pending": self._restart_timer is not None,
            }

    def set_auto_restart(self, enabled: bool):
        with self.lock:
            self.auto_restart = bool(enabled)
            if not self.auto_restart and self._restart_timer is not None:
                self._restart_timer.cancel()
                self._restart_timer = None
        log(f"Автоматический перезапуск при сбое: {'включён' if enabled else 'выключен'}.")

    def set_search_url(self, url: str):
        """Новый, основной способ задать URL с фильтрами - БЕЗ браузера.

        Раньше сайт требовал авторизацию (через v3bl.goszakup.gov.kz), и
        URL с фильтрами приходилось вытаскивать из открытой вкладки Edge
        после ручного логина (см. launch_browser/confirm_ready ниже - эти
        методы оставлены как резервный способ, но больше не используются
        по умолчанию). Теперь обычный goszakup.gov.kz отдаёт результаты
        поиска анонимным (не авторизованным) запросам без каких-либо кук -
        достаточно скопировать ссылку с фильтрами прямо из адресной строки
        браузера и вставить её в интерфейс."""
        url = (url or "").strip()
        if not url:
            return False, "Пустой URL"
        if "goszakup.gov.kz" not in url:
            return False, "Похоже, это не ссылка на goszakup.gov.kz"
        with self.lock:
            if self.state not in ("idle", "ready", "stopped", "error"):
                return False, f"Нельзя менять URL в состоянии '{self.state}'."
            self.base_url = url
            self.state = "ready"
            self.error_message = None
        log(f"URL с фильтрами сохранён: {url}")
        if "search/announce" not in url and "announce" not in url:
            log("[!] Похоже, это не страница результатов поиска — проверьте ссылку.")
        return True, url

    def launch_browser(self):
        """РЕЗЕРВНЫЙ способ (не используется по умолчанию) - на случай,
        если goszakup.gov.kz снова потребует авторизацию для поиска, как
        было раньше через v3bl.goszakup.gov.kz. Открывает окно Edge для
        ручного логина, дальше используется confirm_ready()."""
        if not SELENIUM_AVAILABLE:
            self.state = "error"
            self.error_message = "Модуль selenium не установлен (pip install selenium)."
            log("[!] selenium не установлен.")
            return False
        with self.lock:
            self.state = "launching"
        ok, reason = try_auto_launch_edge()
        if not ok and reason != "already_running":
            with self.lock:
                self.state = "error"
                self.error_message = f"Не удалось открыть Edge ({reason})."
            return False
        try:
            self.driver = attach_to_running_edge()
        except Exception as e:
            with self.lock:
                self.state = "error"
                self.error_message = f"Не удалось подключиться к Edge: {e}"
            log(f"[!] Не удалось подключиться к Edge: {e}")
            return False
        with self.lock:
            self.state = "waiting_tab"
        log("Подключение к Edge успешно. Авторизуйтесь на goszakup.gov.kz, настройте "
            "фильтры и нажмите 'Найти', затем нажмите 'Я готов' в интерфейсе.")
        return True

    def confirm_ready(self):
        """Вызывается по кнопке 'Я готов' — ищет среди вкладок ту, что
        относится к goszakup.gov.kz, и запоминает её как базовый URL."""
        if self.driver is None:
            self.error_message = "Браузер ещё не запущен."
            return False, self.error_message
        base_url = find_and_switch_to_goszakup_tab(self.driver)
        if base_url is None:
            msg = ("Не нашёл вкладку goszakup.gov.kz среди открытых. Убедитесь, что "
                   "вкладка с результатами поиска открыта в этом же окне Edge.")
            log(f"[!] {msg}")
            return False, msg
        with self.lock:
            self.base_url = base_url
            self.state = "ready"
        log(f"Базовый URL с фильтрами: {base_url}")
        if "search/announce" not in base_url and "announce" not in base_url:
            log("[!] Похоже, это не страница результатов поиска — проверьте фильтры.")
        return True, base_url

    def start_run(self):
        with self.lock:
            if self.state not in ("ready", "stopped", "error"):
                return False, f"Нельзя запустить сбор в состоянии '{self.state}'."
            if self.base_url is None:
                return False, "Сначала укажите ссылку с фильтрами."
            self.state = "running"
            self.error_message = None
            # Новая сессия сбора - обнуляем счётчики "Статистика сессии".
            # Итоги ПРЕДЫДУЩЕЙ сессии при этом не теряются: они уже
            # сохранены в run_stats.json (см. _save_run_stats в finally
            # блоке _run_loop) и доступны через last_finished_at, пока не
            # завершится эта новая сессия.
            for key in self.counts:
                self.counts[key] = 0
        self.stop_event = threading.Event()
        self.keepalive_stop = threading.Event()
        self.thread = threading.Thread(target=self._run_loop, daemon=True)
        self.thread.start()
        return True, "started"

    def stop_run(self):
        with self.lock:
            if self._restart_timer is not None:
                self._restart_timer.cancel()
                self._restart_timer = None
                if self.state != "running":
                    log("Запланированный автоматический перезапуск отменён пользователем.")
                    self.state = "stopped"
                    return True, "auto-restart cancelled"
            if self.state != "running":
                return False, f"Сбор не запущен (состояние: {self.state})."
            self.state = "stopping"
        log("Останавливаю сбор по команде из интерфейса...")
        self.stop_event.set()
        return True, "stopping"

    def _sweep_existing_tenders(self, session: requests.Session):
        """Перед началом нового сбора перепроверяет статус уже сохранённых
        (в предыдущих запусках) тендеров в обоих Excel-файлах: тендеры,
        ставшие 'Завершено'/'Отменено', удаляются из таблицы (и из отметки
        'Просмотрено', если она была) - чтобы таблица не засорялась старыми
        закрытыми закупками. У остальных заодно обновляется столбец
        'Статус' на актуальный, если он изменился.

        Делает по одному GET-запросу на каждый ещё сохранённый тендер
        (параллельно, DETAIL_WORKERS потоков) - при большом количестве уже
        накопленных строк это может занять заметное время; прогресс виден
        в живом логе."""
        status_col = HEADERS.index("Статус") + 1
        url_col = HEADERS.index("Ссылка на объявление") + 1

        for sink in (sink_main, sink_astana):
            if self.stop_event.is_set():
                return

            with sink.lock:
                snapshot = []
                for row_idx in range(2, sink.worksheet.max_row + 1):
                    url_val = sink.worksheet.cell(row=row_idx, column=url_col).value
                    status_val = sink.worksheet.cell(row=row_idx, column=status_col).value
                    if url_val:
                        snapshot.append((row_idx, str(url_val), status_val or ""))

            if not snapshot:
                continue

            log(f"Проверяю статусы {len(snapshot)} уже сохранённых тендеров "
                f"в {os.path.basename(sink.path)}...")

            def check_one(url):
                try:
                    thread_session = get_thread_session(session)
                    html = fetch_html(thread_session, url, attempts=3)
                    detail = parse_detail_page(html, url)
                    return _extract_detail_status(detail.get("description", ""))
                except Exception:
                    return None

            new_statuses = {}
            executor = ThreadPoolExecutor(max_workers=DETAIL_WORKERS)
            future_map = {executor.submit(check_one, item[1]): item for item in snapshot}
            try:
                for future in as_completed(future_map):
                    if self.stop_event.is_set():
                        break
                    row_idx = future_map[future][0]
                    try:
                        new_statuses[row_idx] = future.result()
                    except Exception:
                        new_statuses[row_idx] = None
            finally:
                # wait=False + cancel_futures - не ждём завершения ещё не
                # стартовавших задач (а их могут быть сотни при большой
                # таблице), чтобы "Остановить" в интерфейсе срабатывало
                # быстро, а не только после того как довыполнится вся
                # текущая пачка запросов.
                executor.shutdown(wait=False, cancel_futures=True)

            removed_here = 0
            updated_here = 0
            rows_to_delete = []
            with sink.lock:
                for row_idx, url_val, old_status in snapshot:
                    new_status = new_statuses.get(row_idx)
                    if not new_status:
                        continue
                    if new_status != old_status:
                        sink.worksheet.cell(row=row_idx, column=status_col).value = new_status
                        updated_here += 1
                    if "заверш" in new_status.lower() or "отмен" in new_status.lower():
                        rows_to_delete.append((row_idx, url_val))

                for row_idx, url_val in sorted(rows_to_delete, key=lambda x: x[0], reverse=True):
                    sink.worksheet.delete_rows(row_idx)
                    tender_id = _extract_tender_id_from_url(url_val)
                    if tender_id:
                        sink.existing_ids.discard(tender_id)
                        remove_viewed(tender_id)
                    removed_here += 1

            if removed_here or updated_here:
                sink.save_now()

            with self.lock:
                self.counts["removed_completed"] += removed_here
                self.counts["updated_statuses"] += updated_here

            log(f"    {os.path.basename(sink.path)}: обновлено статусов {updated_here}, "
                f"удалено завершённых/отменённых {removed_here}.")

    def _run_loop(self):
        driver = self.driver
        base_url = self.base_url
        seen_ids = set()
        page = 1
        try:
            _apply_active_keywords()
            seen_ids = load_seen_ids()
            log(f"Уже обработано в предыдущих запусках: {len(seen_ids)} объявлений.")
            init_workbooks_for_run()

            session = requests.Session()
            if driver is not None:
                # Резервный путь через Edge (см. launch_browser) - куки из
                # авторизованной сессии браузера синхронизируются на каждой
                # странице. ВАЖНО: если гостевой доступ (см. ветку else)
                # работает - используйте его, не этот путь. На практике
                # куки из браузера (даже с пустой/гостевой сессией) иногда
                # заставляют сайт отдавать урезанный/другой ответ, чем
                # совсем анонимный запрос без кук вообще - веб-сервер,
                # похоже, как-то по-своему интерпретирует такие куки.
                try:
                    ua = driver.execute_script("return navigator.userAgent;")
                except Exception:
                    ua = "Mozilla/5.0"
                session.headers.update({
                    "User-Agent": ua, "Referer": base_url, "Accept-Language": "ru-RU,ru;q=0.9",
                })
                session.cookies.update(get_cookies_dict(driver))
            else:
                # Основной путь (без браузера): goszakup.gov.kz отдаёт
                # результаты поиска анонимным запросам без каких-либо кук -
                # заголовки ниже подобраны и проверены как рабочие для
                # такого анонимного доступа. НЕ добавляйте сюда Cookie -
                # именно отсутствие кук здесь и работает надёжно.
                session.headers.update({
                    "User-Agent": (
                        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) "
                        "AppleWebKit/537.36 (KHTML, like Gecko) "
                        "Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0"
                    ),
                    "Accept-Language": "ru-RU,ru;q=0.9",
                    "Accept": "text/html,application/xhtml+xml,application/xml;q=0.9,*/*;q=0.8",
                })

            if not self.stop_event.is_set():
                self._sweep_existing_tenders(session)
                _save_run_stats(dict(self.counts), "running", None)

            keepalive_thread = None
            if driver is not None:
                keepalive_thread = threading.Thread(
                    target=browser_keepalive_loop, args=(driver, base_url, self.keepalive_stop), daemon=True,
                )
                keepalive_thread.start()

            while page <= MAX_PAGES_SAFETY_LIMIT and not self.stop_event.is_set():
                with self.lock:
                    self.counts["page"] = page
                page_url = build_page_url(base_url, page)
                log(f"--- Страница {page}: {page_url}")

                if driver is not None:
                    session.cookies.update(get_cookies_dict(driver))

                html = None
                attempt_num = 0
                while html is None and not self.stop_event.is_set():
                    attempt_num += 1
                    try:
                        html = fetch_html(session, page_url)
                    except Exception as e:
                        log_error_traceback(f"list page fetch, page={page}, attempt={attempt_num}")
                        wait_sec = min(10 * attempt_num, 120)
                        log(f"    [!] Страница {page} не загружается ({e.__class__.__name__}), "
                            f"попытка {attempt_num}. Жду {wait_sec} сек...")
                        # ждём короткими интервалами, чтобы быстро реагировать на "стоп"
                        for _ in range(wait_sec):
                            if self.stop_event.is_set():
                                break
                            time.sleep(1)
                        if driver is not None:
                            session.cookies.update(get_cookies_dict(driver))

                if self.stop_event.is_set():
                    break

                polite_sleep(REQUEST_DELAY_SEC)
                dump_debug_html(f"list_page_{page}.html", html)

                rows = parse_list_page(html)
                empty_retry = 0
                while not rows and empty_retry < EMPTY_PAGE_MAX_RETRIES and not self.stop_event.is_set():
                    empty_retry += 1
                    wait_sec = min(15 * empty_retry, 90)
                    log(f"    [!] Страница {page} вернула 0 объявлений - возможно, временная заминка "
                        f"сайта, а не конец списка. Повторная проверка {empty_retry}/{EMPTY_PAGE_MAX_RETRIES} "
                        f"через {wait_sec} сек...")
                    for _ in range(wait_sec):
                        if self.stop_event.is_set():
                            break
                        time.sleep(1)
                    if self.stop_event.is_set():
                        break
                    try:
                        if driver is not None:
                            session.cookies.update(get_cookies_dict(driver))
                        html = fetch_html(session, page_url)
                        rows = parse_list_page(html)
                    except Exception:
                        log_error_traceback(f"empty page retry, page={page}, attempt={empty_retry}")
                        rows = []

                if not rows:
                    log("Больше объявлений действительно не найдено (после повторных проверок) - "
                        "это либо конец списка по вашим фильтрам, либо сбился фильтр/URL. Останавливаюсь.")
                    break

                new_rows = [
                    r for r in rows
                    if r["id"] not in seen_ids and r.get("status", "") != "Завершено"
                ]
                log(f"    Найдено на странице: {len(rows)}, новых: {len(new_rows)}")

                found_on_this_page = 0
                executor = ThreadPoolExecutor(max_workers=DETAIL_WORKERS)
                future_to_row = {executor.submit(process_tender, session, row): row for row in new_rows}
                try:
                    for future in as_completed(future_to_row):
                        if self.stop_event.is_set():
                            log("    Останавливаюсь по команде - прерываю обработку текущей "
                                "страницы (уже запущенные запросы доработают в фоне, новые не стартуют).")
                            break
                        row = future_to_row[future]
                        with self.lock:
                            self.counts["checked"] += 1
                        try:
                            result = future.result()
                        except Exception as e:
                            with self.lock:
                                self.counts["errors"] += 1
                            log(f"    [!] Ошибка при обработке {row['id']}: {e} (см. {ERROR_LOG_FILE})")
                            log_error_traceback(f"tender_id={row['id']} url={row.get('url', '')}")
                            result = None

                        with _seen_ids_lock:
                            seen_ids.add(row["id"])
                            if result:
                                save_seen_ids(seen_ids)

                        if result:
                            is_astana = bool(result.pop("is_astana", False))
                            was_added = append_it_result(result, astana=is_astana)
                            if was_added:
                                with self.lock:
                                    self.counts["found_it"] += 1
                                    if is_astana:
                                        self.counts["found_astana"] += 1
                                found_on_this_page += 1
                                tag = "IT/Астана" if is_astana else "IT"
                                log(f"    [{tag}] {row['number_anno']}: {row['title'][:70]}")
                            else:
                                with self.lock:
                                    self.counts["skipped_duplicates"] += 1
                                log(f"    [= дубль] {row['number_anno']}: {row['title'][:70]}")

                        if self.counts["checked"] % 20 == 0:
                            with _seen_ids_lock:
                                save_seen_ids(seen_ids)
                            _save_run_stats(dict(self.counts), "running", None)
                finally:
                    executor.shutdown(wait=False, cancel_futures=True)

                if new_rows and found_on_this_page == 0:
                    log(f"    На странице {page} IT-тендеров не найдено.")

                with _seen_ids_lock:
                    save_seen_ids(seen_ids)
                page += 1

        except Exception:
            log("Непредвиденная ошибка в цикле сбора:")
            log(traceback.format_exc())
            log_error_traceback("main loop")
            with self.lock:
                self.error_message = "Непредвиденная ошибка — см. лог."
        finally:
            self.keepalive_stop.set()
            with _seen_ids_lock:
                save_seen_ids(seen_ids)
            save_workbooks_now()
            should_auto_restart = False
            with self.lock:
                self.state = "stopped" if self.error_message is None else "error"
                self.last_finished_at = time.strftime("%Y-%m-%d %H:%M:%S")
                _save_run_stats(dict(self.counts), self.state, self.last_finished_at)
                # Авто-перезапуск только если сбор реально прервался с
                # ошибкой (не был штатно завершён - страницы кончились или
                # достигнут лимит - и не был остановлен пользователем
                # кнопкой "Остановить", что всегда выставляет stop_event).
                if self.auto_restart and self.state == "error" and not self.stop_event.is_set():
                    should_auto_restart = True
            log("=" * 50)
            log(f"Проверено: {self.counts['checked']}, найдено IT: {self.counts['found_it']} "
                f"(из них Астана: {self.counts['found_astana']}), ошибок: {self.counts['errors']}")
            log(f"Удалено завершённых/отменённых из таблицы: {self.counts['removed_completed']}, "
                f"обновлено статусов: {self.counts['updated_statuses']}")
            if should_auto_restart:
                with self.lock:
                    self.auto_restart_count += 1
                    attempt_no = self.auto_restart_count
                delay = 30
                log(f"[авто-перезапуск] Сбор прервался с ошибкой. Перезапускаю через "
                    f"{delay} сек (попытка №{attempt_no}). Чтобы отменить — нажмите «Остановить».")
                timer = threading.Timer(delay, self._trigger_auto_restart)
                timer.daemon = True
                with self.lock:
                    self._restart_timer = timer
                timer.start()
            else:
                log("Сбор остановлен.")

    def _trigger_auto_restart(self):
        with self.lock:
            self._restart_timer = None
            if self.state != "error":
                # Состояние успели поменять вручную (например, отмена) -
                # автоматически перезапускать больше не нужно.
                return
        ok, info = self.start_run()
        if not ok:
            log(f"[авто-перезапуск] Не удалось перезапустить: {info}")


runner = ScraperRunner()


# ============================================================================
# ============ ИИ-ПЕРЕСМОТР УЖЕ СОБРАННЫХ ДАННЫХ (по кнопке) =================
# ============================================================================
# В отличие от ai_check_relevance() внутри process_tender (который смотрит
# на КАЖДЫЙ тендер в реальном времени во время сбора), это - разовое
# ручное действие: "прогнать текущий критерий по тому, что уже лежит в
# Тендерах/Пограничных". Полезно, например, после того как критерий был
# изменён/уточнён и хочется пересмотреть уже накопленное, не запуская
# новый полный сбор.
#
#   - "Пограничные": то, что ИИ сочтёт релевантным - ПЕРЕНОСИТСЯ в основную
#     таблицу (убирается из пограничных). Это безопасно - тендер просто
#     "находится", а не удаляется.
#   - "Тендеры": ничего не удаляется автоматически - несоответствующее
#     критерию только ПОМЕЧАЕТСЯ (столбец "Обоснование ИИ" + пометка в
#     ключевых словах), решение об удалении остаётся за пользователем
#     (кнопка "Удалить" уже есть в интерфейсе).

_ai_rescan_lock = threading.Lock()
_ai_rescan_state = {
    "running": False, "checked": 0, "total": 0,
    "promoted": 0, "flagged": 0, "failed": 0, "finished_at": None, "error": None,
}
_ai_rescan_warned = False


def get_ai_rescan_state() -> dict:
    with _ai_rescan_lock:
        return dict(_ai_rescan_state)


def _note_ai_rescan_failure(ai_result):
    """Вызывается после каждого ai_check_relevance() внутри пересмотра.
    Если проверка не удалась (ai_result is None) - считает это в счётчик
    'failed' и ОДИН раз (не на каждую из потенциально сотен строк) выводит
    в видимый лог настоящую причину - чтобы 'проверено 70, везде 0' не
    выглядело загадочно, а сразу было понятно, что дело в ключе/модели/
    лимите, а не в том, что просто ничего не подошло."""
    global _ai_rescan_warned
    if ai_result is not None:
        return
    with _ai_rescan_lock:
        _ai_rescan_state["failed"] += 1
        already_warned = _ai_rescan_warned
        _ai_rescan_warned = True
    if not already_warned:
        reason = _last_ai_error or "причина неизвестна"
        log(f"[!] Проверка ИИ не удалась: {reason}")
        log("[!] Проверьте ключ/модель на вкладке «ИИ-проверка» — дальнейшие "
            "такие же сбои в этом прогоне будут просто считаться (см. «не удалось» в статистике).")


def start_ai_rescan(targets):
    """targets: множество/список из 'borderline' и/или 'tenders'."""
    settings = get_ai_settings()
    if not settings.get("enabled"):
        return False, "ИИ-проверка выключена — включите её на вкладке «ИИ-проверка»."
    if not ai_provider_ready(settings.get("provider")):
        return False, "Провайдер ИИ не готов (нет ключа) — проверьте вкладку «ИИ-проверка»."
    targets = set(t for t in targets if t in ("borderline", "tenders"))
    if not targets:
        return False, "Не выбрано, что пересматривать."

    with _ai_rescan_lock:
        if _ai_rescan_state["running"]:
            return False, "Проверка уже идёт."
        _ai_rescan_state.update({
            "running": True, "checked": 0, "total": 0,
            "promoted": 0, "flagged": 0, "failed": 0, "finished_at": None, "error": None,
        })
        global _ai_rescan_warned
        _ai_rescan_warned = False

    thread = threading.Thread(target=_ai_rescan_worker, args=(targets, settings), daemon=True)
    thread.start()
    return True, "started"


def _rescan_session() -> requests.Session:
    session = requests.Session()
    session.headers.update({
        "User-Agent": (
            "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
            "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0"
        ),
        "Accept-Language": "ru-RU,ru;q=0.9",
    })
    return session


def _ai_rescan_worker(targets, settings):
    log("=" * 50)
    log(f"[ИИ-пересмотр] Начинаю пересмотр уже собранных данных ({', '.join(sorted(targets))})...")
    session = _rescan_session()
    try:
        if "borderline" in targets:
            _ai_rescan_borderline(session, settings)
        if "tenders" in targets:
            _ai_rescan_tenders(session, settings)
    except Exception as e:
        log_error_traceback(f"ai rescan worker: {e}")
        with _ai_rescan_lock:
            _ai_rescan_state["error"] = str(e)
    finally:
        with _ai_rescan_lock:
            _ai_rescan_state["running"] = False
            _ai_rescan_state["finished_at"] = time.strftime("%Y-%m-%d %H:%M:%S")
            snap = dict(_ai_rescan_state)
        log(f"[ИИ-пересмотр] Готово. Проверено: {snap['checked']}, "
            f"перенесено из пограничных: {snap['promoted']}, помечено в тендерах: {snap['flagged']}, "
            f"не удалось проверить: {snap['failed']}.")


def _ai_rescan_borderline(session: requests.Session, settings: dict):
    if not os.path.exists(BORDERLINE_LOG_FILE):
        log("[ИИ-пересмотр] Пограничных кандидатов нет.")
        return
    with open(BORDERLINE_LOG_FILE, "r", encoding="utf-8") as f:
        raw_lines = f.readlines()

    entries = []
    for line in raw_lines:
        stripped = line.rstrip("\n")
        if not stripped:
            continue
        parts = stripped.split("\t")
        if len(parts) < 5:
            continue
        ts, score_part, number_anno, url, title = parts[:5]
        extra = {}
        for p in parts[5:]:
            if "=" in p:
                k, _, v = p.partition("=")
                extra[k] = v
        entries.append({
            "line": line, "number_anno": number_anno, "url": url, "title": title,
            "amount": extra.get("amount", ""), "start_date": extra.get("start_date", ""),
            "end_date": extra.get("end_date", ""), "keywords": extra.get("keywords", ""),
        })

    if not entries:
        log("[ИИ-пересмотр] Пограничных кандидатов нет.")
        return

    log(f"[ИИ-пересмотр] Пограничных к проверке: {len(entries)}")
    with _ai_rescan_lock:
        _ai_rescan_state["total"] += len(entries)

    kept_lines = []
    for entry in entries:
        with _ai_rescan_lock:
            _ai_rescan_state["checked"] += 1

        description = ""
        try:
            html = fetch_html(session, entry["url"], attempts=2)
            detail = parse_detail_page(html, entry["url"])
            description = detail.get("description", "")
        except Exception:
            pass  # если карточка не открылась - проверяем хотя бы по названию

        ai_result = ai_check_relevance(
            entry["title"], description, settings["criteria"], settings["model"], settings["provider"],
        )
        _note_ai_rescan_failure(ai_result)
        polite_sleep((1.0, 1.5))  # не долбить API чаще лимита бесплатного тарифа

        if ai_result and ai_result["relevant"]:
            status = _extract_detail_status(description) if description else ""
            is_astana = bool(description) and ASTANA_CITY_RE.search(description[:4000]) is not None
            keywords = [k.strip() for k in entry["keywords"].split(",") if k.strip()]
            keywords.append("[ИИ] релевантно по критерию (пересмотр)")
            row = {
                "number_anno": entry["number_anno"], "title": entry["title"],
                "customer_name": "", "organizer_name": "", "method": "",
                "amount": entry["amount"], "status": status,
                "start_date": entry["start_date"], "end_date": entry["end_date"],
                "url": entry["url"], "matched_keywords": keywords, "documents": [],
                "description": description,
                "it_type": detect_it_type(entry["title"], description, ""),
                "priority": detect_priority(entry["title"], description, ""),
                "ai_reasoning": ai_result["reasoning"],
            }
            was_added = append_it_result(row, astana=is_astana)
            if was_added:
                with _ai_rescan_lock:
                    _ai_rescan_state["promoted"] += 1
                tag = "IT/Астана" if is_astana else "IT"
                log(f"    [{tag}] {entry['number_anno']}: {entry['title'][:70]} (перенесён из пограничных)")
            # тендер найден и (если новый) уже записан - не сохраняем строку
            # обратно в borderline, независимо от was_added (если он там
            # почему-то уже был - дублировать в логе всё равно незачем)
        else:
            kept_lines.append(entry["line"])

    with open(BORDERLINE_LOG_FILE, "w", encoding="utf-8") as f:
        f.writelines(kept_lines)


def _ai_rescan_tenders(session: requests.Session, settings: dict):
    title_col = HEADERS.index("Название лота/объявления") + 1
    desc_col = HEADERS.index("Краткое описание/ТЗ (выдержка)") + 1
    ai_col = HEADERS.index("Обоснование ИИ") + 1
    keywords_col = HEADERS.index("Совпавшие IT-ключевые слова") + 1

    for sink in (sink_main, sink_astana):
        with sink.lock:
            if sink.worksheet is None:
                sink.init_for_run()
            snapshot = []
            for row_idx in range(2, sink.worksheet.max_row + 1):
                title = sink.worksheet.cell(row=row_idx, column=title_col).value
                if not title:
                    continue
                desc = sink.worksheet.cell(row=row_idx, column=desc_col).value or ""
                snapshot.append((row_idx, title, desc))

        if not snapshot:
            continue

        log(f"[ИИ-пересмотр] {os.path.basename(sink.path)}: проверяю {len(snapshot)} записей...")
        with _ai_rescan_lock:
            _ai_rescan_state["total"] += len(snapshot)

        changed = False
        for row_idx, title, desc in snapshot:
            with _ai_rescan_lock:
                _ai_rescan_state["checked"] += 1

            ai_result = ai_check_relevance(
                title, desc, settings["criteria"], settings["model"], settings["provider"],
            )
            _note_ai_rescan_failure(ai_result)
            polite_sleep((1.0, 1.5))  # не долбить API чаще лимита бесплатного тарифа
            if not ai_result:
                continue  # проверка не удалась - строку не трогаем

            with sink.lock:
                sink.worksheet.cell(row=row_idx, column=ai_col).value = ai_result["reasoning"]
                if not ai_result["relevant"]:
                    old_kw = sink.worksheet.cell(row=row_idx, column=keywords_col).value or ""
                    marker = "[ИИ: не соответствует критерию]"
                    if marker not in old_kw:
                        new_kw = f"{old_kw}, {marker}".strip(", ")
                        sink.worksheet.cell(row=row_idx, column=keywords_col).value = new_kw
            changed = True

            if ai_result and not ai_result["relevant"]:
                with _ai_rescan_lock:
                    _ai_rescan_state["flagged"] += 1
                log(f"    [!] {title[:60]} — ИИ считает несоответствующим критерию (не удалено, только помечено)")

        if changed:
            sink.save_now()
