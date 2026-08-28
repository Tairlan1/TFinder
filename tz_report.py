# -*- coding: utf-8 -*-
"""
tz_report.py — отдельный инструмент (не часть веб-приложения).

Берёт Excel-файл, выгруженный веб-интерфейсом (goszakup_it_tenders.xlsx,
goszakup_it_tenders_astana.xlsx или izbrannye_tendery.xlsx — во всех трёх
есть столбец "Документы (названия и ссылки)"), и для каждого тендера:

  1. Разбирает список вложений и определяет среди них ТЗ (техническое
     задание) — по названию файла, а если явного совпадения нет — берёт
     наиболее вероятный документ (pdf/doc/docx, не смета/чертёж/протокол).
  2. Скачивает его в папку рядом (по умолчанию ./ТЗ) и добавляет в Excel
     столбец "Файл ТЗ" — гиперссылку на скачанный файл.
  3. Если указан флаг --analyze — прогоняет текст ТЗ через Claude и
     записывает краткую выжимку "что конкретно нужно" в столбец
     "Анализ ТЗ (ИИ)".

БЕЗОПАСНОСТЬ ДАННЫХ: по умолчанию скрипт НЕ трогает исходный файл — делает
его копию "<имя>_с_ТЗ.xlsx" и работает с ней. Это важно, если веб-интерфейс
(app.py) может параллельно записывать в тот же исходный xlsx во время
сбора — тогда правки этого скрипта могли бы быть перезаписаны. Если вы
точно знаете, что сбор сейчас не идёт, можно передать --in-place.

УСТАНОВКА:
    pip install requests openpyxl pdfplumber python-docx
    # для --analyze дополнительно:
    pip install anthropic

ПРИМЕРЫ ЗАПУСКА:
    python tz_report.py goszakup_it_tenders.xlsx
    python tz_report.py goszakup_it_tenders.xlsx --analyze
    python tz_report.py izbrannye_tendery.xlsx --output-dir ТЗ_избранное --analyze
    python tz_report.py goszakup_it_tenders.xlsx --analyze --model claude-haiku-4-5-20251001
    python tz_report.py goszakup_it_tenders.xlsx --limit 5   # тестовый прогон на первых 5 строках

Ключ Anthropic для --analyze берётся из переменной окружения
ANTHROPIC_API_KEY (или передайте явно через --api-key).
"""

import argparse
import os
import re
import shutil
import sys
import time
import uuid

import requests
from openpyxl import load_workbook
from openpyxl.styles import Font

DOC_LINK_RE = re.compile(r"(.+?)\s*\((https?://[^\s)]+)\)")

# Слова, по которым определяем, что вложение — именно ТЗ (а не конкурсная
# документация целиком, смета, чертежи и т.п.).
TZ_NAME_HINTS = [
    "тз", "техническое задание", "тех.задание", "техзадание", "тех задание",
    "terms of reference", "technical specification",
    "техникалық тапсырма", "тапсырма",
]
EXCLUDE_NAME_HINTS = [
    "смета", "чертеж", "чертёж", "протокол", "проект договора", "рисунок",
    "форма заявки", "доверенность",
]

REQUEST_HEADERS = {
    "User-Agent": (
        "Mozilla/5.0 (Windows NT 10.0; Win64; x64) AppleWebKit/537.36 "
        "(KHTML, like Gecko) Chrome/126.0.0.0 Safari/537.36 Edg/126.0.0.0"
    ),
    "Accept-Language": "ru-RU,ru;q=0.9",
}

ANALYSIS_PROMPT = """Ты помогаешь специалисту быстро понять техническое задание (ТЗ) тендера \
на госзакупках Казахстана (может быть про IT/разработку ПО или про проектирование/AutoCAD/\
строительство — определи сам по содержанию).

Прочитай текст ТЗ ниже и напиши краткую выжимку на русском языке (не более 8-10 пунктов, списком):
- Что конкретно нужно сделать/поставить (предмет закупки)
- Ключевые технические требования/объёмы/масштаб работ
- Сроки выполнения, если указаны
- Требования к квалификации/опыту/допускам исполнителя, если есть
- Любые особо важные условия (гарантии, этапность, обязательные стандарты/ГОСТы/СНиПы, лицензии)

Если текст обрезан или неполон — напиши это явно отдельной строкой в конце. Не выдумывай то, \
чего нет в тексте.

Текст ТЗ:
---
{text}
---
"""


def log(msg):
    print(msg, flush=True)


def parse_documents_cell(cell_value):
    """'name1 (url1); name2 (url2)' -> [{"name":..., "url":...}, ...]"""
    if not cell_value:
        return []
    docs = []
    for part in str(cell_value).split(";"):
        part = part.strip()
        if not part:
            continue
        m = DOC_LINK_RE.match(part)
        if m:
            docs.append({"name": m.group(1).strip(), "url": m.group(2).strip()})
    return docs


def pick_tz_document(docs):
    """Выбирает наиболее вероятный документ ТЗ из списка вложений."""
    if not docs:
        return None
    for d in docs:
        name_lower = d["name"].lower()
        if any(hint in name_lower for hint in TZ_NAME_HINTS):
            return d
    candidates = [
        d for d in docs
        if d["url"].lower().split("?")[0].endswith((".pdf", ".doc", ".docx"))
        and not any(h in d["name"].lower() for h in EXCLUDE_NAME_HINTS)
    ]
    if candidates:
        return candidates[0]
    return docs[0]


def safe_filename(text, max_len=80):
    text = re.sub(r"[^\w\-. ]+", "_", str(text), flags=re.UNICODE).strip()
    return text[:max_len] or uuid.uuid4().hex


def download_file(url, dest_path, attempts=4):
    for attempt in range(attempts):
        try:
            with requests.get(url, headers=REQUEST_HEADERS, stream=True, timeout=30) as resp:
                resp.raise_for_status()
                with open(dest_path, "wb") as f:
                    for chunk in resp.iter_content(chunk_size=8192):
                        f.write(chunk)
            return True
        except Exception as e:
            wait = 3 * (attempt + 1)
            log(f"    [!] Не удалось скачать (попытка {attempt + 1}/{attempts}): {e}. Жду {wait}с...")
            if attempt < attempts - 1:
                time.sleep(wait)
    return False


def extract_text(path):
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            import pdfplumber
            parts = []
            with pdfplumber.open(path) as pdf:
                for page in pdf.pages:
                    parts.append(page.extract_text() or "")
            return "\n".join(parts)
        elif ext == ".docx":
            import docx
            d = docx.Document(path)
            parts = [p.text for p in d.paragraphs]
            for table in d.tables:
                for row in table.rows:
                    for cell in row.cells:
                        parts.append(cell.text)
            return "\n".join(parts)
        elif ext == ".doc":
            log("    [!] .doc (старый формат Word) — авто-извлечение текста не поддерживается "
                "(нужен модуль вроде textract с системными зависимостями). Файл всё равно "
                "скачан и доступен по ссылке в Excel, просто без анализа ИИ.")
            return ""
    except Exception as e:
        log(f"    [!] Не удалось извлечь текст из {os.path.basename(path)}: {e}")
    return ""


def analyze_with_ai(text, model, api_key):
    if not text.strip():
        return ""
    try:
        import anthropic
    except ImportError:
        log("    [!] Модуль 'anthropic' не установлен (pip install anthropic) — анализ пропущен.")
        return ""

    client = anthropic.Anthropic(api_key=api_key)
    max_chars = 60000  # разумный запас под контекст модели
    truncated = text[:max_chars]
    suffix = "\n\n[текст обрезан — документ длиннее]" if len(text) > max_chars else ""
    try:
        resp = client.messages.create(
            model=model,
            max_tokens=800,
            messages=[{"role": "user", "content": ANALYSIS_PROMPT.format(text=truncated + suffix)}],
        )
        return "".join(getattr(block, "text", "") for block in resp.content).strip()
    except Exception as e:
        log(f"    [!] Ошибка при обращении к ИИ: {e}")
        return ""


def main():
    parser = argparse.ArgumentParser(
        description="Скачивает ТЗ тендеров из Excel-файла и (опционально) анализирует их через ИИ."
    )
    parser.add_argument("xlsx_path", help="Путь к Excel-файлу (goszakup_it_tenders.xlsx / "
                                           "_astana.xlsx / izbrannye_tendery.xlsx)")
    parser.add_argument("--output-dir", default="ТЗ",
                         help="Папка для скачанных файлов ТЗ (по умолчанию: ./ТЗ)")
    parser.add_argument("--analyze", action="store_true",
                         help="Прогонять текст ТЗ через ИИ (Claude) и писать выжимку в отдельный столбец")
    parser.add_argument("--model", default="claude-sonnet-5",
                         help="Модель Claude для анализа (по умолчанию claude-sonnet-5; "
                              "для экономии на больших объёмах — claude-haiku-4-5-20251001)")
    parser.add_argument("--api-key", default=None,
                         help="Anthropic API-ключ (по умолчанию берётся из ANTHROPIC_API_KEY)")
    parser.add_argument("--in-place", action="store_true",
                         help="Изменить исходный файл напрямую вместо создания копии "
                              "<имя>_с_ТЗ.xlsx. ОСТОРОЖНО: не используйте, если веб-интерфейс "
                              "может сейчас параллельно писать в этот же файл (идёт сбор) — "
                              "он периодически пересохраняет книгу и сотрёт эти изменения.")
    parser.add_argument("--limit", type=int, default=None,
                         help="Обработать только первые N строк (для тестового прогона)")
    args = parser.parse_args()

    api_key = args.api_key or os.environ.get("ANTHROPIC_API_KEY")
    if args.analyze and not api_key:
        log("[!] --analyze указан, но не найден ANTHROPIC_API_KEY (и не передан --api-key). Останавливаюсь.")
        sys.exit(1)

    if not os.path.exists(args.xlsx_path):
        log(f"[!] Файл не найден: {args.xlsx_path}")
        sys.exit(1)

    if args.in_place:
        working_path = args.xlsx_path
    else:
        base, ext = os.path.splitext(args.xlsx_path)
        working_path = f"{base}_с_ТЗ{ext}"
        shutil.copy2(args.xlsx_path, working_path)
        log(f"Работаю с копией (исходный файл не трогаю): {working_path}")

    os.makedirs(args.output_dir, exist_ok=True)

    log(f"Открываю {working_path}...")
    wb = load_workbook(working_path)
    ws = wb.active

    header_row = [cell.value for cell in ws[1]]

    def col_index(name, create=False):
        if name in header_row:
            return header_row.index(name) + 1
        if create:
            idx = len(header_row) + 1
            ws.cell(row=1, column=idx, value=name)
            header_row.append(name)
            return idx
        return None

    docs_col = col_index("Документы (названия и ссылки)")
    number_col = col_index("№ объявления")

    if docs_col is None:
        log("[!] В файле нет столбца 'Документы (названия и ссылки)' — нечего скачивать. "
            "Убедитесь, что используете файл, выгруженный веб-интерфейсом (там уже собраны "
            "ссылки на вложения) — goszakup_it_tenders.xlsx, _astana.xlsx или izbrannye_tendery.xlsx.")
        sys.exit(1)

    tz_file_col = col_index("Файл ТЗ", create=True)
    tz_status_col = col_index("Статус скачивания ТЗ", create=True)
    analysis_col = col_index("Анализ ТЗ (ИИ)", create=True) if args.analyze else None

    total_rows = ws.max_row - 1
    processed = 0
    downloaded = 0
    analyzed = 0

    for row_idx in range(2, ws.max_row + 1):
        if args.limit and processed >= args.limit:
            log(f"Достигнут лимит --limit {args.limit}, останавливаюсь.")
            break
        processed += 1
        number_anno = ws.cell(row=row_idx, column=number_col).value if number_col else str(row_idx)
        docs_cell = ws.cell(row=row_idx, column=docs_col).value
        docs = parse_documents_cell(docs_cell)

        log(f"[{processed}/{total_rows}] {number_anno}: вложений {len(docs)}")

        if not docs:
            ws.cell(row=row_idx, column=tz_status_col, value="нет вложений")
            continue

        tz_doc = pick_tz_document(docs)
        if not tz_doc:
            ws.cell(row=row_idx, column=tz_status_col, value="ТЗ не определено")
            continue

        ext = os.path.splitext(tz_doc["url"].split("?")[0])[1].lower()
        if ext not in (".pdf", ".doc", ".docx"):
            ext = ".pdf"
        doc_name_no_ext = os.path.splitext(tz_doc["name"])[0]
        filename = f"{safe_filename(number_anno)}_{safe_filename(doc_name_no_ext)}{ext}"
        dest_path = os.path.join(args.output_dir, filename)

        if os.path.exists(dest_path):
            log(f"    Уже скачано ранее: {filename}")
        else:
            ok = download_file(tz_doc["url"], dest_path)
            if not ok:
                ws.cell(row=row_idx, column=tz_status_col, value="ошибка скачивания")
                continue
            downloaded += 1
            log(f"    Скачано: {filename}")

        cell = ws.cell(row=row_idx, column=tz_file_col, value=filename)
        cell.hyperlink = "file:///" + os.path.abspath(dest_path).replace("\\", "/")
        cell.font = Font(color="0563C1", underline="single")
        ws.cell(row=row_idx, column=tz_status_col, value="ок")

        if args.analyze:
            text = extract_text(dest_path)
            summary = analyze_with_ai(text, args.model, api_key)
            if summary:
                ws.cell(row=row_idx, column=analysis_col, value=summary)
                analyzed += 1
                log(f"    Анализ ИИ готов ({len(summary)} симв.)")

        # Сохраняем прогресс по ходу дела - на случай обрыва посреди большого файла.
        if processed % 5 == 0:
            wb.save(working_path)

    wb.save(working_path)
    log("=" * 60)
    summary_line = f"Обработано строк: {processed}, скачано ТЗ: {downloaded}"
    if args.analyze:
        summary_line += f", проанализировано ИИ: {analyzed}"
    log(summary_line)
    log(f"Файлы ТЗ сохранены в: {os.path.abspath(args.output_dir)}")
    log(f"Excel обновлён: {os.path.abspath(working_path)}")


if __name__ == "__main__":
    main()
