# -*- coding: utf-8 -*-
"""
finder.py — автоматизированный парсер IT-тендеров goszakup.gov.kz
(однофайловая версия).

ПЕРЕД ЗАПУСКОМ:
    pip install selenium beautifulsoup4 requests openpyxl pdfplumber python-docx xlrd

    Скачайте msedgedriver.exe под версию вашего Edge:
    https://developer.microsoft.com/microsoft-edge/tools/webdriver/
    Положите его рядом с finder.py (или укажите путь в EDGE_DRIVER_PATH ниже).

КАК ЗАПУСКАТЬ:
    1) Просто запустите:  python finder.py

       Скрипт сам откроет ОТДЕЛЬНОЕ окно Edge со своим собственным профилем
       (папка AUTOMATION_PROFILE_DIR рядом со скриптом, по умолчанию
       ".edge_automation_profile"). Это НЕ ваш обычный профиль Edge — он не
       использует ваши обычные вкладки/куки/пароли и не конфликтует с уже
       открытым обычным Edge. Можно держать оба браузера открытыми
       одновременно.

    2) В открывшемся отдельном окне Edge: авторизуйтесь на goszakup.gov.kz
       (при первом запуске нужно будет залогиниться заново — это отдельный
       "чистый" профиль), настройте фильтры на странице "Поиск объявлений",
       нажмите "Найти".

       При следующих запусках логин обычно сохраняется в этом профиле,
       заново вводить пароль не потребуется.

    3) Вернитесь в консоль и нажмите Enter, когда список результатов открыт
       в этом отдельном окне Edge - дальше парсер работает автоматически.

Результат: goszakup_it_tenders.xlsx (дозаписывается построчно, без потери
данных при сбое) + goszakup_it_tenders_astana.xlsx (отдельно, тендеры по
Астане — см. ниже) + checkpoint_seen_ids.json (для дедупликации между
запусками) + errors.log (полные traceback'и всех ошибок обработки
объявлений).
"""

import json
import os
import random
import re
import subprocess
import sys
import threading
import time
import traceback
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode, urljoin

import requests
from bs4 import BeautifulSoup
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService

KEYWORDS_FILE = "it_keywords.json"


# ============================================================================
# ================================ НАСТРОЙКИ ================================
# ============================================================================

# --- Подключение к Edge ---
EDGE_DEBUG_ADDRESS = "127.0.0.1:9222"
EDGE_DRIVER_PATH = "msedgedriver.exe"          # путь к msedgedriver, если не в PATH
EDGE_BROWSER_PATH = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
AUTO_LAUNCH_EDGE = True                         # True -> finder.py сам попробует открыть Edge с портом отладки

# Отдельный профиль Edge только для этого скрипта - НЕ ваш обычный профиль.
# Благодаря этому окно автоматизации не конфликтует с уже открытым обычным
# Edge и не имеет доступа к вашим обычным вкладкам/паролям/истории.
# Можно указать абсолютный путь, если хотите хранить его не рядом со скриптом.
AUTOMATION_PROFILE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), ".edge_automation_profile"
)

# --- Поведение парсера ---
COUNT_RECORD_PER_PAGE = 2000
MAX_PAGES_SAFETY_LIMIT = 2000
REQUEST_DELAY_SEC = (0.3, 0.7)          # пауза между страницами списка (requests, не браузер - можно быстрее)
DETAIL_PAGE_DELAY_SEC = (0.15, 0.4)     # пауза перед запросом карточки тендера (на поток)
REQUEST_TIMEOUT_SEC = 20                # таймаут одного HTTP-запроса (requests)
DETAIL_WORKERS = 4                     # сколько карточек тендеров обрабатывать параллельно
# Раньше стояло 1 (по сути однопоточная обработка), хотя вся инфраструктура
# (своя requests.Session на поток, блокировки на запись в Excel/чекпоинт,
# уникальные имена временных файлов) уже была готова к параллельной работе.
# 4 - разумный старт; если сайт начнёт отвечать ошибками/капчей, уменьшите
# это число обратно.

SCAN_ATTACHED_DOCUMENTS = True          # если по названию+описанию IT не найден - проверять вложения (резервная проверка)
MAX_DOCS_TO_SCAN_IF_NO_MATCH = 3        # сколько вложений максимум скачивать на такую резервную проверку
MAX_DOCUMENT_SIZE_MB = 25
DOWNLOAD_TMP_DIR = "tmp_downloads"
DELETE_DOWNLOADED_DOCS_AFTER_SCAN = True

# Раз в сколько НОВЫХ найденных IT-тендеров сохранять Excel на диск
# (файл держится открытым в памяти между сохранениями - это сильно быстрее,
# чем пересохранять весь .xlsx на каждой строке).
XLSX_SAVE_EVERY_N_MATCHES = 5

# --- Вывод / отладка ---
OUTPUT_XLSX = "goszakup_it_tenders.xlsx"
CHECKPOINT_FILE = "checkpoint_seen_ids.json"
ERROR_LOG_FILE = "errors.log"           # сюда пишутся ПОЛНЫЕ traceback'и всех ошибок обработки
DEBUG_MODE = False                      # True -> сохранять HTML каждой страницы в debug_html/
DEBUG_DUMP_DIR = "debug_html"

# --- IT-ключевые слова для классификации (регистр не важен) ---
# ВАЖНО: объявления на goszakup.gov.kz часто публикуются на казахском языке
# (например, "Біріңғай цифрлық білім беру экожүйесін дамыту бойынша
# жұмыстар" - "Работы по развитию единой цифровой образовательной
# экосистемы"). Раньше все списки ниже содержали только русские слова, из-за
# чего казахскоязычные IT-тендеры почти всегда пропускались - не потому что
# они не про IT, а просто потому что скрипт не узнавал казахские термины.
# Поэтому к каждому русскому списку ниже добавлены казахские соответствия.
IT_STRONG_KEYWORDS = [
    # --- русский ---
    "разработка программного обеспечения",
    "разработка информационной системы",
    "информационная система",
    "автоматиз",  # основа слова - покрывает "автоматизация"/"автоматизации"/"автоматизировать"
    "внедрение информационной системы",
    "доработка по",
    "модернизация информационной системы",
    "crm", "erp", "bpm",
    "api",
    "веб-портал",
    "веб-приложени",
    "мобильное приложени",
    "sql", "postgresql", "mysql",
    "oracle",
    "база данных",
    "1с", "1c",
    "кибербезопасност",
    "информационная безопасност",
    "эцп",
    "искусственный интеллект",
    "машинное обучение",
    "нейросет",
    # --- казахский ---
    "бағдарламалық қамтамасыз етуді әзірлеу",   # разработка программного обеспечения
    "ақпараттық жүйені әзірлеу",                 # разработка информационной системы
    "ақпараттық жүйе",                           # информационная система
    "автоматтандыру",                            # автоматизация
    "ақпараттық жүйені енгізу",                  # внедрение информационной системы
    "ақпараттық жүйені жаңғырту",                # модернизация информационной системы
    "веб-портал",
    "веб-қосымша",                               # веб-приложение
    "мобильді қосымша",                          # мобильное приложение
    "деректер қоры",                             # база данных
    "киберқауіпсіздік",                          # кибербезопасность
    "ақпараттық қауіпсіздік",                    # информационная безопасность
    "жасанды интеллект",                         # искусственный интеллект
    "машиналық оқыту",                           # машинное обучение
    "нейрондық желі",                            # нейросеть
    "цифрлық экожүйе",                           # цифровая экосистема
    "цифрлық білім беру",                        # цифровое образование
    "бірыңғай ақпараттық жүйе",                  # единая информационная система
]

IT_MEDIUM_KEYWORDS = [
    # --- русский ---
    "программн",
    "программное обеспечение",
    "сопровождение по",
    "техническая поддержка по",
    "сервер",
    "серверное оборудование",
    "сетевое оборудование",
    "виртуализац",
    "облачн",
    "cloud",
    "хостинг",
    "домен",
    "firewall",
    "helpdesk",
    "программист",
    "системный администратор",
    # --- казахский ---
    "бағдарламалық қамтамасыз ету",              # программное обеспечение
    "бағдарламалық жасақтама",                   # программное обеспечение (второй вариант)
    "сервер",
    "желілік жабдық",                            # сетевое оборудование
    "виртуализация",
    "бұлтты",                                    # облачный
    "хостинг",
    "домен",
    "бағдарламашы",                              # программист
    "жүйелік әкімші",                            # системный администратор
    "платформа",
]

IT_WEAK_KEYWORDS = [
    # --- русский ---
    "лицензи",
    "информационн",
    "цифров",
    "электронн",
    # --- казахский ---
    "лицензия",
    "ақпараттық",                                # информационный
    "цифрлық",                                   # цифровой
    "электрондық",                               # электронный
]

# Слова-действия, которые описывают именно РАЗРАБОТКУ/СОЗДАНИЕ/ДОРАБОТКУ ПО
# или построение системы автоматизации "с нуля" или на основе существующей.
# Это ядро новой логики классификации (см. classify_is_it): тендер
# засчитывается как релевантный, только если такое действие явно
# упоминается - иначе, например, обычная закупка серверов или лицензий
# (без единого слова про разработку/создание/доработку/автоматизацию)
# больше не будет считаться "тем самым" IT-тендером, даже набрав немного
# баллов по другим ключевым словам.
# ПРИМЕЧАНИЕ ПРО ОКОНЧАНИЯ: сравнение везде идёт как простая проверка
# подстроки (kw in text), без морфологического разбора. Полное слово вида
# "автоматизация" НЕ совпадёт как подстрока с "автоматизации"/"автоматизирова
# ть" - это разные окончания. Поэтому русские слова-действия ниже заданы в
# виде ОСНОВЫ слова (без окончания) - так одно правило сразу покрывает все
# падежи/формы: "разработ" находит и "разработка", и "разработку", и
# "разработать", и "разработчик"; "автоматиз" находит и "автоматизация", и
# "автоматизации", и "автоматизировать".
DEV_ACTION_KEYWORDS = [
    # --- русский (основы слов, без окончаний) ---
    "разработ",        # разработка/разработку/разработать/разработчик...
    "создан", "создат",  # создание/создать/создана/создать...
    "доработ",         # доработка/доработку/доработать...
    "модерниз",        # модернизация/модернизации/модернизировать...
    "внедр",           # внедрение/внедрить/внедрена...
    "автоматиз",       # автоматизация/автоматизации/автоматизировать...
    "проектирование информационной системы",
    # --- казахский ---
    "әзірлеу",       # разработка / разработать (в т.ч. "әзірлеу жөніндегі жұмыстар")
    "дамыту",        # развитие/разработка ("... дамыту бойынша жұмыстар")
    "жасау",         # создание/создать
    "құру",          # создание/построение (системы, платформы и т.п.)
    "жетілдіру",     # доработка/усовершенствование
    "жаңғырту",      # модернизация
    "енгізу",        # внедрение
    "автоматтандыру",  # автоматизация
]

# Слова, характерные именно для СТРОИТЕЛЬНО-МОНТАЖНЫХ/пусконаладочных
# работ и физической установки оборудования - а не для разработки софта.
# Пользователь ищет тендеры именно на разработку/доработку ПО или создание
# систем автоматизации, а НЕ на монтаж оборудования, поэтому такие
# тендеры (если в них при этом нет ни одного DEV_ACTION_KEYWORDS) должны
# отсекаться, даже если формально упоминают "сервер"/"сетевое оборудование"
# и получают немного баллов по IT_MEDIUM_KEYWORDS.
MOUNTING_EXCLUDE_KEYWORDS = [
    # --- русский ---
    "монтаж",
    "монтажные работы",
    "монтажно-наладочные работы",
    "пусконаладочные работы",
    "демонтаж",
    "поставка и установка",
    "установка оборудования",
    "прокладка кабеля",
    "строительно-монтажные работы",
    "структурированная кабельная система",
    "видеонаблюдени",
    "охранная сигнализация",
    "слаботочные системы",
    "ремонт оргтехники",
    "техническое обслуживание оборудования",
    # --- казахский ---
    "монтаждау",                       # монтаж
    "монтаждау жұмыстары",             # монтажные работы
    "жабдықты орнату",                 # установка оборудования
    "жабдықты жеткізу және орнату",    # поставка и установка оборудования
    "күрделі жөндеу",                  # капитальный ремонт
    "ағымдағы жөндеу",                 # текущий ремонт
    "бейнебақылау жүйесін орнату",     # установка системы видеонаблюдения
]

IT_NEGATIVE_KEYWORDS = [
    # --- русский ---
    "мебель",
    "стол",
    "стул",
    "шкаф",
    "дверь",
    "окно",
    "жалюзи",
    "строительство",
    "ремонт здания",
    "капитальный ремонт",
    "асфальт",
    "бетон",
    "кирпич",
    "труба",
    "водоснабжение",
    "канализация",
    "отопление",
    "спецодежда",
    "бумага",
    "канцелярские товары",
    "хозтовары",
    # --- казахский ---
    "жиһаз",           # мебель
    "құрылыс",         # строительство
    "асфальт",
    "бетон",
    "кірпіш",          # кирпич
    "құбыр",           # труба
    "сумен жабдықтау", # водоснабжение
    "кәріз",           # канализация
    "жылыту",          # отопление
    "арнайы киім",     # спецодежда
    "қағаз",           # бумага
    "кеңсе тауарлары", # канцелярские товары
]

IT_CONTEXT_PHRASES = [
    # --- русский ---
    "техническое задание",
    "разработка информационной системы",
    "создание информационной системы",
    "оказание ит услуг",
    "оказание услуг по сопровождению",
    "оказание услуг по разработке",
    "внедрение программного обеспечения",
    "модернизация информационной системы",
    "автоматизация процессов",
    # --- казахский ---
    "техникалық тапсырма",                   # техническое задание
    "ақпараттық жүйені әзірлеу",             # разработка информационной системы
    "ақпараттық жүйені құру",                # создание информационной системы
    "ит қызметтерін көрсету",                # оказание ИТ услуг
    "бағдарламалық қамтамасыз етуді енгізу", # внедрение программного обеспечения
    "процестерді автоматтандыру",            # автоматизация процессов
    "цифрлық экожүйені дамыту",              # развитие цифровой экосистемы
]

# Фразы, которые часто вызывают ложное определение IT-тендера.
# Они будут вырезаться из текста перед поиском ключевых слов.

EXCLUDE_FALSE_POSITIVE_PHRASES = [
    # --- русский ---
    "информационные услуги",
    "информационное сопровождение",
    "информационно-разъяснительная работа",
    "информационно-консультационные услуги",
    "информационное обеспечение мероприятий",
    "размещение информации",
    "изготовление информационных стендов",
    "информационные материалы",
    "полиграфическая продукция",
    "освещение деятельности",
    "публикация материалов",
    "рекламные услуги",
    "услуги по информированию населения",
    # --- казахский ---
    "ақпараттық қызметтер",              # информационные услуги
    "ақпараттық сүйемелдеу",             # информационное сопровождение
    "ақпараттық материалдар",            # информационные материалы
    "баспа өнімі",                       # полиграфическая продукция
    "жарнамалық қызметтер",              # рекламные услуги
    "қоғамды ақпараттандыру жөніндегі қызметтер",  # услуги по информированию населения
]

KNOWN_STATUSES = [
    "Опубликовано (прием ценовых предложений)",
    "Опубликовано (ожидание проведения аукциона)",
    "Опубликовано (дополнение заявок)",
    "Опубликовано (прием заявок)",
    "Опубликовано (проверка заявок)",
    "Ожидание камерального контроля",
    "Ожидание проведения контроля качества",
    "Отправлено на контроль качества",
    "Пройдено контроль качества",
    "Рассмотрение дополнений заявок",
    "Формирование протокола промежуточных итогов",
    "Формирование протокола преддопуска",
    "Формирование протокола допуска",
    "Формирование протокола итогов",
    "Принятие решение о пересмотре итогов",
    "Принятие решение об исполнении уведомления",
    "Изменена документация",
    "Пересмотр итогов",
    "На обжаловании",
    "Приостановлено",
    "Отказ от закупки",
    "Рассмотрение заявок",
    "Опубликовано",
    "Завершено",
    "Отменено",
]

KNOWN_METHODS = [
    "Второй этап конкурса с использованием рамочного соглашения",
    "Первый этап конкурса с использованием рамочного соглашения",
    "Из одного источника по несостоявшимся закупкам не ГЗ",
    "Из одного источника по несостоявшимся закупкам",
    "Конкурс с предварительным квалификационным отбором",
    "Конкурс с применением двухэтапных процедур",
    "Конкурс с использованием рейтингово-балльной системы",
    "Конкурс с применением специального порядка",
    "Тендер с использованием рейтингово-балльной системы",
    "Тендер с применением особого порядка",
    "Государственные закупки с применением особого порядка",
    "Запрос ценовых предложений (не ГЗ new)",
    "Запрос ценовых предложений",
    "Открытый конкурс",
    "Аукцион (до 2022)",
    "Аукцион (с 2022)",
    "Аукцион (не ГЗ) с 2022 г",
    "Аукцион",
    "Закупка жилища",
    "Закупка по государственному социальному заказу",
    "Конкурс по строительству «под ключ»",
    "Тендер по строительству «под ключ»",
    "Тендер",
]

if os.path.exists(KEYWORDS_FILE):
    with open(KEYWORDS_FILE, "r", encoding="utf-8") as f:
        data = json.load(f)

    IT_STRONG_KEYWORDS = data.get("strong", IT_STRONG_KEYWORDS)
    IT_MEDIUM_KEYWORDS = data.get("medium", IT_MEDIUM_KEYWORDS)
    IT_WEAK_KEYWORDS = data.get("weak", IT_WEAK_KEYWORDS)
    IT_NEGATIVE_KEYWORDS = data.get("negative", IT_NEGATIVE_KEYWORDS)
    IT_CONTEXT_PHRASES = data.get("context", IT_CONTEXT_PHRASES)
    # Новые ключи - необязательные, только если хотите переопределить
    # список слов-действий ("разработка"/"әзірлеу" и т.п.) или список
    # монтажно-строительных слов-исключений своим набором.
    DEV_ACTION_KEYWORDS = data.get("dev_action", DEV_ACTION_KEYWORDS)
    MOUNTING_EXCLUDE_KEYWORDS = data.get("mounting_exclude", MOUNTING_EXCLUDE_KEYWORDS)

ANNOUNCE_LINK_RE = re.compile(r"/ru/announce/index/(\d+)")
NUMBER_ANNO_RE = re.compile(r"\b(\d{5,}-\d+)\b")
_STOP_LOOKAHEAD = (
    r"(?=\s*\d{4}-\d{2}-\d{2}|\s*[\d\s\u00a0]{1,20}[.,]\d{2}\s*(?:$|[А-ЯЁ])"
    # "Способ:" тут раньше не совпадал с реальной меткой на сайте
    # ("Способ закупки:"), из-за чего регулярка не останавливалась и
    # захватывала способ закупки внутрь поля "Заказчик". Теперь стоп-слово
    # ловит "Способ" с любым текстом до двоеточия (например "Способ закупки:").
    r"|Заказчик:|Организатор:|Способ[^:\n]{0,30}:|Статус:|$)"
)
ORGANIZER_RE = re.compile(r"Организатор:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
CUSTOMER_RE = re.compile(r"Заказчик:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
DOC_EXTENSIONS = (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".rtf", ".zip", ".rar", ".txt")

HEADERS = [
    "№ объявления",
    "Название лота/объявления",
    "Заказчик",
    "Организатор",
    "Способ закупки",
    "Сумма, тг.",
    "Статус",
    "Начало приёма заявок",
    "Окончание приёма заявок",
    "Ссылка на объявление",
    "Совпавшие IT-ключевые слова",
    "Документы (названия и ссылки)",
    "Тип IT",
    "Приоритет",
    "Краткое описание/ТЗ (выдержка)",
]


# ============================================================================
# ============================ ВСПОМОГАТЕЛЬНОЕ ==============================
# ============================================================================

def polite_sleep(delay_range):
    time.sleep(random.uniform(*delay_range))


def dump_debug_html(name: str, html: str):
    if not DEBUG_MODE:
        return
    os.makedirs(DEBUG_DUMP_DIR, exist_ok=True)
    with open(os.path.join(DEBUG_DUMP_DIR, name), "w", encoding="utf-8") as f:
        f.write(html)


def log_error_traceback(header: str):
    """Пишет полный traceback последнего исключения в ERROR_LOG_FILE.
    Раньше в консоль выводился только str(e), из-за чего было невозможно
    понять, в какой именно строке кода реально падает ошибка - все ошибки
    выглядели одинаково ('NoneType' object has no attribute 'get'), хотя
    происходили в разных местах."""
    try:
        with open(ERROR_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(f"\n--- {time.strftime('%Y-%m-%d %H:%M:%S')} | {header} ---\n")
            f.write(traceback.format_exc())
            f.write("\n")
    except Exception:
        pass


# ============================================================================
# ============================== БРАУЗЕР =====================================
# ============================================================================

def _edge_debug_port_alive() -> bool:
    try:
        r = requests.get(f"http://{EDGE_DEBUG_ADDRESS}/json/version", timeout=1.5)
        return r.ok
    except Exception:
        return False


def try_auto_launch_edge():
    """Открывает ОТДЕЛЬНОЕ окно Edge со своим профилем (AUTOMATION_PROFILE_DIR).

    Это не обычный профиль пользователя, поэтому новый процесс msedge.exe
    не блокируется уже открытым "обычным" Edge и не мешает ему - оба могут
    работать одновременно. Логин на goszakup.gov.kz нужно будет выполнить
    в этом отдельном окне (один раз - профиль на диске сохранит сессию
    между запусками).
    """
    if _edge_debug_port_alive():
        print("Порт отладки уже отвечает - использую уже открытое окно автоматизации.")
        return

    if not AUTO_LAUNCH_EDGE:
        print("[!] AUTO_LAUNCH_EDGE выключен. Откройте Edge вручную с флагом")
        print(f"    --remote-debugging-port={EDGE_DEBUG_ADDRESS.split(':')[-1]}")
        print(f"    --user-data-dir=\"{AUTOMATION_PROFILE_DIR}\"")
        return

    if not os.path.exists(EDGE_BROWSER_PATH):
        print(f"[!] Не нашёл Edge по пути {EDGE_BROWSER_PATH}. "
              f"Откройте браузер вручную с флагом --remote-debugging-port=9222 "
              f"и --user-data-dir=\"{AUTOMATION_PROFILE_DIR}\".")
        return

    os.makedirs(AUTOMATION_PROFILE_DIR, exist_ok=True)
    port = EDGE_DEBUG_ADDRESS.split(":")[-1]
    print("Открываю отдельное окно Edge для автоматизации (свой профиль,")
    print(f"не связан с вашим обычным Edge): {AUTOMATION_PROFILE_DIR}")
    subprocess.Popen([
        EDGE_BROWSER_PATH,
        f"--remote-debugging-port={port}",
        f"--user-data-dir={AUTOMATION_PROFILE_DIR}",
        "--profile-directory=Default",
        "--no-first-run",
        "--no-default-browser-check",
        "https://goszakup.gov.kz/ru/announce/index",
    ])

    # Ждём, пока порт отладки реально откроется (до ~60 секунд).
    for i in range(120):
        time.sleep(0.5)
        if _edge_debug_port_alive():
            print("Отдельное окно Edge запущено, порт отладки открылся успешно.")
            return
        if i > 0 and i % 10 == 0:
            print(f"    ...жду запуск Edge ({i // 2} сек)")

    # Так как профиль отдельный и свой, единственная реальная причина сбоя
    # здесь - сам файл профиля уже заблокирован ДРУГИМ процессом finder.py /
    # тем же автоматизационным окном, оставшимся в фоне.
    print("\n" + "!" * 70)
    print("Порт отладки не открылся. Возможно, окно автоматизации уже было")
    print("открыто ранее и осталось висеть в фоне (проверьте диспетчер задач")
    print("на процесс msedge.exe, использующий папку профиля):")
    print(f"    {AUTOMATION_PROFILE_DIR}")
    print("\nЧТО СДЕЛАТЬ:")
    print("  1) Закройте все окна Edge, у которых в заголовке/адресной строке")
    print("     видно, что это отдельный профиль автоматизации, либо выполните")
    print("       taskkill /IM msedge.exe /F")
    print("     (это не тронет ваш обычный Edge, если он уже был закрыт при")
    print("      выполнении этой команды - она закрывает ВСЕ окна msedge.exe,")
    print("      так что обычный Edge тоже закроется и его нужно будет")
    print("      открыть заново отдельно).")
    print("  2) Запустите finder.py снова.")
    print("!" * 70 + "\n")


def attach_to_running_edge() -> webdriver.Edge:
    options = EdgeOptions()
    options.add_experimental_option("debuggerAddress", EDGE_DEBUG_ADDRESS)
    service = EdgeService(executable_path=EDGE_DRIVER_PATH)
    return webdriver.Edge(service=service, options=options)


def get_cookies_dict(driver: webdriver.Edge) -> dict:
    return {c["name"]: c["value"] for c in driver.get_cookies()}


# ============================================================================
# ======================== ПАРСИНГ СТРАНИЦЫ СПИСКА ==========================
# ============================================================================

def build_page_url(base_url: str, page: int) -> str:
    parts = urlsplit(base_url)

    # сохраняем ВСЕ параметры, включая повторяющиеся filter[status][]
    params = parse_qsl(parts.query, keep_blank_values=True)

    new_params = []

    for k, v in params:
        if k not in ("page", "count_record"):
            new_params.append((k, v))

    new_params.append(("page", str(page)))
    new_params.append(("count_record", str(COUNT_RECORD_PER_PAGE)))

    return urlunsplit((
        parts.scheme,
        parts.netloc,
        parts.path,
        urlencode(new_params),
        ""
    ))


def _find_status(row_text: str) -> str:
    for status in KNOWN_STATUSES:
        if status in row_text:
            return status
    return ""


def _find_method(row_text: str) -> str:
    for method in KNOWN_METHODS:
        if method in row_text:
            return method
    return ""


def _clean_amount(raw: str) -> str:
    """Иногда при схлопывании строки таблицы в одну строку текста соседняя
    ячейка (например, короткий числовой код) слипается с суммой без
    настоящего разделителя, и регулярка захватывает всё как одно число
    (например, "00 150 000.00" вместо "150 000.00"). Отрезаем такой
    мусорный ведущий фрагмент из одних нулей."""
    if not raw:
        return raw
    s = raw.strip()
    m = re.match(r"^0{1,3}[\s\u00a0]+(\d{1,3}(?:[\s\u00a0]\d{3})*[.,]\d{2})$", s)
    if m:
        return m.group(1)
    return s


def _extract_amount(row, row_text: str) -> str:
    """Извлекает сумму из строки таблицы.

    Раньше сумма всегда бралась регуляркой по ВСЕЙ схлопнутой в одну строку
    текстовой строке (row_text), выбирая последнее совпадение вида
    "123 456.78". Это хрупко: если в строке есть другое число такого же
    вида (например, часть кода лота или БИН с разделителем), сумма могла
    быть выбрана неверно.

    Теперь, если у нас есть доступ к исходным ячейкам таблицы (row - это
    <tr>), сначала пробуем найти ячейку, весь текст которой ЦЕЛИКОМ похож
    на денежную сумму - это гораздо надёжнее, чем искать подстроку в общем
    тексте строки. Если это не сработало (не <tr>, либо ни одна ячейка не
    подошла) - используем старый regex-по-всей-строке способ как fallback.
    """
    money_cell_re = re.compile(r"^[\d\s\u00a0]{1,20}[.,]\d{2}$")
    if row is not None and getattr(row, "name", None) == "tr":
        cells = row.find_all(["td", "th"])
        # Идём с конца - сумма обычно ближе к концу строки таблицы
        # (после названия/заказчика/организатора).
        for cell in reversed(cells):
            cell_text = cell.get_text(strip=True)
            if money_cell_re.match(cell_text):
                return _clean_amount(cell_text)

    amount_matches = re.findall(r"[\d\s\u00a0]{1,20}[.,]\d{2}", row_text)
    if amount_matches:
        return _clean_amount(amount_matches[-1].strip())
    return ""


def _truncate_leaked_text(text: str, *markers: str) -> str:
    """ORGANIZER_RE/CUSTOMER_RE останавливаются по текстовому лейблу
    "Способ:"/"Статус:" - но, судя по всему, сайт способ закупки выводит
    БЕЗ отдельного лейбла перед ним, просто сплошным текстом сразу после
    заказчика/организатора. Из-за этого способ закупки/статус иногда
    целиком попадает внутрь поля "Заказчик" или "Организатор". Способ и
    статус мы и так независимо находим через _find_method/_find_status,
    так что просто обрезаем поле по первому вхождению уже найденного
    текста способа/статуса, если он туда просочился."""
    if not text:
        return text
    cut_at = len(text)
    for marker in markers:
        if marker:
            idx = text.find(marker)
            if idx != -1:
                cut_at = min(cut_at, idx)
    return text[:cut_at].strip(" ,.-\u00a0")


def parse_list_page(html: str) -> list:
    soup = BeautifulSoup(html, "html.parser")
    results = []
    seen_ids_on_page = set()

    for a_tag in soup.find_all("a", href=ANNOUNCE_LINK_RE):
        match = ANNOUNCE_LINK_RE.search(a_tag.get("href", ""))
        if not match:
            continue
        tender_id = match.group(1)
        if tender_id in seen_ids_on_page:
            continue
        seen_ids_on_page.add(tender_id)

        title = a_tag.get_text(strip=True)
        if not title:
            continue

        row = a_tag.find_parent("tr") or a_tag.find_parent(["div", "li"])
        row_text = row.get_text(" ", strip=True) if row else a_tag.get_text(" ", strip=True)

        number_anno_match = NUMBER_ANNO_RE.search(row_text)
        number_anno = number_anno_match.group(1) if number_anno_match else tender_id

        organizer_match = ORGANIZER_RE.search(row_text)
        organizer_name = organizer_match.group(1).strip() if organizer_match else ""

        customer_match = CUSTOMER_RE.search(row_text)
        customer_name = customer_match.group(1).strip() if customer_match else ""

        status = _find_status(row_text)
        method = _find_method(row_text)

        organizer_name = _truncate_leaked_text(organizer_name, method, status)
        customer_name = _truncate_leaked_text(customer_name, method, status)

        amount = _extract_amount(row, row_text)

        dates = re.findall(r"\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(?::\d{2})?", row_text)
        start_date = dates[0] if len(dates) >= 1 else ""
        end_date = dates[1] if len(dates) >= 2 else ""

        results.append({
            "id": tender_id,
            "number_anno": number_anno,
            "title": title,
            "url": f"https://goszakup.gov.kz/ru/announce/index/{tender_id}",
            "organizer_name": organizer_name,
            "customer_name": customer_name,
            "status": status,
            "method": method,
            "amount": amount,
            "start_date": start_date,
            "end_date": end_date,
        })

    return results


# ============================================================================
# ======================= ПАРСИНГ ДЕТАЛЬНОЙ СТРАНИЦЫ =========================
# ============================================================================

# Теги, которые почти никогда не относятся к содержанию самого объявления,
# а являются "обвязкой" сайта (меню, футер, выпадающие списки фильтров и
# категорий и т.п.) - именно из-за них раньше в текст объявления попадали
# названия ВСЕХ категорий закупок (включая IT), и почти всё определялось
# как "IT-тендер".
_CHROME_TAGS = ("script", "style", "nav", "header", "footer", "select", "option", "noscript", "iframe")
_CHROME_HINT_RE = re.compile(
    r"(menu|navbar|nav-|footer|header|sidebar|breadcrumb|filter|search-form|"
    r"cookie|lang-switch|top-panel|topline|langs)",
    re.I,
)


def _strip_page_chrome(soup: BeautifulSoup) -> None:
    for tag_name in _CHROME_TAGS:
        for tag in soup.find_all(tag_name):
            tag.decompose()
    # Второй проход - по классам/id, которые явно указывают на меню/футер/фильтры,
    # даже если это не семантические теги <nav>/<footer>, а div/ul и т.п.
    #
    # ВАЖНО: soup.find_all(True) возвращает список ВСЕХ тегов, вычисленный
    # один раз в начале. Но decompose() внутри этого же цикла удаляет не
    # только сам тег, но и всех его потомков - у них при этом .attrs
    # становится None. Если такой потомок ещё встретится дальше в этом же
    # предвычисленном списке, вызов tag.get(...) упадёт с AttributeError
    # ('NoneType' object has no attribute 'get'), т.к. .get() внутри bs4
    # обращается к self.attrs.get(...). Поэтому пропускаем уже
    # декомпозированные (в т.ч. косвенно, как потомки) теги.
    for tag in soup.find_all(True):
        if getattr(tag, "decomposed", False):
            continue
        cls = " ".join(tag.get("class", []) or [])
        tag_id = tag.get("id", "") or ""
        if (cls and _CHROME_HINT_RE.search(cls)) or (tag_id and _CHROME_HINT_RE.search(tag_id)):
            tag.decompose()


def parse_detail_page(html: str, page_url: str) -> dict:
    soup = BeautifulSoup(html, "html.parser")

    # Ссылки на документы собираем ДО чистки - вложения иногда лежат в блоках,
    # которые могут случайно попасть под чистку по классам.
    documents = []
    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        href_lower = href.lower()
        if href_lower.endswith(DOC_EXTENSIONS) or "download" in href_lower or "/file/" in href_lower:
            full_url = urljoin(page_url, href)
            doc_name = a_tag.get_text(strip=True) or full_url.split("/")[-1]
            documents.append({"name": doc_name, "url": full_url})

    _strip_page_chrome(soup)

    main_container = (
        soup.find("div", attrs={"id": re.compile(r"content", re.I)})
        or soup.find("main")
        or soup.find("table")
        or soup.find("body")
    )
    description_text = main_container.get_text("\n", strip=True) if main_container else ""

    customer_bin = ""
    bin_match = re.search(r"БИН[:\s]*([0-9]{12})", description_text)
    if bin_match:
        customer_bin = bin_match.group(1)

    return {"description": description_text, "documents": documents, "customer_bin": customer_bin}


# ============================================================================
# ==================== СКАЧИВАНИЕ И ЧТЕНИЕ ДОКУМЕНТОВ ========================
# ============================================================================

def _safe_filename(url: str) -> str:
    name = url.split("?")[0].rstrip("/").split("/")[-1]
    name = re.sub(r"[^\w\.\-]+", "_", name)
    if not name:
        name = uuid.uuid4().hex
    return f"{uuid.uuid4().hex[:8]}_{name}"


def download_file(url: str, session: requests.Session, cookies: dict):
    os.makedirs(DOWNLOAD_TMP_DIR, exist_ok=True)
    local_path = os.path.join(DOWNLOAD_TMP_DIR, _safe_filename(url))
    try:
        with session.get(url, cookies=cookies, stream=True, timeout=30) as resp:
            resp.raise_for_status()
            total = 0
            max_bytes = MAX_DOCUMENT_SIZE_MB * 1024 * 1024
            with open(local_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    total += len(chunk)
                    if total > max_bytes:
                        f.close()
                        os.remove(local_path)
                        return None
                    f.write(chunk)
        return local_path
    except Exception:
        if os.path.exists(local_path):
            try:
                os.remove(local_path)
            except OSError:
                pass
        return None


def _extract_pdf(path):
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docx(path):
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_doc_legacy(path):
    try:
        import textract
        return textract.process(path).decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_xlsx(path):
    try:
        import openpyxl as _openpyxl
        wb = _openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        parts.append(str(cell))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_xls_legacy(path):
    try:
        import xlrd
        wb = xlrd.open_workbook(path)
        parts = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for cell in sheet.row(row_idx):
                    parts.append(str(cell.value))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".doc":
            return _extract_doc_legacy(path)
        elif ext in (".xlsx", ".xlsm"):
            return _extract_xlsx(path)
        elif ext == ".xls":
            return _extract_xls_legacy(path)
        elif ext in (".txt", ".rtf"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        return ""
    except Exception:
        return ""


def get_document_text(url: str, session: requests.Session, cookies: dict) -> str:
    path = download_file(url, session, cookies)
    if not path:
        return ""
    try:
        return extract_text_from_file(path)
    finally:
        if DELETE_DOWNLOADED_DOCS_AFTER_SCAN and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass


# ============================================================================
# ================================ КЛАССИФИКАЦИЯ =============================
# ============================================================================

def _normalize(text: str) -> str:
    return (text or "").lower().replace("ё", "е")

# Слова/фразы, которые сами по себе означают именно ПРОДЛЕНИЕ существующей
# лицензии/подписки/поддержки, а не новую IT-закупку.
LICENSE_EXCLUDE_PATTERNS = [
    # --- русский ---
    "продление лиценз",
    "продление права использования",
    "продление права пользования",
    "продление подписки",
    "подписка",
    "renewal",
    "subscription",
    "annual support",
    "техническая поддержка программного обеспечения",
    "продление технической поддержки",
    "продление сопровождения",
    "продление доступа",
    "право использования",
    "неисключительная лицензия",
    "лицензия на использование",
    "продление лицензии",
    "продление лицензий",
    "продление лицензии на",
    "предоставление лицензии",
    "право пользования по",
    "техническая поддержка и сопровождение",
    # --- казахский ---
    "лицензияны ұзарту",              # продление лицензии
    "жазылымды ұзарту",               # продление подписки
    "жазылым",                        # подписка
    "пайдалану құқығы",               # право использования
    "қолдау көрсету қызметтерін ұзарту",  # продление услуг поддержки
]

# ВАЖНО: раньше сюда же были добавлены названия антивирусных вендоров
# (kaspersky, eset, bitdefender, crowdstrike и т.п.), и is_license_renewal()
# считал тендер продлением лицензии, как только встречала ЛЮБОЕ из этих
# слов - даже без единого слова про "продление"/"подписку". Из-за этого
# полностью новая закупка "Услуги по внедрению системы защиты информации
# на базе Kaspersky Endpoint Security" тоже помечалась как продление и
# выбрасывалась, хотя это обычный (и часто приоритетный) IT-тендер.
# Названия вендоров сами по себе НЕ являются признаком продления - признак
# продления это наличие слов выше (подписка/продление/право использования
# и т.д.), которые срабатывают независимо от конкретного бренда. Поэтому
# бренды из списка-триггера убраны; при необходимости их можно
# использовать отдельно только как вспомогательную информацию, но не как
# причину исключения тендера.

def is_license_renewal(*texts):
    text = _normalize(" ".join(filter(None, texts)))

    if any(_normalize(p) in text for p in LICENSE_EXCLUDE_PATTERNS):
        # Если рядом со словом про продление/подписку встречается явный
        # глагол разработки/создания/доработки (в т.ч. на казахском - см.
        # DEV_ACTION_KEYWORDS), это, скорее всего, не просто продление, а
        # реальная разработка/доработка - не исключаем такой тендер.
        if not any(_normalize(x) in text for x in DEV_ACTION_KEYWORDS):
            return True

    return False

def classify_is_it(title="", description="", documents=""):
    """Определяет, релевантен ли тендер.

    Раньше классификатор считал тендер IT-релевантным при наборе баллов по
    ЛЮБЫМ IT-словам (например, "сервер" + "лицензия" уже давали баллы),
    включая закупки оборудования/монтажа, где явно нет ни слова про
    разработку. Из-за этого, с одной стороны, могли пролезать чисто
    монтажные/поставочные тендеры, а с другой - целиком пропускались
    казахскоязычные тендеры на разработку (не было казахских ключевых
    слов вообще).

    Теперь логика переориентирована конкретно под запрос "ищу тендеры на
    разработку с нуля или доработку существующего софта / создание системы
    автоматизации, а не на монтажные работы":

      1) тендер должен явно упоминать действие разработки/создания/
         доработки/модернизации/внедрения/автоматизации - на русском ИЛИ
         казахском (см. DEV_ACTION_KEYWORDS). Без этого сигнала тендер
         НЕ считается релевантным, даже если наберёт много баллов по
         прочим IT-словам (сервер, лицензия и т.п.) - такие тендеры чаще
         всего про закупку/обслуживание оборудования, а не про разработку;
      2) если при этом в тексте преобладает монтажно-строительная лексика
         (MOUNTING_EXCLUDE_KEYWORDS) - это сильный штраф к баллу, а если
         дополнительно нет НИ ОДНОГО DEV_ACTION_KEYWORDS - тендер отсекается
         сразу (см. правило (1), это и так сработает).
    """
    title = _normalize(title)
    description = _normalize(description)
    documents = _normalize(documents)

    score = 0
    matched = []

    for phrase in EXCLUDE_FALSE_POSITIVE_PHRASES:
        title = title.replace(_normalize(phrase), "")
        description = description.replace(_normalize(phrase), "")
        documents = documents.replace(_normalize(phrase), "")

    def scan(text, strong, medium, weak, multiplier):
        nonlocal score, matched

        for kw in strong:
            if _normalize(kw) in text:
                matched.append(kw)
                score += 10 * multiplier

        for kw in medium:
            if _normalize(kw) in text:
                matched.append(kw)
                score += 4 * multiplier

        for kw in weak:
            if _normalize(kw) in text:
                matched.append(kw)
                score += 1 * multiplier

    scan(title, IT_STRONG_KEYWORDS, IT_MEDIUM_KEYWORDS, IT_WEAK_KEYWORDS, 3)
    scan(description, IT_STRONG_KEYWORDS, IT_MEDIUM_KEYWORDS, IT_WEAK_KEYWORDS, 2)
    scan(documents, IT_STRONG_KEYWORDS, IT_MEDIUM_KEYWORDS, IT_WEAK_KEYWORDS, 1)

    for phrase in IT_CONTEXT_PHRASES:
        if _normalize(phrase) in description:
            matched.append(phrase)
            score += 8

    # Раньше негативные ключевые слова (мебель, стройка, спецодежда и т.п.)
    # проверялись ТОЛЬКО в description, хотя точно такие же слова в title
    # или в текстах вложенных документов ничем не отличаются по смыслу -
    # они точно так же должны понижать оценку. Теперь проверяем по всему
    # объединённому тексту (title + description + documents).
    combined_text = " ".join([title, description, documents])
    for kw in IT_NEGATIVE_KEYWORDS:
        if _normalize(kw) in combined_text:
            score -= 5

    # --- Ядро новой логики: обязателен явный сигнал "это про разработку" ---
    dev_hits = [kw for kw in DEV_ACTION_KEYWORDS if _normalize(kw) in combined_text]
    if dev_hits:
        matched.extend(dev_hits)
        # Считаем один раз по всему объединённому тексту (а не по каждой
        # из трёх зон с разным множителем, как для strong/medium/weak) -
        # само наличие глагола разработки важно как факт, а не то, в каком
        # именно поле (title/description/documents) он встретился.
        score += 10

    # --- Штраф/отсечение монтажно-строительной лексики ---
    mounting_hits = [kw for kw in MOUNTING_EXCLUDE_KEYWORDS if _normalize(kw) in combined_text]
    if mounting_hits:
        score -= 8 * len(set(mounting_hits))

    matched = sorted(set(matched))

    is_it = bool(dev_hits) and score >= IT_SCORE_THRESHOLD

    return is_it, score, matched

def detect_it_type(*text_fragments):
    text = _normalize(" ".join(t for t in text_fragments if t))

    if any(x in text for x in [
        "разработка", "доработка", "внедрение", "crm", "erp", "api", "веб", "мобильное приложени",
        # казахский
        "әзірлеу", "жетілдіру", "енгізу", "веб-қосымша", "мобильді қосымша", "жасау", "құру",
    ]):
        return "ПО"

    if any(x in text for x in [
        "сервер", "коммутатор", "маршрутизатор", "сетевое оборудование", "скс", "лвс", "firewall",
        # казахский
        "желілік жабдық",
    ]):
        return "Оборудование"

    if any(x in text for x in [
        "техническая поддержка", "сопровождение", "аутсорсинг", "администрирование", "helpdesk",
        # казахский
        "техникалық қолдау", "сүйемелдеу", "әкімшілендіру",
    ]):
        return "Услуги"

    return "Другое"

def detect_priority(*text_fragments):
    text = _normalize(" ".join(t for t in text_fragments if t))

    score = 0

    high = [
        "разработка", "внедрение", "информационная система", "erp", "crm", "api",
        "искусственный интеллект", "модернизация",
        # казахский
        "әзірлеу", "енгізу", "ақпараттық жүйе", "жасанды интеллект", "жаңғырту", "дамыту",
        "автоматтандыру",
    ]

    medium = [
        "сервер", "сетевое оборудование", "база данных", "виртуализация", "sql", "postgresql", "1с",
        # казахский
        "деректер қоры", "желілік жабдық",
    ]

    low = [
        "техническая поддержка", "сопровождение", "лицензи",
        # казахский
        "техникалық қолдау", "сүйемелдеу", "лицензия",
    ]

    score += sum(5 for w in high if w in text)
    score += sum(3 for w in medium if w in text)
    score += sum(1 for w in low if w in text)

    if score >= 10:
        return "Высокий"

    if score >= 5:
        return "Средний"

    return "Низкий"


# --- Порог классификации IT-тендера и логирование пограничных случаев ---
# Порог IT_SCORE_THRESHOLD подобран вручную и остаётся некоторым
# компромиссом. Чтобы его можно было спокойно донастраивать в будущем, не
# теряя молча потенциально релевантные тендеры, всё что набрало заметный,
# но недостаточный балл (BORDERLINE_SCORE_MIN <= score < IT_SCORE_THRESHOLD),
# логируется отдельно в BORDERLINE_LOG_FILE - это НЕ IT-тендеры в основном
# понимании скрипта и они никуда не попадают в Excel, но их стоит иногда
# просматривать глазами, чтобы видеть, не сироп ли порог отсекает что-то
# нужное.
IT_SCORE_THRESHOLD = 12
BORDERLINE_SCORE_MIN = 6
BORDERLINE_LOG_FILE = "borderline_candidates.log"
_borderline_lock = threading.Lock()


def log_borderline_candidate(row: dict, score: int, matched: list):
    try:
        with _borderline_lock:
            with open(BORDERLINE_LOG_FILE, "a", encoding="utf-8") as f:
                f.write(
                    f"{time.strftime('%Y-%m-%d %H:%M:%S')}\t"
                    f"score={score}\t"
                    f"{row.get('number_anno', '')}\t"
                    f"{row.get('url', '')}\t"
                    f"{row.get('title', '')[:120]}\t"
                    f"keywords={', '.join(matched)}\n"
                )
    except Exception:
        pass


# ============================================================================
# ============================ ГОРОД АСТАНА ==================================
# ============================================================================

# Тендеры, которые относятся к Астане (заказчик/организатор/описание),
# выделяются в ОТДЕЛЬНЫЙ Excel-файл (OUTPUT_XLSX_ASTANA) и НЕ дублируются
# в основном файле (OUTPUT_XLSX) - см. append_it_result() ниже.
ASTANA_CITY_RE = re.compile(
    r"\b(?:г\.?\s*)?астан[аеу]\b|\bnur[\s\-]?sultan\b|\bнур[\s\-]?султан\w*\b",
    re.I,
)


def is_astana_tender(row: dict, detail: dict) -> bool:
    """Определяет, относится ли тендер к городу Астана.

    Смотрим на заказчика, организатора и (на всякий случай, если в
    заказчике/организаторе город не упомянут явно) на начало описания
    объявления, где обычно указывается адрес/город заказчика. "Нур-Султан"
    учитывается как прежнее название Астаны на случай старых объявлений.
    """
    haystack = " ".join([
        row.get("customer_name", "") or "",
        row.get("organizer_name", "") or "",
        (detail.get("description", "") or "")[:2000],
    ])
    return bool(ASTANA_CITY_RE.search(haystack))


# ============================================================================
# ================================ EXCEL =====================================
# ============================================================================

class ExcelSink:
    """Инкапсулирует один выходной Excel-файл: открытую в памяти книгу,
    лист, блокировку и счётчик несохранённых строк.

    Раньше вся эта логика (глобальные _workbook/_worksheet/_xlsx_lock/
    _rows_since_save + функции init_workbook_for_run/save_workbook_now/
    append_row) существовала в единственном экземпляре как набор глобальных
    переменных модуля. Чтобы добавить второй выходной файл (тендеры по
    Астане) без дублирования кода и не путая состояние одного файла с
    другим, эта логика вынесена в класс - теперь просто создаётся два
    независимых экземпляра ExcelSink с разными путями и optional-суффиксом
    в имени листа.
    """

    def __init__(self, path: str, sheet_title: str = "IT-тендеры"):
        self.path = path
        self.sheet_title = sheet_title
        self.lock = threading.Lock()
        self.workbook = None
        self.worksheet = None
        self.rows_since_save = 0

    def _build_new_workbook(self):
        wb = Workbook()
        ws = wb.active
        ws.title = self.sheet_title
        ws.append(HEADERS)
        header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
        for col_idx in range(1, len(HEADERS) + 1):
            cell = ws.cell(row=1, column=col_idx)
            cell.font = Font(bold=True, color="FFFFFF")
            cell.fill = header_fill
            cell.alignment = Alignment(vertical="center", wrap_text=True)
        widths = [16, 45, 35, 30, 22, 16, 24, 18, 18, 30, 30, 18, 18, 45, 50]
        for i, w in enumerate(widths, start=1):
            ws.column_dimensions[get_column_letter(i)].width = w
        ws.freeze_panes = "A2"
        return wb, ws

    def init_for_run(self):
        """Загружает существующий Excel в память (если есть) либо создаёт
        новый. Вызывается один раз в начале main() - дальше книга не
        перечитывается с диска на каждой строке."""
        if os.path.exists(self.path):
            self.workbook = load_workbook(self.path)
            self.worksheet = (
                self.workbook[self.sheet_title]
                if self.sheet_title in self.workbook.sheetnames
                else self.workbook.active
            )
        else:
            self.workbook, self.worksheet = self._build_new_workbook()
            self.workbook.save(self.path)

    def save_now(self):
        with self.lock:
            if self.workbook is not None:
                self.workbook.save(self.path)
                self.rows_since_save = 0

    def append_row(self, row: dict):
        """Потокобезопасно добавляет строку в открытую в памяти книгу.
        Сохранение на диск происходит не на каждой строке, а раз в
        XLSX_SAVE_EVERY_N_MATCHES найденных тендеров (плюс всегда в конце
        и при остановке/ошибке - см. main())."""
        values = [
            row.get("number_anno", ""),
            row.get("title", ""),
            row.get("customer_name", ""),
            row.get("organizer_name", ""),
            row.get("method", ""),
            row.get("amount", ""),
            row.get("status", ""),
            row.get("start_date", ""),
            row.get("end_date", ""),
            row.get("url", ""),
            ", ".join(row.get("matched_keywords", [])),
            "; ".join(row.get("documents", [])),
            row.get("it_type", ""),
            row.get("priority", ""),
            (row.get("description", "") or "")[:1500],
        ]

        with self.lock:
            self.worksheet.append(values)
            last_row = self.worksheet.max_row
            url_cell = self.worksheet.cell(row=last_row, column=HEADERS.index("Ссылка на объявление") + 1)
            if row.get("url"):
                url_cell.hyperlink = row["url"]
                url_cell.font = Font(color="0563C1", underline="single")
            self.rows_since_save += 1
            need_save = self.rows_since_save >= XLSX_SAVE_EVERY_N_MATCHES

        if need_save:
            self.save_now()


# --- Выходные файлы ---
OUTPUT_XLSX_ASTANA = "goszakup_it_tenders_astana.xlsx"

sink_main = ExcelSink(OUTPUT_XLSX, sheet_title="IT-тендеры")
sink_astana = ExcelSink(OUTPUT_XLSX_ASTANA, sheet_title="IT-тендеры (Астана)")


def init_workbooks_for_run():
    sink_main.init_for_run()
    sink_astana.init_for_run()


def save_workbooks_now():
    sink_main.save_now()
    sink_astana.save_now()


def append_it_result(result: dict, astana: bool):
    """Кладёт найденный IT-тендер РОВНО в один файл:
    - если тендер по Астане -> только в sink_astana;
    - иначе -> только в sink_main.
    Дублирования между файлами нет намеренно, как просили."""
    if astana:
        sink_astana.append_row(result)
    else:
        sink_main.append_row(result)


# ============================================================================
# ============================== ЧЕКПОИНТ ====================================
# ============================================================================

_seen_ids_lock = threading.Lock()


def load_seen_ids() -> set:
    if not os.path.exists(CHECKPOINT_FILE):
        return set()
    try:
        with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("seen_ids", []))
    except Exception:
        return set()


def save_seen_ids(seen_ids: set):
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump({"seen_ids": sorted(seen_ids)}, f, ensure_ascii=False, indent=2)


# ============================================================================
# ============================== ОСНОВНОЙ ЦИКЛ ===============================
# ============================================================================

# ---------------------------------------------------------------------------
# ВАЖНО: requests.Session (и его cookie jar) не рассчитан на одновременное
# использование из нескольких потоков (несколько потоков одновременно
# читают/пишут одну и ту же cookie jar на каждый HTTP-запрос). Раньше все
# DETAIL_WORKERS потоков делили ОДИН session, что и приводило к массовым
# случайным ошибкам вида "AttributeError: 'NoneType' object has no
# attribute 'get'" (гонка внутри http.cookiejar/urllib3).
#
# Решение: у каждого потока - своя requests.Session (thread-local),
# куки в неё каждый раз копируются (дешёвая операция) из "главной"
# сессии, авторизованной через cookies браузера. Сама "главная" сессия
# при этом мутируется только из главного потока, между страницами -
# то есть в момент, когда ThreadPoolExecutor ещё/уже не работает.
# ---------------------------------------------------------------------------
_thread_local = threading.local()


def get_thread_session(base_session: requests.Session) -> requests.Session:
    sess = getattr(_thread_local, "session", None)
    if sess is None:
        sess = requests.Session()
        sess.headers.update(dict(base_session.headers))
        _thread_local.session = sess
    # Синхронизируем куки из главной сессии (безопасно: главный поток
    # не мутирует base_session.cookies, пока работает executor).
    sess.cookies.update(base_session.cookies.get_dict())
    return sess


def fetch_html(session: requests.Session, url: str, attempts: int = 6) -> str:
    """GET с повторными попытками. Используется вместо Selenium для карточек
    тендеров и страниц списка - это на порядок быстрее, т.к. не требует
    полноценного открытия страницы в браузере, и позволяет работать в
    несколько потоков одновременно.

    v3bl.goszakup.gov.kz - судя по всему, бета/тестовый поддомен и иногда
    подвисает на несколько секунд-минут (ConnectTimeout), поэтому здесь
    делается несколько попыток с постепенно растущей паузой (5, 10, 15,
    20, 25 сек), прежде чем сдаться - раньше пауза была всего 0.5-1.5 сек,
    чего для реальных сетевых заминок недостаточно."""
    last_exc = None
    for attempt in range(attempts):
        try:
            resp = session.get(url, timeout=REQUEST_TIMEOUT_SEC)
            resp.raise_for_status()
            resp.encoding = resp.encoding or "utf-8"
            return resp.text
        except Exception as e:
            last_exc = e
            if attempt < attempts - 1:
                wait_sec = 5 * (attempt + 1)
                print(f"    [!] Сетевая заминка ({e.__class__.__name__}), "
                      f"жду {wait_sec} сек и пробую снова "
                      f"(попытка {attempt + 2}/{attempts})...")
                time.sleep(wait_sec)
    raise last_exc


def process_tender(session: requests.Session, row: dict):
    if "заверш" in row.get("status", "").lower():
        return None
    # Собственная сессия для этого потока - см. комментарий выше.
    thread_session = get_thread_session(session)

    polite_sleep(DETAIL_PAGE_DELAY_SEC)
    html = fetch_html(thread_session, row["url"])
    dump_debug_html(f"detail_{row['id']}.html", html)

    detail = parse_detail_page(html, row["url"])

    if is_license_renewal(
        row.get("title", ""),
        detail.get("description", ""),
        " ".join(d["name"] for d in detail["documents"]),
    ):
        return None

    doc_display = [f"{d['name']} ({d['url']})" for d in detail["documents"]]

    # Сначала - быстрая проверка по названию и (уже очищенному от меню/футера)
    # описанию объявления. Для подавляющего большинства НЕ-IT тендеров этого
    # достаточно, и мы вообще не тратим время на скачивание вложений.
    doc_names = " ".join(d["name"] for d in detail["documents"])

    is_it, score, matched_keywords = classify_is_it(
        row.get("title", ""),
        detail.get("description", ""),
        doc_names
    )

    # Резервная проверка: если по названию/описанию IT не найден, но есть
    # прикреплённые документы - заглядываем в них (вдруг IT-требование есть
    # только в самом ТЗ). Ограничено MAX_DOCS_TO_SCAN_IF_NO_MATCH файлами,
    # чтобы не тормозить на объявлениях, где документов много, а к IT они
    # заведомо не относятся.
    if not is_it and SCAN_ATTACHED_DOCUMENTS and detail["documents"]:
        cookies = thread_session.cookies.get_dict()
        for doc in detail["documents"][:MAX_DOCS_TO_SCAN_IF_NO_MATCH]:
            text = get_document_text(doc["url"], thread_session, cookies)
            if text:
                is_it, score, matched_keywords = classify_is_it(
                    doc["name"],
                    text,
                    doc["name"]
                )
                if is_it:
                    break

    if not is_it:
        # Пограничные случаи (заметный, но недостаточный балл) логируем
        # отдельно для последующего ручного просмотра - см. комментарий
        # у IT_SCORE_THRESHOLD/BORDERLINE_SCORE_MIN выше.
        if BORDERLINE_SCORE_MIN <= score < IT_SCORE_THRESHOLD:
            log_borderline_candidate(row, score, matched_keywords)
        return None

    it_type = detect_it_type(
        row.get("title", ""),
        detail.get("description", ""),
        doc_names
    )

    priority = detect_priority(
        row.get("title", ""),
        detail.get("description", ""),
        doc_names
    )

    return {
        "number_anno": row.get("number_anno", ""),
        "title": row.get("title", ""),
        "customer_name": row.get("customer_name") or row.get("organizer_name", ""),
        "customer_bin": detail.get("customer_bin", ""),
        "organizer_name": row.get("organizer_name", ""),
        "method": row.get("method", ""),
        "amount": row.get("amount", ""),
        "status": row.get("status", ""),
        "start_date": row.get("start_date", ""),
        "end_date": row.get("end_date", ""),
        "url": row.get("url", ""),
        "matched_keywords": matched_keywords,
        "documents": doc_display,
        "description": detail.get("description", ""),
        "it_type": it_type,
        "priority": priority,
        "is_astana": is_astana_tender(row, detail),
    }


FINDER_VERSION = (
    "2024-07-v10 (добавлены казахские ключевые слова; классификация "
    "переориентирована на разработку/доработку ПО и системы автоматизации "
    "- требуется явный сигнал DEV_ACTION_KEYWORDS; монтажно-строительные "
    "тендеры без такого сигнала отсекаются; реальный keep-alive браузерной "
    "вкладки; убраны бренды антивирусов из триггера продления лицензии; "
    "негативные ключевые слова проверяются по всему тексту; сумма "
    "извлекается по ячейкам таблицы; логирование пограничных кандидатов; "
    "отдельный Excel по Астане; поднято число потоков обработки карточек; "
    "убран дублирующий импорт)"
)


def browser_keepalive_loop(driver, base_url: str, stop_event: threading.Event, interval_sec: int = 180):
    """Периодически "трогает" сервер goszakup.gov.kz прямо из JS-контекста
    открытой вкладки Edge (через fetch() с credentials: 'include'), чтобы
    сайт не считал сессию бездействующей и не разлогинивал пользователя,
    пока requests-сессия работает параллельно в фоне.

    Раньше в комментариях и в FINDER_VERSION заявлялось, что такой поток
    существует ("keep-alive вкладки против разрыва сессии"), но фактически
    ни функции, ни запуска потока в коде не было - обновлялись только куки
    requests.Session перед каждой страницей списка, что НЕ эквивалентно
    поддержанию активности серверной сессии, если сайт разлогинивает по
    таймауту бездействия именно в браузере.

    Это НЕ полная навигация/refresh страницы - вкладка в браузере не
    дёргается, фильтры/прокрутка/результаты поиска не сбрасываются.
    """
    while not stop_event.wait(interval_sec):
        try:
            driver.execute_script(
                "fetch(arguments[0], {credentials: 'include', cache: 'no-store'}).catch(function(){});",
                base_url,
            )
        except Exception:
            # Не критично - на следующей странице списка куки всё равно
            # обновятся заново из driver.get_cookies().
            pass


def find_and_switch_to_goszakup_tab(driver) -> str:
    """Edge может открыть несколько вкладок (например, стартовую страницу
    MSN плюс ту, что мы указали при запуске) - Selenium по умолчанию
    остаётся на первой из них, а не обязательно на той, где пользователь
    искал тендеры. Эта функция ищет среди ВСЕХ открытых вкладок ту, что
    относится к goszakup.gov.kz, и переключает драйвер на неё.
    Возвращает URL найденной вкладки или None, если не нашла."""
    found_url = None
    found_handle = None
    for handle in driver.window_handles:
        try:
            driver.switch_to.window(handle)
            url = driver.current_url
        except Exception:
            continue
        if "goszakup.gov.kz" in url:
            found_url = url
            found_handle = handle
            # Предпочитаем вкладку именно со страницей поиска/результатов,
            # если открыто несколько вкладок goszakup - но в целом любая
            # вкладка этого домена лучше, чем ничего.
            if "search/announce" in url or "announce" in url:
                break
    if found_handle is not None:
        driver.switch_to.window(found_handle)
    return found_url


def main():
    print("=" * 70)
    print("Парсер IT-тендеров goszakup.gov.kz (finder.py)")
    print(f"Версия скрипта: {FINDER_VERSION}")
    print("=" * 70)

    try_auto_launch_edge()

    print("\nПодключаюсь к отдельному окну Edge (remote debugging)...")
    try:
        driver = attach_to_running_edge()
    except Exception as e:
        print(f"\nНе удалось подключиться к Edge: {e}")
        print("\nПроверьте:")
        print("  1) отдельное окно Edge для автоматизации реально открылось")
        print(f"     (профиль: {AUTOMATION_PROFILE_DIR});")
        print("  2) версия msedgedriver.exe совпадает с версией вашего Edge")
        print("     (Edge -> Настройки -> О Microsoft Edge, сверить номер версии;")
        print("      скачать драйвер: https://developer.microsoft.com/microsoft-edge/tools/webdriver/).")
        sys.exit(1)

    print("Подключение успешно.\n")
    print("Убедитесь, что в ОТДЕЛЬНОМ окне Edge (не в вашем обычном браузере!) вы:")
    print("  1) авторизовались на goszakup.gov.kz;")
    print("  2) настроили нужные фильтры на странице 'Поиск объявлений';")
    print("  3) нажали 'Найти' и видите список результатов.\n")
    input("Когда всё готово - нажмите Enter здесь, чтобы начать сбор данных...")

    base_url = find_and_switch_to_goszakup_tab(driver)
    attempts_left = 3
    while base_url is None and attempts_left > 0:
        print("\n[!] Не нашёл среди открытых вкладок ни одной с goszakup.gov.kz.")
        print("    Убедитесь, что вкладка с результатами поиска действительно")
        print("    открыта в ЭТОМ ЖЕ окне Edge (том, что запустил finder.py),")
        print("    и что вы уже нажали 'Найти'.")
        input("    Когда исправите - нажмите Enter, чтобы попробовать снова...")
        base_url = find_and_switch_to_goszakup_tab(driver)
        attempts_left -= 1

    if base_url is None:
        print("\nТак и не нашёл вкладку goszakup.gov.kz. Останавливаюсь.")
        sys.exit(1)

    print(f"\nБазовый URL с вашими фильтрами: {base_url}")
    if "search/announce" not in base_url and "announce" not in base_url:
        print("[!] Похоже, это не страница результатов поиска объявлений -")
        print("    проверьте, что перед запуском вы дошли до списка тендеров.")

    seen_ids = load_seen_ids()
    print(f"Уже обработано в предыдущих запусках: {len(seen_ids)} объявлений.\n")

    init_workbooks_for_run()

    # Дальше список и карточки тендеров забираем через обычный requests,
    # используя куки из уже залогиненного отдельного окна Edge - это
    # значительно быстрее, чем гонять настоящий браузер по каждой странице.
    #
    # Это "главная" сессия: список страниц (single-threaded) читает её
    # напрямую, а каждый воркер-поток в ThreadPoolExecutor берёт себе
    # ОТДЕЛЬНУЮ сессию через get_thread_session(session) (см. process_tender).
    session = requests.Session()
    session.headers.update({
        "User-Agent": driver.execute_script("return navigator.userAgent;"),
        "Referer": base_url,
        "Accept-Language": "ru-RU,ru;q=0.9",
    })
    session.cookies.update(get_cookies_dict(driver))

    # Фоновый поток, который периодически "трогает" настоящую вкладку Edge
    # через JS fetch(), чтобы сервер не отключал сессию по таймауту
    # бездействия, пока requests-сессия работает в фоне (см.
    # browser_keepalive_loop выше).
    keepalive_stop = threading.Event()
    keepalive_thread = threading.Thread(
        target=browser_keepalive_loop,
        args=(driver, base_url, keepalive_stop),
        daemon=True,
    )
    keepalive_thread.start()

    total_found_it = 0
    total_found_astana = 0
    total_checked = 0
    total_errors = 0
    page = 1

    try:
        while page <= MAX_PAGES_SAFETY_LIMIT:
            page_url = build_page_url(base_url, page)
            print(page_url)
            print(f"--- Страница {page}: {page_url}")

            # Куки на всякий случай обновляем каждую страницу списка (дёшево,
            # без навигации - просто читаем текущее состояние окна логина),
            # чтобы устойчиво работать, даже если сайт периодически ротирует
            # токены сессии. Это происходит ДО запуска пула потоков на эту
            # страницу, поэтому конкурентной записи в session.cookies с
            # чтением из воркер-потоков здесь не возникает.
            session.cookies.update(get_cookies_dict(driver))

            # Ждём, пока страница списка реально не загрузится - без ограничения
            # по числу попыток. Пауза между попытками растёт (10с, 20с, 30с...),
            # но не больше 2 минут, чтобы не ждать бесконечно долго между
            # проверками при действительно долгой недоступности сайта.
            html = None
            attempt_num = 0
            while html is None:
                attempt_num += 1
                try:
                    html = fetch_html(session, page_url)
                except Exception as e:
                    log_error_traceback(f"list page fetch, page={page}, url={page_url}, attempt={attempt_num}")
                    wait_sec = min(10 * attempt_num, 120)
                    print(f"    [!] Страница {page} долго не загружается "
                          f"({e.__class__.__name__}), попытка {attempt_num}. "
                          f"Жду {wait_sec} сек и пробую снова...")
                    time.sleep(wait_sec)
                    # На всякий случай обновим куки перед повторной попыткой -
                    # вдруг сессия за это время истекла.
                    session.cookies.update(get_cookies_dict(driver))

            polite_sleep(REQUEST_DELAY_SEC)
            dump_debug_html(f"list_page_{page}.html", html)

            rows = parse_list_page(html)
            if not rows:
                empty_dump_path = os.path.abspath("empty_page_debug.html")
                try:
                    with open(empty_dump_path, "w", encoding="utf-8") as f:
                        f.write(html)
                except Exception:
                    pass
                print("Больше объявлений не найдено (либо сайт отдал страницу без "
                      "объявлений по другой причине - истекшая сессия, сброшенный "
                      "фильтр, капча и т.п.). Останавливаюсь.")
                print(f"HTML этой страницы сохранён в: {empty_dump_path}")
                print("Откройте этот файл в браузере/текстовом редакторе и сверьте:")
                print("  - если там видна нормальная страница поиска, но БЕЗ строк "
                      "с объявлениями - значит объявлений по вашему фильтру "
                      "действительно больше нет;")
                print("  - если там форма логина, капча, сообщение об ошибке или "
                      "просто исходная (нефильтрованная) страница поиска - значит "
                      "сессия/фильтр сбросились, и нужно заново пройти шаги 1-3 в "
                      "отдельном окне Edge и перезапустить скрипт.")
                break

            new_rows = [
                r for r in rows
                if r["id"] not in seen_ids
                and r.get("status", "") != "Завершено"
            ]
            print(f"    Найдено на странице: {len(rows)}, новых: {len(new_rows)}")

            # Карточки тендеров этой страницы обрабатываем параллельно
            # (DETAIL_WORKERS потоков) - самая медленная часть (сеть,
            # разбор HTML, иногда скачивание вложений) больше не блокирует
            # друг друга по очереди.
            found_on_this_page = 0
            with ThreadPoolExecutor(max_workers=DETAIL_WORKERS) as executor:
                future_to_row = {executor.submit(process_tender, session, row): row for row in new_rows}
                for future in as_completed(future_to_row):
                    row = future_to_row[future]
                    total_checked += 1
                    try:
                        result = future.result()
                    except Exception as e:
                        total_errors += 1
                        print(f"    [!] Ошибка при обработке {row['id']}: {e}  (подробности в {ERROR_LOG_FILE})")
                        log_error_traceback(f"tender_id={row['id']} url={row.get('url', '')}")
                        result = None

                    with _seen_ids_lock:
                        seen_ids.add(row["id"])

                    if result:
                        is_astana = bool(result.pop("is_astana", False))
                        append_it_result(result, astana=is_astana)
                        total_found_it += 1
                        found_on_this_page += 1
                        if is_astana:
                            total_found_astana += 1
                            print(f"    [IT/Астана] {row['number_anno']}: {row['title'][:70]}")
                        else:
                            print(f"    [IT] {row['number_anno']}: {row['title'][:70]}")

                    if total_checked % 20 == 0:
                        with _seen_ids_lock:
                            save_seen_ids(seen_ids)

            if new_rows and found_on_this_page == 0:
                print(f"    На странице {page} IT-тендеров не найдено.")

            with _seen_ids_lock:
                save_seen_ids(seen_ids)
            page += 1

    except KeyboardInterrupt:
        print("\n\nОстановлено пользователем (Ctrl+C). Сохраняю прогресс...")
    except Exception:
        print("\n\nНепредвиденная ошибка:")
        traceback.print_exc()
        log_error_traceback("main loop")
    finally:
        keepalive_stop.set()
        with _seen_ids_lock:
            save_seen_ids(seen_ids)
        save_workbooks_now()
        print("\n" + "=" * 70)
        print(f"Проверено объявлений за сессию: {total_checked}")
        print(f"Ошибок при обработке: {total_errors}")
        print(f"Найдено релевантных IT-тендеров всего: {total_found_it}")
        print(f"  из них по Астане (отдельный файл): {total_found_astana}")
        print(f"  остальные (основной файл): {total_found_it - total_found_astana}")
        print(f"Основной файл: {os.path.abspath(OUTPUT_XLSX)}")
        print(f"Файл по Астане: {os.path.abspath(OUTPUT_XLSX_ASTANA)}")
        if total_errors:
            print(f"Подробные traceback'и ошибок: {os.path.abspath(ERROR_LOG_FILE)}")
        print("=" * 70)


if __name__ == "__main__":
    main()