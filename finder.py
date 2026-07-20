# -*- coding: utf-8 -*-
r"""
finder.py — автоматизированный парсер IT-тендеров goszakup.gov.kz
(однофайловая версия).

ПЕРЕД ЗАПУСКОМ:
    pip install selenium beautifulsoup4 requests openpyxl pdfplumber python-docx xlrd

    Скачайте msedgedriver.exe под версию вашего Edge:
    https://developer.microsoft.com/microsoft-edge/tools/webdriver/
    Положите его рядом с finder.py (или укажите путь в EDGE_DRIVER_PATH ниже).

КАК ЗАПУСКАТЬ:
    1) Просто запустите:  python finder.py

       Скрипт сам откроет ОТДЕЛЬНОЕ окно Edge со своим собственным профилем
       (папка AUTOMATION_PROFILE_DIR рядом со скриптом, по умолчанию
       ".edge_automation_profile"). Это НЕ ваш обычный профиль Edge — он не
       использует ваши обычные вкладки/куки/пароли и не конфликтует с уже
       открытым обычным Edge. Можно держать оба браузера открытыми
       одновременно.

    2) В открывшемся отдельном окне Edge: авторизуйтесь на goszakup.gov.kz
       (при первом запуске нужно будет залогиниться заново — это отдельный
       "чистый" профиль), настройте фильтры на странице "Поиск объявлений",
       нажмите "Найти".

       При следующих запусках логин обычно сохраняется в этом профиле,
       заново вводить пароль не потребуется.

    3) Вернитесь в консоль и нажмите Enter, когда список результатов открыт
       в этом отдельном окне Edge - дальше парсер работает автоматически.

Результат: goszakup_it_tenders.xlsx (дозаписывается построчно, без потери
данных при сбое) + checkpoint_seen_ids.json (для дедупликации между запусками)
+ errors.log (полные traceback'и всех ошибок обработки объявлений).
"""

import json
import os
import random
import re
import subprocess
import sys
import threading
import time
import traceback
import uuid
from concurrent.futures import ThreadPoolExecutor, as_completed
from urllib.parse import urlsplit, urlunsplit, parse_qsl, urlencode, urljoin

import requests
from bs4 import BeautifulSoup
from openpyxl import Workbook, load_workbook
from openpyxl.styles import Font, Alignment, PatternFill
from openpyxl.utils import get_column_letter
from selenium import webdriver
from selenium.webdriver.edge.options import Options as EdgeOptions
from selenium.webdriver.edge.service import Service as EdgeService


# ============================================================================
# ================================ НАСТРОЙКИ ================================
# ============================================================================

# --- Подключение к Edge ---
EDGE_DEBUG_ADDRESS = "127.0.0.1:9222"
EDGE_DRIVER_PATH = "msedgedriver.exe"          # путь к msedgedriver, если не в PATH
EDGE_BROWSER_PATH = r"C:\Program Files (x86)\Microsoft\Edge\Application\msedge.exe"
AUTO_LAUNCH_EDGE = True                         # True -> finder.py сам попробует открыть Edge с портом отладки

# Отдельный профиль Edge только для этого скрипта - НЕ ваш обычный профиль.
# Благодаря этому окно автоматизации не конфликтует с уже открытым обычным
# Edge и не имеет доступа к вашим обычным вкладкам/паролям/истории.
# Можно указать абсолютный путь, если хотите хранить его не рядом со скриптом.
AUTOMATION_PROFILE_DIR = os.path.join(
    os.path.dirname(os.path.abspath(__file__)), ".edge_automation_profile"
)

# --- Поведение парсера ---
COUNT_RECORD_PER_PAGE = 100
MAX_PAGES_SAFETY_LIMIT = 2000
REQUEST_DELAY_SEC = (0.3, 0.7)          # пауза между страницами списка (requests, не браузер - можно быстрее)
DETAIL_PAGE_DELAY_SEC = (0.15, 0.4)     # пауза перед запросом карточки тендера (на поток)
REQUEST_TIMEOUT_SEC = 20                # таймаут одного HTTP-запроса (requests)
DETAIL_WORKERS = 6                      # сколько карточек тендеров обрабатывать параллельно

SCAN_ATTACHED_DOCUMENTS = True          # если по названию+описанию IT не найден - проверять вложения (резервная проверка)
MAX_DOCS_TO_SCAN_IF_NO_MATCH = 3        # сколько вложений максимум скачивать на такую резервную проверку
MAX_DOCUMENT_SIZE_MB = 25
DOWNLOAD_TMP_DIR = "tmp_downloads"
DELETE_DOWNLOADED_DOCS_AFTER_SCAN = True

# Раз в сколько НОВЫХ найденных IT-тендеров сохранять Excel на диск
# (файл держится открытым в памяти между сохранениями - это сильно быстрее,
# чем пересохранять весь .xlsx на каждой строке).
XLSX_SAVE_EVERY_N_MATCHES = 5

# --- Вывод / отладка ---
OUTPUT_XLSX = "goszakup_it_tenders.xlsx"
CHECKPOINT_FILE = "checkpoint_seen_ids.json"
ERROR_LOG_FILE = "errors.log"           # сюда пишутся ПОЛНЫЕ traceback'и всех ошибок обработки
DEBUG_MODE = False                      # True -> сохранять HTML каждой страницы в debug_html/
DEBUG_DUMP_DIR = "debug_html"

# --- IT-ключевые слова для классификации (регистр не важен) ---
IT_KEYWORDS = [
    "информационн", "информатизац", "программн", "программное обеспечение",
    "лицензи", "автоматизац", "цифровизац", "цифровая трансформ",
    "разработка по", "разработка программного", "разработка веб", "разработка сайта",
    "разработка приложени", "разработка систем", "разработка модул",
    "внедрение", "сопровождение по", "сопровождение информационн", "техническая поддержка по",
    "техническая поддержка информационн", "доработка по", "доработка систем",
    "модернизация информационн", "модернизация систем", "апгрейд систем",
    "информационная система", "информационн систем", "аис",
    "автоматизированная информационная система", "геоинформационн", "гис ",
    "система электронного документооборота", "сэд", "crm", "erp", "бпм", "bpm",
    "система управления", "личный кабинет", "портал ", "веб-портал",
    "веб-разработка", "веб разработка", "веб-сайт", "веб-приложени", "мобильное приложени",
    "мобильн приложени", "ios приложени", "android приложени", "фронтенд", "бэкенд",
    "frontend", "backend",
    "api ", " api.", "интеграция систем", "интеграционная шина", "esb",
    "база данных", "баз данных", "субд", "sql", "postgresql", "mysql",
    "oracle database", "nosql", "хранилище данных", "дата-центр", "дата центр",
    "дашборд", "bi-систем", "аналитическая платформа",
    "кибербезопасност", "информационная безопасност", "иб систем", "защита информации",
    "антивирус", "средства защиты информации", "сзи ", "межсетевой экран", "firewall",
    "проникновени", "пентест", "аудит информационной безопасности", "шифрование",
    "криптограф", "электронная цифровая подпись", "эцп",
    "1с ", "1c ", "1с:", "1c:", "конфигурация 1с", "программист 1с",
    "серверное оборудование", "сервер (", "серверов", "сетевое оборудование",
    "вычислительн", "цод ", "виртуализац", "облачн", "cloud",
    "хостинг", "домен", "wi-fi", "лвс", "локальная вычислительная сеть",
    "структурированная кабельная система", "скс",
    "it-услуг", "ит-услуг", "ит услуги", "it услуги", "аутсорсинг ит", "аутстаффинг ит",
    "техническая поддержка пользователей", "служба поддержки", "helpdesk", "хелпдеск",
    "искусственный интеллект", "машинное обучение", "нейросет", "чат-бот", "чатбот",
    "телекоммуникацион", "программист", "системный администратор", "разработчик по",
]

EXCLUDE_FALSE_POSITIVE_PHRASES = [
    "информационный стенд",
    "информационная доска",
    "информационный щит",
    "информационная табличка",
]

KNOWN_STATUSES = [
    "Опубликовано (прием ценовых предложений)",
    "Опубликовано (ожидание проведения аукциона)",
    "Опубликовано (дополнение заявок)",
    "Опубликовано (прием заявок)",
    "Опубликовано (проверка заявок)",
    "Ожидание камерального контроля",
    "Ожидание проведения контроля качества",
    "Отправлено на контроль качества",
    "Пройдено контроль качества",
    "Рассмотрение дополнений заявок",
    "Формирование протокола промежуточных итогов",
    "Формирование протокола преддопуска",
    "Формирование протокола допуска",
    "Формирование протокола итогов",
    "Принятие решение о пересмотре итогов",
    "Принятие решение об исполнении уведомления",
    "Изменена документация",
    "Пересмотр итогов",
    "На обжаловании",
    "Приостановлено",
    "Отказ от закупки",
    "Рассмотрение заявок",
    "Опубликовано",
    "Завершено",
    "Отменено",
]

KNOWN_METHODS = [
    "Второй этап конкурса с использованием рамочного соглашения",
    "Первый этап конкурса с использованием рамочного соглашения",
    "Из одного источника по несостоявшимся закупкам не ГЗ",
    "Из одного источника по несостоявшимся закупкам",
    "Конкурс с предварительным квалификационным отбором",
    "Конкурс с применением двухэтапных процедур",
    "Конкурс с использованием рейтингово-балльной системы",
    "Конкурс с применением специального порядка",
    "Тендер с использованием рейтингово-балльной системы",
    "Тендер с применением особого порядка",
    "Государственные закупки с применением особого порядка",
    "Запрос ценовых предложений (не ГЗ new)",
    "Запрос ценовых предложений",
    "Открытый конкурс",
    "Аукцион (до 2022)",
    "Аукцион (с 2022)",
    "Аукцион (не ГЗ) с 2022 г",
    "Аукцион",
    "Закупка жилища",
    "Закупка по государственному социальному заказу",
    "Конкурс по строительству «под ключ»",
    "Тендер по строительству «под ключ»",
    "Тендер",
]

ANNOUNCE_LINK_RE = re.compile(r"/ru/announce/index/(\d+)")
NUMBER_ANNO_RE = re.compile(r"\b(\d{5,}-\d+)\b")
_STOP_LOOKAHEAD = (
    r"(?=\s*\d{4}-\d{2}-\d{2}|\s*[\d\s\u00a0]{1,20}[.,]\d{2}\s*(?:$|[А-ЯЁ])"
    # "Способ:" тут раньше не совпадал с реальной меткой на сайте
    # ("Способ закупки:"), из-за чего регулярка не останавливалась и
    # захватывала способ закупки внутрь поля "Заказчик". Теперь стоп-слово
    # ловит "Способ" с любым текстом до двоеточия (например "Способ закупки:").
    r"|Заказчик:|Организатор:|Способ[^:\n]{0,30}:|Статус:|$)"
)
ORGANIZER_RE = re.compile(r"Организатор:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
CUSTOMER_RE = re.compile(r"Заказчик:\s*(.+?)" + _STOP_LOOKAHEAD, re.S)
DOC_EXTENSIONS = (".pdf", ".doc", ".docx", ".xls", ".xlsx", ".rtf", ".zip", ".rar", ".txt")

HEADERS = [
    "№ объявления", "Название лота/объявления", "Заказчик",
    "Организатор", "Способ закупки", "Сумма, тг.", "Статус",
    "Начало приёма заявок", "Окончание приёма заявок", "Ссылка на объявление",
    "Совпавшие IT-ключевые слова", "Документы (названия и ссылки)",
    "Краткое описание/ТЗ (выдержка)",
]


# ============================================================================
# ============================ ВСПОМОГАТЕЛЬНОЕ ==============================
# ============================================================================

def polite_sleep(delay_range):
    time.sleep(random.uniform(*delay_range))


def dump_debug_html(name: str, html: str):
    if not DEBUG_MODE:
        return
    os.makedirs(DEBUG_DUMP_DIR, exist_ok=True)
    with open(os.path.join(DEBUG_DUMP_DIR, name), "w", encoding="utf-8") as f:
        f.write(html)


def log_error_traceback(header: str):
    """Пишет полный traceback последнего исключения в ERROR_LOG_FILE.
    Раньше в консоль выводился только str(e), из-за чего было невозможно
    понять, в какой именно строке кода реально падает ошибка - все ошибки
    выглядели одинаково ('NoneType' object has no attribute 'get'), хотя
    происходили в разных местах."""
    try:
        with open(ERROR_LOG_FILE, "a", encoding="utf-8") as f:
            f.write(f"\n--- {time.strftime('%Y-%m-%d %H:%M:%S')} | {header} ---\n")
            f.write(traceback.format_exc())
            f.write("\n")
    except Exception:
        pass


# ============================================================================
# ============================== БРАУЗЕР =====================================
# ============================================================================

def _edge_debug_port_alive() -> bool:
    try:
        r = requests.get(f"http://{EDGE_DEBUG_ADDRESS}/json/version", timeout=1.5)
        return r.ok
    except Exception:
        return False


def try_auto_launch_edge():
    """Открывает ОТДЕЛЬНОЕ окно Edge со своим профилем (AUTOMATION_PROFILE_DIR).

    Это не обычный профиль пользователя, поэтому новый процесс msedge.exe
    не блокируется уже открытым "обычным" Edge и не мешает ему - оба могут
    работать одновременно. Логин на goszakup.gov.kz нужно будет выполнить
    в этом отдельном окне (один раз - профиль на диске сохранит сессию
    между запусками).
    """
    if _edge_debug_port_alive():
        print("Порт отладки уже отвечает - использую уже открытое окно автоматизации.")
        return

    if not AUTO_LAUNCH_EDGE:
        print("[!] AUTO_LAUNCH_EDGE выключен. Откройте Edge вручную с флагом")
        print(f"    --remote-debugging-port={EDGE_DEBUG_ADDRESS.split(':')[-1]}")
        print(f"    --user-data-dir=\"{AUTOMATION_PROFILE_DIR}\"")
        return

    if not os.path.exists(EDGE_BROWSER_PATH):
        print(f"[!] Не нашёл Edge по пути {EDGE_BROWSER_PATH}. "
              f"Откройте браузер вручную с флагом --remote-debugging-port=9222 "
              f"и --user-data-dir=\"{AUTOMATION_PROFILE_DIR}\".")
        return

    os.makedirs(AUTOMATION_PROFILE_DIR, exist_ok=True)
    port = EDGE_DEBUG_ADDRESS.split(":")[-1]
    print("Открываю отдельное окно Edge для автоматизации (свой профиль,")
    print(f"не связан с вашим обычным Edge): {AUTOMATION_PROFILE_DIR}")
    subprocess.Popen([
        EDGE_BROWSER_PATH,
        f"--remote-debugging-port={port}",
        f"--user-data-dir={AUTOMATION_PROFILE_DIR}",
        "--profile-directory=Default",
        "--no-first-run",
        "--no-default-browser-check",
        "https://goszakup.gov.kz/ru/announce/index",
    ])

    # Ждём, пока порт отладки реально откроется (до ~60 секунд).
    for i in range(120):
        time.sleep(0.5)
        if _edge_debug_port_alive():
            print("Отдельное окно Edge запущено, порт отладки открылся успешно.")
            return
        if i > 0 and i % 10 == 0:
            print(f"    ...жду запуск Edge ({i // 2} сек)")

    # Так как профиль отдельный и свой, единственная реальная причина сбоя
    # здесь - сам файл профиля уже заблокирован ДРУГИМ процессом finder.py /
    # тем же автоматизационным окном, оставшимся в фоне.
    print("\n" + "!" * 70)
    print("Порт отладки не открылся. Возможно, окно автоматизации уже было")
    print("открыто ранее и осталось висеть в фоне (проверьте диспетчер задач")
    print("на процесс msedge.exe, использующий папку профиля):")
    print(f"    {AUTOMATION_PROFILE_DIR}")
    print("\nЧТО СДЕЛАТЬ:")
    print("  1) Закройте все окна Edge, у которых в заголовке/адресной строке")
    print("     видно, что это отдельный профиль автоматизации, либо выполните")
    print("       taskkill /IM msedge.exe /F")
    print("     (это не тронет ваш обычный Edge, если он уже был закрыт при")
    print("      выполнении этой команды - она закрывает ВСЕ окна msedge.exe,")
    print("      так что обычный Edge тоже закроется и его нужно будет")
    print("      открыть заново отдельно).")
    print("  2) Запустите finder.py снова.")
    print("!" * 70 + "\n")


def attach_to_running_edge() -> webdriver.Edge:
    options = EdgeOptions()
    options.add_experimental_option("debuggerAddress", EDGE_DEBUG_ADDRESS)
    service = EdgeService(executable_path=EDGE_DRIVER_PATH)
    return webdriver.Edge(service=service, options=options)


def get_cookies_dict(driver: webdriver.Edge) -> dict:
    return {c["name"]: c["value"] for c in driver.get_cookies()}


# ============================================================================
# ======================== ПАРСИНГ СТРАНИЦЫ СПИСКА ==========================
# ============================================================================

def build_page_url(base_url: str, page: int) -> str:
    parts = urlsplit(base_url)
    query = dict(parse_qsl(parts.query))
    query["page"] = str(page)
    query["count_record"] = str(COUNT_RECORD_PER_PAGE)
    return urlunsplit((parts.scheme, parts.netloc, parts.path, urlencode(query), ""))


def _find_status(row_text: str) -> str:
    for status in KNOWN_STATUSES:
        if status in row_text:
            return status
    return ""


def _find_method(row_text: str) -> str:
    for method in KNOWN_METHODS:
        if method in row_text:
            return method
    return ""


def _clean_amount(raw: str) -> str:
    """Иногда при схлопывании строки таблицы в одну строку текста соседняя
    ячейка (например, короткий числовой код) слипается с суммой без
    настоящего разделителя, и регулярка захватывает всё как одно число
    (например, "00 150 000.00" вместо "150 000.00"). Отрезаем такой
    мусорный ведущий фрагмент из одних нулей."""
    if not raw:
        return raw
    s = raw.strip()
    m = re.match(r"^0{1,3}[\s\u00a0]+(\d{1,3}(?:[\s\u00a0]\d{3})*[.,]\d{2})$", s)
    if m:
        return m.group(1)
    return s


def _truncate_leaked_text(text: str, *markers: str) -> str:
    """ORGANIZER_RE/CUSTOMER_RE останавливаются по текстовому лейблу
    "Способ:"/"Статус:" - но, судя по всему, сайт способ закупки выводит
    БЕЗ отдельного лейбла перед ним, просто сплошным текстом сразу после
    заказчика/организатора. Из-за этого способ закупки/статус иногда
    целиком попадает внутрь поля "Заказчик" или "Организатор". Способ и
    статус мы и так независимо находим через _find_method/_find_status,
    так что просто обрезаем поле по первому вхождению уже найденного
    текста способа/статуса, если он туда просочился."""
    if not text:
        return text
    cut_at = len(text)
    for marker in markers:
        if marker:
            idx = text.find(marker)
            if idx != -1:
                cut_at = min(cut_at, idx)
    return text[:cut_at].strip(" ,.-\u00a0")


def parse_list_page(html: str) -> list:
    soup = BeautifulSoup(html, "html.parser")
    results = []
    seen_ids_on_page = set()

    for a_tag in soup.find_all("a", href=ANNOUNCE_LINK_RE):
        match = ANNOUNCE_LINK_RE.search(a_tag.get("href", ""))
        if not match:
            continue
        tender_id = match.group(1)
        if tender_id in seen_ids_on_page:
            continue
        seen_ids_on_page.add(tender_id)

        title = a_tag.get_text(strip=True)
        if not title:
            continue

        row = a_tag.find_parent("tr") or a_tag.find_parent(["div", "li"])
        row_text = row.get_text(" ", strip=True) if row else a_tag.get_text(" ", strip=True)

        number_anno_match = NUMBER_ANNO_RE.search(row_text)
        number_anno = number_anno_match.group(1) if number_anno_match else tender_id

        organizer_match = ORGANIZER_RE.search(row_text)
        organizer_name = organizer_match.group(1).strip() if organizer_match else ""

        customer_match = CUSTOMER_RE.search(row_text)
        customer_name = customer_match.group(1).strip() if customer_match else ""

        status = _find_status(row_text)
        method = _find_method(row_text)

        organizer_name = _truncate_leaked_text(organizer_name, method, status)
        customer_name = _truncate_leaked_text(customer_name, method, status)

        amount = ""
        amount_matches = re.findall(r"[\d\s\u00a0]{1,20}[.,]\d{2}", row_text)
        if amount_matches:
            amount = _clean_amount(amount_matches[-1].strip())

        dates = re.findall(r"\d{4}-\d{2}-\d{2}\s+\d{2}:\d{2}(?::\d{2})?", row_text)
        start_date = dates[0] if len(dates) >= 1 else ""
        end_date = dates[1] if len(dates) >= 2 else ""

        results.append({
            "id": tender_id,
            "number_anno": number_anno,
            "title": title,
            "url": f"https://goszakup.gov.kz/ru/announce/index/{tender_id}",
            "organizer_name": organizer_name,
            "customer_name": customer_name,
            "status": status,
            "method": method,
            "amount": amount,
            "start_date": start_date,
            "end_date": end_date,
        })

    return results


# ============================================================================
# ======================= ПАРСИНГ ДЕТАЛЬНОЙ СТРАНИЦЫ =========================
# ============================================================================

# Теги, которые почти никогда не относятся к содержанию самого объявления,
# а являются "обвязкой" сайта (меню, футер, выпадающие списки фильтров и
# категорий и т.п.) - именно из-за них раньше в текст объявления попадали
# названия ВСЕХ категорий закупок (включая IT), и почти всё определялось
# как "IT-тендер".
_CHROME_TAGS = ("script", "style", "nav", "header", "footer", "select", "option", "noscript", "iframe")
_CHROME_HINT_RE = re.compile(
    r"(menu|navbar|nav-|footer|header|sidebar|breadcrumb|filter|search-form|"
    r"cookie|lang-switch|top-panel|topline|langs)",
    re.I,
)


def _strip_page_chrome(soup: BeautifulSoup) -> None:
    for tag_name in _CHROME_TAGS:
        for tag in soup.find_all(tag_name):
            tag.decompose()
    # Второй проход - по классам/id, которые явно указывают на меню/футер/фильтры,
    # даже если это не семантические теги <nav>/<footer>, а div/ul и т.п.
    #
    # ВАЖНО: soup.find_all(True) возвращает список ВСЕХ тегов, вычисленный
    # один раз в начале. Но decompose() внутри этого же цикла удаляет не
    # только сам тег, но и всех его потомков - у них при этом .attrs
    # становится None. Если такой потомок ещё встретится дальше в этом же
    # предвычисленном списке, вызов tag.get(...) упадёт с AttributeError
    # ('NoneType' object has no attribute 'get'), т.к. .get() внутри bs4
    # обращается к self.attrs.get(...). Поэтому пропускаем уже
    # декомпозированные (в т.ч. косвенно, как потомки) теги.
    for tag in soup.find_all(True):
        if getattr(tag, "decomposed", False):
            continue
        cls = " ".join(tag.get("class", []) or [])
        tag_id = tag.get("id", "") or ""
        if (cls and _CHROME_HINT_RE.search(cls)) or (tag_id and _CHROME_HINT_RE.search(tag_id)):
            tag.decompose()


def parse_detail_page(html: str, page_url: str) -> dict:
    soup = BeautifulSoup(html, "html.parser")

    # Ссылки на документы собираем ДО чистки - вложения иногда лежат в блоках,
    # которые могут случайно попасть под чистку по классам.
    documents = []
    for a_tag in soup.find_all("a", href=True):
        href = a_tag["href"]
        href_lower = href.lower()
        if href_lower.endswith(DOC_EXTENSIONS) or "download" in href_lower or "/file/" in href_lower:
            full_url = urljoin(page_url, href)
            doc_name = a_tag.get_text(strip=True) or full_url.split("/")[-1]
            documents.append({"name": doc_name, "url": full_url})

    _strip_page_chrome(soup)

    main_container = (
        soup.find("div", attrs={"id": re.compile(r"content", re.I)})
        or soup.find("main")
        or soup.find("table")
        or soup.find("body")
    )
    description_text = main_container.get_text("\n", strip=True) if main_container else ""

    customer_bin = ""
    bin_match = re.search(r"БИН[:\s]*([0-9]{12})", description_text)
    if bin_match:
        customer_bin = bin_match.group(1)

    return {"description": description_text, "documents": documents, "customer_bin": customer_bin}


# ============================================================================
# ==================== СКАЧИВАНИЕ И ЧТЕНИЕ ДОКУМЕНТОВ ========================
# ============================================================================

def _safe_filename(url: str) -> str:
    name = url.split("?")[0].rstrip("/").split("/")[-1]
    name = re.sub(r"[^\w\.\-]+", "_", name)
    if not name:
        name = uuid.uuid4().hex
    return f"{uuid.uuid4().hex[:8]}_{name}"


def download_file(url: str, session: requests.Session, cookies: dict):
    os.makedirs(DOWNLOAD_TMP_DIR, exist_ok=True)
    local_path = os.path.join(DOWNLOAD_TMP_DIR, _safe_filename(url))
    try:
        with session.get(url, cookies=cookies, stream=True, timeout=30) as resp:
            resp.raise_for_status()
            total = 0
            max_bytes = MAX_DOCUMENT_SIZE_MB * 1024 * 1024
            with open(local_path, "wb") as f:
                for chunk in resp.iter_content(chunk_size=8192):
                    total += len(chunk)
                    if total > max_bytes:
                        f.close()
                        os.remove(local_path)
                        return None
                    f.write(chunk)
        return local_path
    except Exception:
        if os.path.exists(local_path):
            try:
                os.remove(local_path)
            except OSError:
                pass
        return None


def _extract_pdf(path):
    try:
        import pdfplumber
        parts = []
        with pdfplumber.open(path) as pdf:
            for page in pdf.pages:
                parts.append(page.extract_text() or "")
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_docx(path):
    try:
        import docx
        d = docx.Document(path)
        parts = [p.text for p in d.paragraphs]
        for table in d.tables:
            for row in table.rows:
                for cell in row.cells:
                    parts.append(cell.text)
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_doc_legacy(path):
    try:
        import textract
        return textract.process(path).decode("utf-8", errors="ignore")
    except Exception:
        return ""


def _extract_xlsx(path):
    try:
        import openpyxl as _openpyxl
        wb = _openpyxl.load_workbook(path, data_only=True, read_only=True)
        parts = []
        for ws in wb.worksheets:
            for row in ws.iter_rows(values_only=True):
                for cell in row:
                    if cell is not None:
                        parts.append(str(cell))
        return "\n".join(parts)
    except Exception:
        return ""


def _extract_xls_legacy(path):
    try:
        import xlrd
        wb = xlrd.open_workbook(path)
        parts = []
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for cell in sheet.row(row_idx):
                    parts.append(str(cell.value))
        return "\n".join(parts)
    except Exception:
        return ""


def extract_text_from_file(path: str) -> str:
    ext = os.path.splitext(path)[1].lower()
    try:
        if ext == ".pdf":
            return _extract_pdf(path)
        elif ext == ".docx":
            return _extract_docx(path)
        elif ext == ".doc":
            return _extract_doc_legacy(path)
        elif ext in (".xlsx", ".xlsm"):
            return _extract_xlsx(path)
        elif ext == ".xls":
            return _extract_xls_legacy(path)
        elif ext in (".txt", ".rtf"):
            with open(path, "r", encoding="utf-8", errors="ignore") as f:
                return f.read()
        return ""
    except Exception:
        return ""


def get_document_text(url: str, session: requests.Session, cookies: dict) -> str:
    path = download_file(url, session, cookies)
    if not path:
        return ""
    try:
        return extract_text_from_file(path)
    finally:
        if DELETE_DOWNLOADED_DOCS_AFTER_SCAN and os.path.exists(path):
            try:
                os.remove(path)
            except OSError:
                pass


# ============================================================================
# ================================ КЛАССИФИКАЦИЯ =============================
# ============================================================================

def _normalize(text: str) -> str:
    return (text or "").lower().replace("ё", "е")


def classify_is_it(*text_fragments):
    combined = _normalize(" \n ".join(t for t in text_fragments if t))
    for phrase in EXCLUDE_FALSE_POSITIVE_PHRASES:
        combined = combined.replace(_normalize(phrase), "")

    matched = [kw.strip() for kw in IT_KEYWORDS if _normalize(kw) in combined]
    return len(matched) > 0, matched


# ============================================================================
# ================================ EXCEL =====================================
# ============================================================================

_xlsx_lock = threading.Lock()
_workbook = None       # держим книгу открытой в памяти на весь запуск
_worksheet = None
_rows_since_save = 0


def _build_new_workbook():
    wb = Workbook()
    ws = wb.active
    ws.title = "IT-тендеры"
    ws.append(HEADERS)
    header_fill = PatternFill(start_color="1F4E78", end_color="1F4E78", fill_type="solid")
    for col_idx in range(1, len(HEADERS) + 1):
        cell = ws.cell(row=1, column=col_idx)
        cell.font = Font(bold=True, color="FFFFFF")
        cell.fill = header_fill
        cell.alignment = Alignment(vertical="center", wrap_text=True)
    widths = [16, 45, 35, 30, 22, 16, 24, 18, 18, 30, 30, 45, 50]
    for i, w in enumerate(widths, start=1):
        ws.column_dimensions[get_column_letter(i)].width = w
    ws.freeze_panes = "A2"
    return wb, ws


def init_workbook_for_run():
    """Загружает существующий Excel в память (если есть) либо создаёт новый.
    Вызывается один раз в начале main() - дальше книга не перечитывается
    с диска на каждой строке, что радикально быстрее на больших объёмах."""
    global _workbook, _worksheet
    if os.path.exists(OUTPUT_XLSX):
        _workbook = load_workbook(OUTPUT_XLSX)
        _worksheet = _workbook["IT-тендеры"] if "IT-тендеры" in _workbook.sheetnames else _workbook.active
    else:
        _workbook, _worksheet = _build_new_workbook()
        _workbook.save(OUTPUT_XLSX)


def save_workbook_now():
    global _rows_since_save
    with _xlsx_lock:
        if _workbook is not None:
            _workbook.save(OUTPUT_XLSX)
            _rows_since_save = 0


def append_row(row: dict):
    """Потокобезопасно добавляет строку в открытую в памяти книгу.
    Сохранение на диск происходит не на каждой строке, а раз в
    XLSX_SAVE_EVERY_N_MATCHES найденных тендеров (плюс всегда в конце
    и при остановке/ошибке - см. main())."""
    global _rows_since_save

    values = [
        row.get("number_anno", ""),
        row.get("title", ""),
        row.get("customer_name", ""),
        row.get("organizer_name", ""),
        row.get("method", ""),
        row.get("amount", ""),
        row.get("status", ""),
        row.get("start_date", ""),
        row.get("end_date", ""),
        row.get("url", ""),
        ", ".join(row.get("matched_keywords", [])),
        "; ".join(row.get("documents", [])),
        (row.get("description", "") or "")[:1500],
    ]

    with _xlsx_lock:
        _worksheet.append(values)
        last_row = _worksheet.max_row
        url_cell = _worksheet.cell(row=last_row, column=HEADERS.index("Ссылка на объявление") + 1)
        if row.get("url"):
            url_cell.hyperlink = row["url"]
            url_cell.font = Font(color="0563C1", underline="single")
        _rows_since_save += 1
        need_save = _rows_since_save >= XLSX_SAVE_EVERY_N_MATCHES

    if need_save:
        save_workbook_now()


# ============================================================================
# ============================== ЧЕКПОИНТ ====================================
# ============================================================================

_seen_ids_lock = threading.Lock()


def load_seen_ids() -> set:
    if not os.path.exists(CHECKPOINT_FILE):
        return set()
    try:
        with open(CHECKPOINT_FILE, "r", encoding="utf-8") as f:
            data = json.load(f)
        return set(data.get("seen_ids", []))
    except Exception:
        return set()


def save_seen_ids(seen_ids: set):
    with open(CHECKPOINT_FILE, "w", encoding="utf-8") as f:
        json.dump({"seen_ids": sorted(seen_ids)}, f, ensure_ascii=False, indent=2)


# ============================================================================
# ============================== ОСНОВНОЙ ЦИКЛ ===============================
# ============================================================================

# ---------------------------------------------------------------------------
# ВАЖНО: requests.Session (и его cookie jar) не рассчитан на одновременное
# использование из нескольких потоков (несколько потоков одновременно
# читают/пишут одну и ту же cookie jar на каждый HTTP-запрос). Раньше все
# DETAIL_WORKERS потоков делили ОДИН session, что и приводило к массовым
# случайным ошибкам вида "AttributeError: 'NoneType' object has no
# attribute 'get'" (гонка внутри http.cookiejar/urllib3).
#
# Решение: у каждого потока - своя requests.Session (thread-local),
# куки в неё каждый раз копируются (дешёвая операция) из "главной"
# сессии, авторизованной через cookies браузера. Сама "главная" сессия
# при этом мутируется только из главного потока, между страницами -
# то есть в момент, когда ThreadPoolExecutor ещё/уже не работает.
# ---------------------------------------------------------------------------
_thread_local = threading.local()


def get_thread_session(base_session: requests.Session) -> requests.Session:
    sess = getattr(_thread_local, "session", None)
    if sess is None:
        sess = requests.Session()
        sess.headers.update(dict(base_session.headers))
        _thread_local.session = sess
    # Синхронизируем куки из главной сессии (безопасно: главный поток
    # не мутирует base_session.cookies, пока работает executor).
    sess.cookies.update(base_session.cookies.get_dict())
    return sess


def fetch_html(session: requests.Session, url: str, attempts: int = 6) -> str:
    """GET с повторными попытками. Используется вместо Selenium для карточек
    тендеров и страниц списка - это на порядок быстрее, т.к. не требует
    полноценного открытия страницы в браузере, и позволяет работать в
    несколько потоков одновременно.

    v3bl.goszakup.gov.kz - судя по всему, бета/тестовый поддомен и иногда
    подвисает на несколько секунд-минут (ConnectTimeout), поэтому здесь
    делается несколько попыток с постепенно растущей паузой (5, 10, 15,
    20, 25 сек), прежде чем сдаться - раньше пауза была всего 0.5-1.5 сек,
    чего для реальных сетевых заминок недостаточно."""
    last_exc = None
    for attempt in range(attempts):
        try:
            resp = session.get(url, timeout=REQUEST_TIMEOUT_SEC)
            resp.raise_for_status()
            resp.encoding = resp.encoding or "utf-8"
            return resp.text
        except Exception as e:
            last_exc = e
            if attempt < attempts - 1:
                wait_sec = 5 * (attempt + 1)
                print(f"    [!] Сетевая заминка ({e.__class__.__name__}), "
                      f"жду {wait_sec} сек и пробую снова "
                      f"(попытка {attempt + 2}/{attempts})...")
                time.sleep(wait_sec)
    raise last_exc


def process_tender(session: requests.Session, row: dict):
    # Собственная сессия для этого потока - см. комментарий выше.
    thread_session = get_thread_session(session)

    polite_sleep(DETAIL_PAGE_DELAY_SEC)
    html = fetch_html(thread_session, row["url"])
    dump_debug_html(f"detail_{row['id']}.html", html)

    detail = parse_detail_page(html, row["url"])
    doc_display = [f"{d['name']} ({d['url']})" for d in detail["documents"]]

    # Сначала - быстрая проверка по названию и (уже очищенному от меню/футера)
    # описанию объявления. Для подавляющего большинства НЕ-IT тендеров этого
    # достаточно, и мы вообще не тратим время на скачивание вложений.
    is_it, matched_keywords = classify_is_it(row.get("title", ""), detail.get("description", ""))

    # Резервная проверка: если по названию/описанию IT не найден, но есть
    # прикреплённые документы - заглядываем в них (вдруг IT-требование есть
    # только в самом ТЗ). Ограничено MAX_DOCS_TO_SCAN_IF_NO_MATCH файлами,
    # чтобы не тормозить на объявлениях, где документов много, а к IT они
    # заведомо не относятся.
    if not is_it and SCAN_ATTACHED_DOCUMENTS and detail["documents"]:
        cookies = thread_session.cookies.get_dict()
        for doc in detail["documents"][:MAX_DOCS_TO_SCAN_IF_NO_MATCH]:
            text = get_document_text(doc["url"], thread_session, cookies)
            if text:
                is_it, matched_keywords = classify_is_it(text)
                if is_it:
                    break

    if not is_it:
        return None

    return {
        "number_anno": row.get("number_anno", ""),
        "title": row.get("title", ""),
        "customer_name": row.get("customer_name") or row.get("organizer_name", ""),
        "customer_bin": detail.get("customer_bin", ""),
        "organizer_name": row.get("organizer_name", ""),
        "method": row.get("method", ""),
        "amount": row.get("amount", ""),
        "status": row.get("status", ""),
        "start_date": row.get("start_date", ""),
        "end_date": row.get("end_date", ""),
        "url": row.get("url", ""),
        "matched_keywords": matched_keywords,
        "documents": doc_display,
        "description": detail.get("description", ""),
    }


FINDER_VERSION = "2024-07-v8 (keep-alive вкладки против разрыва сессии + починен сдвиг колонок Excel + чистка суммы + фикс утечки способа закупки в Заказчика)"


KEEPALIVE_INTERVAL_SEC = 180  # как часто "трогаем" настоящую вкладку Edge, чтобы сервер не считал сессию неактивной


def browser_keepalive_loop(driver, stop_event: threading.Event):
    """Периодически перезагружает ТУ ЖЕ вкладку Edge, на которой уже стоит
    авторизованная сессия (обычный GET на её текущий URL - не F5/refresh,
    поэтому не может всплыть диалог "повторно отправить данные формы").

    Смысл: наши фоновые requests-запросы используют те же куки, что и эта
    вкладка, но сервер, судя по всему, не засчитывает их как активность
    реального пользователя и всё равно завершает сессию по таймауту
    бездействия. Реальная навигация в самой вкладке браузера - это
    настоящая активность с точки зрения сервера, и должна сбрасывать этот
    таймер бездействия."""
    while not stop_event.wait(KEEPALIVE_INTERVAL_SEC):
        try:
            current_url = driver.current_url
            driver.get(current_url)
        except Exception:
            # Не критично - просто попробуем в следующий раз через
            # KEEPALIVE_INTERVAL_SEC. Ошибку не логируем, чтобы не забивать
            # errors.log второстепенными сбоями фонового потока.
            pass


def find_and_switch_to_goszakup_tab(driver) -> str:
    """Edge может открыть несколько вкладок (например, стартовую страницу
    MSN плюс ту, что мы указали при запуске) - Selenium по умолчанию
    остаётся на первой из них, а не обязательно на той, где пользователь
    искал тендеры. Эта функция ищет среди ВСЕХ открытых вкладок ту, что
    относится к goszakup.gov.kz, и переключает драйвер на неё.
    Возвращает URL найденной вкладки или None, если не нашла."""
    found_url = None
    found_handle = None
    for handle in driver.window_handles:
        try:
            driver.switch_to.window(handle)
            url = driver.current_url
        except Exception:
            continue
        if "goszakup.gov.kz" in url:
            found_url = url
            found_handle = handle
            # Предпочитаем вкладку именно со страницей поиска/результатов,
            # если открыто несколько вкладок goszakup - но в целом любая
            # вкладка этого домена лучше, чем ничего.
            if "search/announce" in url or "announce" in url:
                break
    if found_handle is not None:
        driver.switch_to.window(found_handle)
    return found_url


def main():
    print("=" * 70)
    print("Парсер IT-тендеров goszakup.gov.kz (finder.py)")
    print(f"Версия скрипта: {FINDER_VERSION}")
    print("=" * 70)

    try_auto_launch_edge()

    print("\nПодключаюсь к отдельному окну Edge (remote debugging)...")
    try:
        driver = attach_to_running_edge()
    except Exception as e:
        print(f"\nНе удалось подключиться к Edge: {e}")
        print("\nПроверьте:")
        print("  1) отдельное окно Edge для автоматизации реально открылось")
        print(f"     (профиль: {AUTOMATION_PROFILE_DIR});")
        print("  2) версия msedgedriver.exe совпадает с версией вашего Edge")
        print("     (Edge -> Настройки -> О Microsoft Edge, сверить номер версии;")
        print("      скачать драйвер: https://developer.microsoft.com/microsoft-edge/tools/webdriver/).")
        sys.exit(1)

    print("Подключение успешно.\n")
    print("Убедитесь, что в ОТДЕЛЬНОМ окне Edge (не в вашем обычном браузере!) вы:")
    print("  1) авторизовались на goszakup.gov.kz;")
    print("  2) настроили нужные фильтры на странице 'Поиск объявлений';")
    print("  3) нажали 'Найти' и видите список результатов.\n")
    input("Когда всё готово - нажмите Enter здесь, чтобы начать сбор данных...")

    base_url = find_and_switch_to_goszakup_tab(driver)
    attempts_left = 3
    while base_url is None and attempts_left > 0:
        print("\n[!] Не нашёл среди открытых вкладок ни одной с goszakup.gov.kz.")
        print("    Убедитесь, что вкладка с результатами поиска действительно")
        print("    открыта в ЭТОМ ЖЕ окне Edge (том, что запустил finder.py),")
        print("    и что вы уже нажали 'Найти'.")
        input("    Когда исправите - нажмите Enter, чтобы попробовать снова...")
        base_url = find_and_switch_to_goszakup_tab(driver)
        attempts_left -= 1

    if base_url is None:
        print("\nТак и не нашёл вкладку goszakup.gov.kz. Останавливаюсь.")
        sys.exit(1)

    print(f"\nБазовый URL с вашими фильтрами: {base_url}")
    if "search/announce" not in base_url and "announce" not in base_url:
        print("[!] Похоже, это не страница результатов поиска объявлений -")
        print("    проверьте, что перед запуском вы дошли до списка тендеров.")

    seen_ids = load_seen_ids()
    print(f"Уже обработано в предыдущих запусках: {len(seen_ids)} объявлений.\n")

    # Фоновый поток, который периодически "трогает" настоящую вкладку Edge,
    # чтобы сервер не отключал сессию по таймауту бездействия (см.
    # browser_keepalive_loop выше).
    keepalive_stop = threading.Event()
    keepalive_thread = threading.Thread(
        target=browser_keepalive_loop, args=(driver, keepalive_stop), daemon=True
    )
    keepalive_thread.start()

    init_workbook_for_run()

    # Дальше список и карточки тендеров забираем через обычный requests,
    # используя куки из уже залогиненного отдельного окна Edge - это
    # значительно быстрее, чем гонять настоящий браузер по каждой странице.
    #
    # Это "главная" сессия: список страниц (single-threaded) читает её
    # напрямую, а каждый воркер-поток в ThreadPoolExecutor берёт себе
    # ОТДЕЛЬНУЮ сессию через get_thread_session(session) (см. process_tender).
    session = requests.Session()
    session.headers.update({
        "User-Agent": driver.execute_script("return navigator.userAgent;"),
        "Referer": base_url,
        "Accept-Language": "ru-RU,ru;q=0.9",
    })
    session.cookies.update(get_cookies_dict(driver))

    total_found_it = 0
    total_checked = 0
    total_errors = 0
    page = 1

    try:
        while page <= MAX_PAGES_SAFETY_LIMIT:
            page_url = build_page_url(base_url, page)
            print(f"--- Страница {page}: {page_url}")

            # Куки на всякий случай обновляем каждую страницу списка (дёшево,
            # без навигации - просто читаем текущее состояние окна логина),
            # чтобы устойчиво работать, даже если сайт периодически ротирует
            # токены сессии. Это происходит ДО запуска пула потоков на эту
            # страницу, поэтому конкурентной записи в session.cookies с
            # чтением из воркер-потоков здесь не возникает.
            session.cookies.update(get_cookies_dict(driver))

            # Ждём, пока страница списка реально не загрузится - без ограничения
            # по числу попыток. Пауза между попытками растёт (10с, 20с, 30с...),
            # но не больше 2 минут, чтобы не ждать бесконечно долго между
            # проверками при действительно долгой недоступности сайта.
            html = None
            attempt_num = 0
            while html is None:
                attempt_num += 1
                try:
                    html = fetch_html(session, page_url)
                except Exception as e:
                    log_error_traceback(f"list page fetch, page={page}, url={page_url}, attempt={attempt_num}")
                    wait_sec = min(10 * attempt_num, 120)
                    print(f"    [!] Страница {page} долго не загружается "
                          f"({e.__class__.__name__}), попытка {attempt_num}. "
                          f"Жду {wait_sec} сек и пробую снова...")
                    time.sleep(wait_sec)
                    # На всякий случай обновим куки перед повторной попыткой -
                    # вдруг сессия за это время истекла.
                    session.cookies.update(get_cookies_dict(driver))

            polite_sleep(REQUEST_DELAY_SEC)
            dump_debug_html(f"list_page_{page}.html", html)

            rows = parse_list_page(html)
            if not rows:
                empty_dump_path = os.path.abspath("empty_page_debug.html")
                try:
                    with open(empty_dump_path, "w", encoding="utf-8") as f:
                        f.write(html)
                except Exception:
                    pass
                print("Больше объявлений не найдено (либо сайт отдал страницу без "
                      "объявлений по другой причине - истекшая сессия, сброшенный "
                      "фильтр, капча и т.п.). Останавливаюсь.")
                print(f"HTML этой страницы сохранён в: {empty_dump_path}")
                print("Откройте этот файл в браузере/текстовом редакторе и сверьте:")
                print("  - если там видна нормальная страница поиска, но БЕЗ строк "
                      "с объявлениями - значит объявлений по вашему фильтру "
                      "действительно больше нет;")
                print("  - если там форма логина, капча, сообщение об ошибке или "
                      "просто исходная (нефильтрованная) страница поиска - значит "
                      "сессия/фильтр сбросились, и нужно заново пройти шаги 1-3 в "
                      "отдельном окне Edge и перезапустить скрипт.")
                break

            new_rows = [r for r in rows if r["id"] not in seen_ids]
            print(f"    Найдено на странице: {len(rows)}, новых: {len(new_rows)}")

            # Карточки тендеров этой страницы обрабатываем параллельно
            # (DETAIL_WORKERS потоков) - самая медленная часть (сеть,
            # разбор HTML, иногда скачивание вложений) больше не блокирует
            # друг друга по очереди.
            found_on_this_page = 0
            with ThreadPoolExecutor(max_workers=DETAIL_WORKERS) as executor:
                future_to_row = {executor.submit(process_tender, session, row): row for row in new_rows}
                for future in as_completed(future_to_row):
                    row = future_to_row[future]
                    total_checked += 1
                    try:
                        result = future.result()
                    except Exception as e:
                        total_errors += 1
                        print(f"    [!] Ошибка при обработке {row['id']}: {e}  (подробности в {ERROR_LOG_FILE})")
                        log_error_traceback(f"tender_id={row['id']} url={row.get('url', '')}")
                        result = None

                    with _seen_ids_lock:
                        seen_ids.add(row["id"])

                    if result:
                        append_row(result)
                        total_found_it += 1
                        found_on_this_page += 1
                        print(f"    [IT] {row['number_anno']}: {row['title'][:70]}")

                    if total_checked % 20 == 0:
                        with _seen_ids_lock:
                            save_seen_ids(seen_ids)

            if new_rows and found_on_this_page == 0:
                print(f"    На странице {page} IT-тендеров не найдено.")

            with _seen_ids_lock:
                save_seen_ids(seen_ids)
            page += 1

    except KeyboardInterrupt:
        print("\n\nОстановлено пользователем (Ctrl+C). Сохраняю прогресс...")
    except Exception:
        print("\n\nНепредвиденная ошибка:")
        traceback.print_exc()
        log_error_traceback("main loop")
    finally:
        keepalive_stop.set()
        with _seen_ids_lock:
            save_seen_ids(seen_ids)
        save_workbook_now()
        print("\n" + "=" * 70)
        print(f"Проверено объявлений за сессию: {total_checked}")
        print(f"Ошибок при обработке: {total_errors}")
        print(f"Найдено релевантных IT-тендеров: {total_found_it}")
        print(f"Результат сохранён в: {os.path.abspath(OUTPUT_XLSX)}")
        if total_errors:
            print(f"Подробные traceback'и ошибок: {os.path.abspath(ERROR_LOG_FILE)}")
        print("=" * 70)


if __name__ == "__main__":
    main()